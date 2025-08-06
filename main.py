import os
import sys
import re
import json
from datetime import datetime
import time
import traceback
import requests
import flet as ft
import logging
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver import ChromeOptions
from webdriver_manager.chrome import ChromeDriverManager
import urllib3
from bs4 import BeautifulSoup
import base64
from io import BytesIO
import matplotlib
matplotlib.use('Agg')  # Неинтерактивный бэкенд

# Проверка и установка необходимых модулей
try:
    import pandas as pd
except ImportError:
    print("❌ Ошибка: Модуль 'pandas' не установлен. Установите: pip install pandas")
    sys.exit(1)

try:
    import openpyxl
except ImportError:
    print("❌ Ошибка: Модуль 'openpyxl' не установлен. Установите: pip install openpyxl")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("⚠️ Предупреждение: Модуль 'python-docx' не установлен. Экспорт в Word будет недоступен. Установите: pip install python-docx")
    Document = None

try:
    import seaborn as sns
except ImportError:
    print("⚠️ Предупреждение: Модуль 'seaborn' не установлен. Графики могут работать некорректно.")
    sns = None

try:
    import matplotlib.pyplot as plt
except ImportError:
    print("❌ Ошибка: Модуль 'matplotlib' не установлен. Установите: pip install matplotlib")
    sys.exit(1)

try:
    from PIL import Image  # Не используется, но оставляем для совместимости
except ImportError:
    print("⚠️ Предупреждение: Модуль 'Pillow' не установлен. Некоторые функции могут работать некорректно.")
    Image = None
import xml.etree.ElementTree as ET
import threading
from collections import Counter
import string
import ast
import re
import hashlib
import pandas as pd
import os
import certifi
import ssl
ssl._create_default_https_context = ssl._create_unverified_context
os.environ['SSL_CERT_FILE'] = certifi.where()
import threading
from urllib.parse import urljoin, urlparse

# Подавление логов Selenium и Chrome
logging.getLogger('selenium').setLevel(logging.WARNING)
logging.getLogger('webdriver_manager').setLevel(logging.WARNING)
logging.getLogger('urllib3').setLevel(logging.WARNING)
logging.getLogger('requests').setLevel(logging.WARNING)

# Подавление логов Chrome
os.environ['WDM_LOG_LEVEL'] = '0'
os.environ['WDM_PRINT_FIRST_LINE'] = 'False'

# Отключение предупреждений urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# Отключение проверки SSL для webdriver_manager
os.environ['WDM_SSL_VERIFY'] = '0'

# Константы
LOG_FILE = "seo_log.txt"
SCREENSHOT_DIR = "screenshots"
REPORT_DIR = "reports"
os.makedirs(SCREENSHOT_DIR, exist_ok=True)
os.makedirs(REPORT_DIR, exist_ok=True)

# Глобальные переменные для экспорта данных
sitemap_export_data = {}

# Импорт SERP Tracker
try:
    from serp_tracker import SERPTracker, run_serp_tracking, run_detailed_site_analysis
except ImportError:
    print("⚠️ Предупреждение: Модуль 'serp_tracker' не найден. Функция трекинга позиций будет недоступна.")
    SERPTracker = None
    run_serp_tracking = None
    run_detailed_site_analysis = None

def log_to_file(text):
    """Записывает лог в файл с временной меткой."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S %Z")
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(f"[{timestamp}] {text}\n")

def create_webdriver(ignore_ssl=False, window_size=None, anti_bot_mode=False):
    """Создает WebDriver с улучшенными настройками для стабильности и обхода блокировок."""
    options = ChromeOptions()
    
    if anti_bot_mode:
        # Режим обхода блокировок
        options.add_argument("--headless=new")  # Добавляем headless режим
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        
        # Эмуляция реального браузера
        options.add_argument("--disable-web-security")
        options.add_argument("--allow-running-insecure-content")
        options.add_argument("--disable-features=VizDisplayCompositor")
        options.add_argument("--disable-features=TranslateUI")
        options.add_argument("--disable-ipc-flooding-protection")
        
        # Случайный User-Agent
        user_agents = [
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36",
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        ]
        import random
        options.add_argument(f"--user-agent={random.choice(user_agents)}")
        
        # Дополнительные настройки для обхода
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-gpu")
        options.add_argument("--disable-software-rasterizer")
        options.add_argument("--disable-background-networking")
        options.add_argument("--disable-background-timer-throttling")
        options.add_argument("--disable-backgrounding-occluded-windows")
        options.add_argument("--disable-renderer-backgrounding")
        options.add_argument("--log-level=3")
        options.add_argument("--mute-audio")
        options.add_argument("--disable-extensions")
        options.add_argument("--disable-plugins")
        options.add_argument("--disable-logging")
        options.add_argument("--silent")
        options.add_argument("--disable-dev-tools")
        
        # Эмуляция реального разрешения
        resolutions = ["1920x1080", "1366x768", "1440x900", "1536x864"]
        if not window_size:
            window_size = random.choice(resolutions)
        
        # Случайные настройки языка и геолокации
        languages = ["ru-RU,ru;q=0.9,en;q=0.8", "en-US,en;q=0.9", "ru-RU,ru;q=0.8,en-US;q=0.5,en;q=0.3"]
        options.add_argument(f"--lang={random.choice(languages)}")
        
    else:
        # Обычный режим
        options.add_argument("--headless=new")
        options.add_argument("--disable-gpu")
        options.add_argument("--disable-software-rasterizer")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-background-networking")
        options.add_argument("--disable-background-timer-throttling")
        options.add_argument("--disable-backgrounding-occluded-windows")
        options.add_argument("--disable-renderer-backgrounding")
        options.add_argument("--disable-features=TranslateUI")
        options.add_argument("--disable-ipc-flooding-protection")
        options.add_argument("--disable-web-security")
        options.add_argument("--disable-features=VizDisplayCompositor")
        options.add_argument("--log-level=3")
        options.add_argument("--mute-audio")
        options.add_argument("--disable-extensions")
        options.add_argument("--disable-plugins")
        options.add_argument("--disable-logging")
        options.add_argument("--silent")
        options.add_argument("--disable-dev-tools")
        options.add_experimental_option('excludeSwitches', ['enable-logging'])
        options.add_experimental_option('useAutomationExtension', False)
    
    options.add_experimental_option("prefs", {
        "profile.default_content_setting_values.notifications": 2,
        "profile.default_content_settings.popups": 0,
        "profile.default_content_setting_values.media_stream": 2,
        "profile.default_content_setting_values.geolocation": 2,
        "profile.default_content_setting_values.images": 2,  # Отключаем изображения для ускорения
        "profile.managed_default_content_settings.images": 2
    })
    
    if window_size:
        options.add_argument(f"--window-size={window_size}")
    
    if ignore_ssl:
        options.add_argument("--ignore-certificate-errors")
        options.add_argument("--ignore-ssl-errors")
    
    try:
        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
        
        if anti_bot_mode:
            # Дополнительные скрипты для обхода детекции
            driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
            driver.execute_script("Object.defineProperty(navigator, 'plugins', {get: () => [1, 2, 3, 4, 5]})")
            driver.execute_script("Object.defineProperty(navigator, 'languages', {get: () => ['ru-RU', 'ru', 'en-US', 'en']})")
        
        driver.set_page_load_timeout(60)
        driver.implicitly_wait(10)
        driver.set_script_timeout(30)
        return driver
    except Exception as e:
        log_to_file(f"Ошибка создания WebDriver: {str(e)}")
        raise e

def try_multiple_access_methods(site_url, ignore_ssl=False):
    """Пытается получить доступ к сайту различными методами."""
    methods_results = []
    
    # Метод 1: Обычный requests
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'ru-RU,ru;q=0.9,en;q=0.8',
            'Accept-Encoding': 'gzip, deflate, br',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
        }
        r = requests.get(site_url, timeout=15, verify=not ignore_ssl, headers=headers, allow_redirects=True)
        methods_results.append(('requests', r.status_code, r.text[:500], r.headers))
    except Exception as e:
        methods_results.append(('requests', 'error', str(e), {}))
    
    # Метод 2: requests с сессией
    try:
        session = requests.Session()
        session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
        })
        r = session.get(site_url, timeout=15, verify=not ignore_ssl, allow_redirects=True)
        methods_results.append(('session', r.status_code, r.text[:500], r.headers))
    except Exception as e:
        methods_results.append(('session', 'error', str(e), {}))
    
    # Метод 3: requests с прокси-заголовками
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'ru-RU,ru;q=0.8,en-US;q=0.5,en;q=0.3',
            'Accept-Encoding': 'gzip, deflate, br',
            'Connection': 'keep-alive',
            'Cache-Control': 'max-age=0',
            'Sec-Fetch-Dest': 'document',
            'Sec-Fetch-Mode': 'navigate',
            'Sec-Fetch-Site': 'none',
            'Sec-Fetch-User': '?1',
        }
        r = requests.get(site_url, timeout=15, verify=not ignore_ssl, headers=headers, allow_redirects=True)
        methods_results.append(('headers', r.status_code, r.text[:500], r.headers))
    except Exception as e:
        methods_results.append(('headers', 'error', str(e), {}))
    
    return methods_results

def check_site_accessibility(site_url, ignore_ssl=False):
    """Проверяет доступность сайта различными методами."""
    log_text = f"\n🔍 Проверка доступности сайта: {site_url}\n"
    
    # Пробуем разные методы доступа
    methods_results = try_multiple_access_methods(site_url, ignore_ssl)
    
    successful_methods = []
    blocked_methods = []
    
    for method, status, content, headers in methods_results:
        log_text += f"\n📡 Метод {method}:\n"
        
        if status == 'error':
            log_text += f"❌ Ошибка: {content}\n"
            blocked_methods.append(method)
        elif status == 200:
            log_text += f"✅ Успешно (статус: {status})\n"
            successful_methods.append(method)
            
            # Проверяем на блокировку по содержимому
            if any(blocked_text in content.lower() for blocked_text in [
                'access denied', 'доступ запрещен', 'blocked', 'заблокирован',
                'cloudflare', 'captcha', 'recaptcha', 'bot', 'robot'
            ]):
                log_text += f"⚠️ Обнаружена блокировка в содержимом\n"
                blocked_methods.append(method)
            else:
                log_text += f"✅ Содержимое доступно\n"
        else:
            log_text += f"⚠️ Статус: {status}\n"
            if status in [403, 429, 503]:
                blocked_methods.append(method)
            else:
                successful_methods.append(method)
    
    # Анализ результатов
    if successful_methods:
        log_text += f"\n✅ Успешные методы: {', '.join(successful_methods)}\n"
        return True, log_text, successful_methods[0]
    else:
        log_text += f"\n❌ Все методы заблокированы: {', '.join(blocked_methods)}\n"
        return False, log_text, None

def parse_summary(summary_content, report_type='full'):
    """Парсит сводку для Excel."""
    if report_type == 'images':
        data = {'Ссылка': [], 'Alt': [], 'Title': [], 'Размер': []}
        current = {}
        for line in summary_content.split('\n'):
            line = line.strip()
            if line.startswith('Ссылка: '):
                if current:
                    data['Ссылка'].append(current.get('src', ''))
                    data['Alt'].append(current.get('alt', ''))
                    data['Title'].append(current.get('title', ''))
                    data['Размер'].append(current.get('size', ''))
                current = {'src': line[8:]}
            elif line.startswith('Alt: '):
                current['alt'] = line[5:]
            elif line.startswith('Title: '):
                current['title'] = line[7:]
            elif line.startswith('Размер: '):
                current['size'] = line[8:]
        if current:
            data['Ссылка'].append(current.get('src', ''))
            data['Alt'].append(current.get('alt', ''))
            data['Title'].append(current.get('title', ''))
            data['Размер'].append(current.get('size', ''))
        
        # Проверяем, что есть данные для экспорта
        if not data['Ссылка']:
            # Если данных нет, создаем базовую структуру
            data['Ссылка'] = ['Нет данных']
            data['Alt'] = ['-']
            data['Title'] = ['-']
            data['Размер'] = ['-']
        
        return data
    elif report_type == 'parser':
        data = {'Ссылка': [], 'Результат': [], 'Статус': []}
        for line in summary_content.split('\n'):
            line = line.strip()
            if line.startswith('- '):
                parts = line[2:].split(': ')
                if len(parts) == 2:
                    url = parts[0]
                    details = parts[1].split(' ')
                    emoji = details[0]
                    status = ' '.join(details[1:])
                    data['Ссылка'].append(url)
                    data['Результат'].append(emoji)
                    data['Статус'].append(status)
        
        # Проверяем, что есть данные для экспорта
        if not data['Ссылка']:
            # Если данных нет, создаем базовую структуру
            data['Ссылка'] = ['Нет данных']
            data['Результат'] = ['-']
            data['Статус'] = ['Нет данных']
        
        return data
    else:
        sections = {'Хорошее': [], 'Проблемы': [], 'Рекомендации': []}
        current_section = None
        for line in summary_content.split('\n'):
            line = line.strip()
            if line.startswith('**Хорошее:**'):
                current_section = 'Хорошее'
            elif line.startswith('**Проблемы:**'):
                current_section = 'Проблемы'
            elif line.startswith('**Рекомендации:**'):
                current_section = 'Рекомендации'
            elif line.startswith('✅ ') or line.startswith('❌ ') or line.startswith('📝 '):
                if current_section:
                    sections[current_section].append(line[2:].strip())
            elif line.startswith('- '):  # Для списков, как в изображениях/ссылках
                if current_section:
                    sections[current_section].append(line[2:].strip())
        max_len = max(len(sections[s]) for s in sections)
        if max_len > 0:  # Проверяем, что есть данные для обработки
            for s in sections:
                sections[s] += [''] * (max_len - len(sections[s]))
        else:
            # Если все списки пустые, создаем базовую структуру
            for s in sections:
                sections[s] = ['']
        return sections

def save_results(site_url, log_content, summary_content, report_type='full', format='excel'):
    """Сохраняет результаты в Excel."""
    report_path = f"{REPORT_DIR}/{report_type}_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    sections = parse_summary(summary_content, report_type)
    
    # Проверяем, что есть данные для сохранения
    if not sections or all(len(v) == 0 for v in sections.values()):
        # Создаем базовую структуру с сообщением об отсутствии данных
        if report_type == 'parser':
            sections = {
                'Ссылка': ['Нет данных для экспорта'],
                'Результат': ['-'],
                'Статус': ['Проверьте настройки парсера']
            }
        elif report_type == 'images':
            sections = {
                'Ссылка': ['Нет данных для экспорта'],
                'Alt': ['-'],
                'Title': ['-'],
                'Размер': ['-']
            }
        else:
            sections = {
                'Хорошее': ['Нет данных для экспорта'],
                'Проблемы': ['Проверьте настройки анализа'],
                'Рекомендации': ['Запустите анализ заново']
            }
    
    df = pd.DataFrame(sections)
    df.to_excel(report_path, index=False)
    return report_path

def check_resource(url, ignore_ssl):
    """Проверяет доступность ресурса и возвращает статус, время, историю редиректов."""
    try:
        response = requests.get(url, timeout=5, verify=not ignore_ssl, allow_redirects=True)
        return url, response.status_code, response.elapsed.total_seconds(), response.history
    except Exception as e:
        return url, f"Error: {str(e)}", 0, []

def get_image_size(url, ignore_ssl):
    """Получает размер изображения в КБ."""
    try:
        response = requests.get(url, timeout=5, verify=not ignore_ssl)
        if response.status_code == 200:
            return len(response.content) / 1024
        return 0
    except:
        return 0

def check_seo_files(site_url, ignore_ssl):
    """Проверяет доступность robots.txt и sitemap.xml."""
    results = []
    for file in ["robots.txt", "sitemap.xml"]:
        url = f"{site_url.rstrip('/')}/{file}"
        try:
            response = requests.get(url, timeout=5, verify=not ignore_ssl)
            results.append((file, response.status_code == 200, response.text if response.status_code == 200 else ""))
        except Exception as e:
            results.append((file, False, str(e)))
    return results

def analyze_robots_txt(robots_content):
    """Анализирует robots.txt на ошибки и рекомендации."""
    errors = []
    positives = []
    found_directives = []
    recommendations = []
    if not robots_content:
        errors.append("robots.txt пустой или недоступен")
        recommendations.append("Создайте robots.txt с базовыми правилами, например: 'User-agent: * \nDisallow: /private/'")
        return errors, positives, found_directives, recommendations

    lines = [line.strip() for line in robots_content.splitlines() if line.strip() and not line.strip().startswith("#")]
    user_agents = set()
    disallow_rules = []
    allow_rules = []
    sitemap_urls = []
    crawl_delays = {}
    host = None

    for line in lines:
        parts = line.split(":", 1)
        if len(parts) != 2:
            continue
        directive, value = parts[0].lower().strip(), parts[1].strip()
        found_directives.append(line)
        if directive == "user-agent":
            user_agents.add(value)
        elif directive == "disallow":
            disallow_rules.append(value)
        elif directive == "allow":
            allow_rules.append(value)
        elif directive == "sitemap":
            sitemap_urls.append(value)
        elif directive == "crawl-delay":
            try:
                ua, delay = value.split(maxsplit=1)
                crawl_delays[ua] = float(delay)
            except ValueError:
                errors.append(f"Некорректный Crawl-delay: {value}")
                recommendations.append("Убедитесь, что Crawl-delay указан как 'User-agent: <bot> <delay>'.")
        elif directive == "host":
            host = value

    if user_agents:
        positives.append(f"Найдено {len(user_agents)} User-agent(s): {', '.join(user_agents)}")
    if disallow_rules:
        positives.append(f"Найдено {len(disallow_rules)} правил Disallow")
    if allow_rules:
        positives.append(f"Найдено {len(allow_rules)} правил Allow")
    if sitemap_urls:
        positives.append(f"Найдено {len(sitemap_urls)} Sitemap URL(s): {', '.join(sitemap_urls)}")
    if host:
        positives.append(f"Найден Host: {host}")
    if crawl_delays:
        positives.append(f"Найдено {len(crawl_delays)} правил Crawl-delay")

    if not user_agents:
        errors.append("Отсутствует секция User-agent")
        recommendations.append("Добавьте 'User-agent: *' для общих правил.")
    if not disallow_rules and not allow_rules:
        errors.append("Отсутствуют правила Disallow или Allow")
        recommendations.append("Добавьте хотя бы одно правило, например 'Disallow: /private/'.")
    if "*" not in user_agents and len(user_agents) > 1:
        errors.append("Отсутствует общий User-agent (*) при множественных агентах")
        recommendations.append("Добавьте 'User-agent: *' для общих настроек.")
    if any(not rule.startswith("/") for rule in disallow_rules + allow_rules):
        errors.append("Правила Disallow/Allow не начинаются с '/'")
        recommendations.append("Убедитесь, что пути начинаются с '/', например 'Disallow: /admin/'.")
    if len(sitemap_urls) > 1 and not all(url.startswith("http") for url in sitemap_urls):
        errors.append("Некоторые Sitemap URL не являются абсолютными")
        recommendations.append("Используйте полные URL, например 'Sitemap: https://example.com/sitemap.xml'.")

    return errors, positives, found_directives, recommendations

def check_robots_summary(site_url, ignore_ssl):
    """Отдельная функция для проверки robots.txt с выводом что хорошо и что плохо."""
    seo_files = check_seo_files(site_url, ignore_ssl)
    robots_status = next((status for file, status, content in seo_files if file == "robots.txt"), False)
    robots_content = next((content for file, status, content in seo_files if file == "robots.txt"), "")
    
    if not robots_status:
        return "❌ robots.txt не найден или недоступен\n📝 Рекомендация: Создайте robots.txt."
    
    errors, positives, found_directives, recommendations = analyze_robots_txt(robots_content)
    
    summary = "### Проверка robots.txt\n\n"
    if positives:
        summary += "**Хорошее:**\n" + "\n".join(f"✅ {p}" for p in positives) + "\n\n"
    if errors:
        summary += "**Проблемы:**\n" + "\n".join(f"❌ {e}" for e in errors) + "\n\n"
    if recommendations:
        summary += "**Рекомендации:**\n" + "\n".join(f"📝 {r}" for r in recommendations) + "\n\n"
    summary += f"**Найденные директивы:** {', '.join(found_directives) if found_directives else 'Нет директив'}\n"
    
    return summary

def get_site_pages(site_url, ignore_ssl, max_pages=15000):
    """Получает список всех страниц на сайте."""
    site_pages = set()
    
    try:
        # Нормализуем базовый URL
        base_url = site_url.rstrip('/')
        
        # Получаем главную страницу
        response = requests.get(site_url, verify=not ignore_ssl, timeout=10)
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Ищем все ссылки на том же домене
            for link in soup.find_all('a', href=True):
                href = link['href'].strip()
                if not href:
                    continue
                    
                # Нормализуем URL
                if href.startswith('/'):
                    # Относительная ссылка
                    full_url = base_url + href
                elif href.startswith(base_url):
                    # Абсолютная ссылка на том же домене
                    full_url = href
                elif href.startswith('http') and base_url in href:
                    # Ссылка на тот же домен
                    full_url = href
                else:
                    # Пропускаем внешние ссылки
                    continue
                
                # Нормализуем URL (убираем дублирующиеся слеши и фрагменты)
                full_url = full_url.rstrip('/')
                if '#' in full_url:
                    full_url = full_url.split('#')[0]
                if '?' in full_url:
                    full_url = full_url.split('?')[0]
                
                # Добавляем только если это валидный URL
                if full_url.startswith('http'):
                    site_pages.add(full_url)
            
            # Добавляем главную страницу
            site_pages.add(base_url)
            
    except Exception as e:
        print(f"Ошибка при получении страниц сайта: {e}")
    
    # Возвращаем отсортированный список
    return sorted(list(site_pages))[:max_pages]

def validate_sitemap(sitemap_content, site_url, ignore_ssl):
    """Валидирует sitemap.xml с поддержкой sitemap index и более подробной проверкой."""
    errors = []
    positives = []
    recommendations = []
    urls_in_sitemap = []
    page_details = []  # Новый список для подробностей по каждой странице
    sitemap_info = {}  # Информация о структуре sitemap
    pages_not_in_sitemap = []  # Страницы на сайте, но не в sitemap
    pages_in_sitemap_not_on_site = []  # Страницы в sitemap, но не на сайте
    
    try:
        root = ET.fromstring(sitemap_content)
        
        # Проверяем тип sitemap
        if 'sitemapindex' in root.tag:
            # Это sitemap index
            positives.append("Обнаружен sitemap index (иерархическая структура)")
            sitemap_elems = list(root.iter('{http://www.sitemaps.org/schemas/sitemap/0.9}sitemap'))
            sitemap_urls = [sitemap_elem.findtext('{http://www.sitemaps.org/schemas/sitemap/0.9}loc') for sitemap_elem in sitemap_elems]
            sitemap_info['type'] = 'sitemapindex'
            sitemap_info['sub_sitemaps'] = sitemap_urls
            positives.append(f"Найдено {len(sitemap_urls)} подчиненных sitemap")
            
            # Рекурсивно обрабатываем все подчиненные sitemap
            all_urls = []
            all_sources = {}
            all_metadata = {}
            all_errors = []
            
            for sub_sitemap_url in sitemap_urls:
                if sub_sitemap_url:
                    sub_result = process_sitemap_recursively(sub_sitemap_url, ignore_ssl)
                    all_urls.extend(sub_result['urls'])
                    all_sources.update(sub_result['sources'])
                    all_metadata.update(sub_result['metadata'])
                    all_errors.extend(sub_result['errors'])
            
            urls_in_sitemap = all_urls
            errors.extend(all_errors)
            
            if all_urls:
                positives.append(f"Всего найдено {len(all_urls)} URL во всех sitemap")
                
                # Обрабатываем все URL и проверяем доступность для всех
                for i, url in enumerate(all_urls):
                    # Получаем метаданные для URL
                    url_metadata = all_metadata.get(url, {})
                    lastmod = url_metadata.get('lastmod', '-')
                    priority = url_metadata.get('priority', '-')
                    changefreq = url_metadata.get('changefreq', '-')
                    
                    # Проверяем доступность для всех URL
                    status = 'не ОК'
                    if url and url.startswith('http'):
                        _, http_status, _, _ = check_resource(url, ignore_ssl)
                        if isinstance(http_status, int) and http_status == 200:
                            status = 'ОК'
                        else:
                            status = f'не ОК ({http_status})'
                    
                    page_details.append({
                        'url': url,
                        'status': status,
                        'lastmod': lastmod,
                        'priority': priority,
                        'changefreq': changefreq,
                        'source_sitemap': all_sources.get(url, 'неизвестно')
                    })
            
        elif 'urlset' in root.tag:
            # Это обычный sitemap
            positives.append("Корневой элемент верный (обычный sitemap)")
            sitemap_info['type'] = 'urlset'
            url_elems = list(root.iter('{http://www.sitemaps.org/schemas/sitemap/0.9}url'))
            urls = [url_elem.findtext('{http://www.sitemaps.org/schemas/sitemap/0.9}loc') for url_elem in url_elems]
            urls_in_sitemap = urls
            positives.append(f"Найдено {len(urls)} URL в sitemap")
            
            if len(urls) > 50000:
                errors.append("Sitemap содержит более 50,000 URL")
                recommendations.append("Разделите sitemap на несколько файлов.")
            
            if not all(url and url.startswith('http') for url in urls if url):
                errors.append("Некоторые URL в sitemap недействительны или относительные")
                recommendations.append("Используйте абсолютные URL в sitemap.")
            
            # Проверка доступности URL в sitemap и сбор подробностей
            for url_elem in url_elems:
                url = url_elem.findtext('{http://www.sitemaps.org/schemas/sitemap/0.9}loc')
                lastmod = url_elem.findtext('{http://www.sitemaps.org/schemas/sitemap/0.9}lastmod') or '-'
                priority = url_elem.findtext('{http://www.sitemaps.org/schemas/sitemap/0.9}priority') or '-'
                changefreq = url_elem.findtext('{http://www.sitemaps.org/schemas/sitemap/0.9}changefreq') or '-'
                status = 'не ОК'
                http_status = None
                if url and url.startswith('http'):
                    _, http_status, _, _ = check_resource(url, ignore_ssl)
                    if isinstance(http_status, int) and http_status == 200:
                        status = 'ОК'
                    else:
                        status = f'не ОК ({http_status})'
                page_details.append({
                    'url': url,
                    'status': status,
                    'lastmod': lastmod,
                    'priority': priority,
                    'changefreq': changefreq,
                    'source_sitemap': 'основной sitemap'
                })
        else:
            errors.append("Неверный корневой элемент sitemap.xml (ожидается urlset или sitemapindex)")
            recommendations.append("Убедитесь, что sitemap соответствует схеме http://www.sitemaps.org/schemas/sitemap/0.9")
        
        # Если есть не ОК, добавить в errors
        broken_urls = [d['url'] for d in page_details if d['status'] != 'ОК']
        if broken_urls:
            # Показываем все проблемные URL в интерфейсе
            errors.append(f"Недоступные URL в sitemap ({len(broken_urls)} URL):")
            for url in broken_urls:
                errors.append(f"  - {url}")
            recommendations.append("Исправьте или удалите недоступные URL из sitemap.")
        
        # Анализируем страницы на сайте vs sitemap
        try:
            site_pages = get_site_pages(site_url, ignore_ssl)
            
            # Нормализуем URL для корректного сравнения
            def normalize_url(url):
                """Нормализует URL для сравнения"""
                if not url:
                    return url
                url = url.rstrip('/')
                if '#' in url:
                    url = url.split('#')[0]
                if '?' in url:
                    url = url.split('?')[0]
                return url
            
            # Нормализуем все URL
            normalized_site_pages = set(normalize_url(url) for url in site_pages)
            normalized_sitemap_urls = set(normalize_url(url) for url in urls_in_sitemap)
            
            # Страницы на сайте, но не в sitemap
            pages_not_in_sitemap = [url for url in site_pages if normalize_url(url) not in normalized_sitemap_urls]
            
            # Страницы в sitemap, но не на сайте (недоступные)
            pages_in_sitemap_not_on_site = [url for url in urls_in_sitemap if normalize_url(url) not in normalized_site_pages and url in broken_urls]
            
            if pages_not_in_sitemap:
                errors.append(f"Страницы на сайте не в sitemap ({len(pages_not_in_sitemap)} URL):")
                for url in pages_not_in_sitemap[:10]:  # Показываем первые 10
                    errors.append(f"  - {url}")
                if len(pages_not_in_sitemap) > 10:
                    errors.append(f"  ... и еще {len(pages_not_in_sitemap) - 10} URL")
                recommendations.append("Добавьте найденные страницы в sitemap.")
            
            if pages_in_sitemap_not_on_site:
                errors.append(f"Страницы в sitemap не на сайте ({len(pages_in_sitemap_not_on_site)} URL):")
                for url in pages_in_sitemap_not_on_site[:10]:  # Показываем первые 10
                    errors.append(f"  - {url}")
                if len(pages_in_sitemap_not_on_site) > 10:
                    errors.append(f"  ... и еще {len(pages_in_sitemap_not_on_site) - 10} URL")
                recommendations.append("Удалите недоступные страницы из sitemap.")
                
        except Exception as e:
            print(f"Ошибка при анализе страниц сайта: {e}")
            
    except ET.ParseError as e:
        errors.append(f"Ошибка парсинга sitemap.xml: {str(e)}")
        recommendations.append("Исправьте XML структуру sitemap.")
        page_details = []
    
    return errors, positives, recommendations, urls_in_sitemap, page_details, sitemap_info, pages_not_in_sitemap, pages_in_sitemap_not_on_site

def check_sitemap_summary(site_url, ignore_ssl):
    """Отдельная функция для проверки sitemap.xml с выводом что хорошо и что плохо и подробностями по страницам."""
    seo_files = check_seo_files(site_url, ignore_ssl)
    sitemap_status = next((status for file, status, content in seo_files if file == "sitemap.xml"), False)
    sitemap_content = next((content for file, status, content in seo_files if file == "sitemap.xml"), "")
    
    if not sitemap_status:
        return "❌ sitemap.xml не найден или недоступен\n📝 Рекомендация: Создайте sitemap.xml."
    
    errors, positives, recommendations, urls_in_sitemap, page_details, sitemap_info, pages_not_in_sitemap, pages_in_sitemap_not_on_site = validate_sitemap(sitemap_content, site_url, ignore_ssl)
    
    # Сохраняем полные данные для экспорта
    global sitemap_export_data
    sitemap_export_data = {
        'urls': urls_in_sitemap,
        'page_details': page_details,
        'sitemap_info': sitemap_info,
        'errors': errors,
        'positives': positives,
        'recommendations': recommendations,
        'pages_not_in_sitemap': pages_not_in_sitemap,
        'pages_in_sitemap_not_on_site': pages_in_sitemap_not_on_site
    }
    
    summary = "### Проверка sitemap.xml\n\n"
    
    # Добавляем информацию о типе sitemap
    if sitemap_info.get('type') == 'sitemapindex':
        summary += "**🔗 Тип sitemap:** Sitemap Index (иерархическая структура)\n\n"
        if sitemap_info.get('sub_sitemaps'):
            summary += "**📋 Подчиненные sitemap:**\n"
            for i, sub_sitemap in enumerate(sitemap_info['sub_sitemaps'][:10], 1):
                summary += f"{i}. {sub_sitemap}\n"
            if len(sitemap_info['sub_sitemaps']) > 10:
                summary += f"... и еще {len(sitemap_info['sub_sitemaps']) - 10} sitemap\n"
            summary += "\n"
    else:
        summary += "**🔗 Тип sitemap:** Обычный sitemap\n\n"
    
    if positives:
        summary += "**Хорошее:**\n" + "\n".join(f"✅ {p}" for p in positives) + "\n\n"
    if errors:
        summary += "**Проблемы:**\n"
        for e in errors:
            if e.startswith("Недоступные URL в sitemap"):
                # Для проблемных URL показываем их отдельно
                summary += f"❌ {e}\n"
            else:
                summary += f"❌ {e}\n"
        summary += "\n"
    if recommendations:
        summary += "**Рекомендации:**\n" + "\n".join(f"📝 {r}" for r in recommendations) + "\n\n"
    
    # Добавляем отдельный раздел с проблемными URL
    broken_urls = [d for d in page_details if d['status'] != 'ОК']
    if broken_urls:
        summary += "**🔴 Проблемные URL (недоступные):**\n"
        for i, d in enumerate(broken_urls, 1):
            summary += f"{i}. {d['url']} - {d['status']}\n"
        summary += "\n"
    
    # Добавляем отдельный раздел с доступными URL
    working_urls = [d for d in page_details if d['status'] == 'ОК']
    if working_urls:
        summary += f"**🟢 Доступные URL: {len(working_urls)} URL**\n\n"
    
    if page_details:
        summary += "**Детализация по страницам (первые 50):**\n"
        # Показываем только первые 50 URL в интерфейсе
        display_details = page_details[:50]
        for d in display_details:
            summary += f"URL: {d['url']}\n"
            summary += f"Статус: {d['status']}\n"
            if 'source_sitemap' in d:
                summary += f"Источник: {d['source_sitemap']}\n"
            if d['lastmod'] != '-':
                summary += f"lastmod: {d['lastmod']}\n"
            if d['priority'] != '-':
                summary += f"priority: {d['priority']}\n"
            if d['changefreq'] != '-':
                summary += f"changefreq: {d['changefreq']}\n"
            summary += "\n"
        
        # Если есть больше URL, показываем информацию
        if len(page_details) > 50:
            summary += f"**... и еще {len(page_details) - 50} URL (используйте экспорт для полного списка)**\n\n"
    return summary

def analyze_keywords(driver, site_url, target_keywords):
    """Анализирует ключевые слова на странице с учетом склонений русских слов."""
    try:
        # Получаем только видимый текст через JavaScript
        visible_text = driver.execute_script("""
            function getVisibleText() {
                // Удаляем скрытые элементы
                const hiddenElements = document.querySelectorAll('script, style, noscript, [style*="display: none"], [style*="visibility: hidden"], .hidden, .invisible');
                hiddenElements.forEach(el => el.style.display = 'none');
                
                // Получаем текст только из видимых элементов
                const walker = document.createTreeWalker(
                    document.body,
                    NodeFilter.SHOW_TEXT,
                    {
                        acceptNode: function(node) {
                            const style = window.getComputedStyle(node.parentElement);
                            if (style.display === 'none' || style.visibility === 'hidden' || style.opacity === '0') {
                                return NodeFilter.FILTER_REJECT;
                            }
                            return NodeFilter.FILTER_ACCEPT;
                        }
                    }
                );
                
                let text = '';
                let node;
                while (node = walker.nextNode()) {
                    text += node.textContent + ' ';
                }
                
                return text;
            }
            return getVisibleText();
        """)
        
        # Очищаем текст
        text = re.sub(r'\s+', ' ', visible_text).strip().lower()
        
        # Удаляем мета-теги и служебную информацию
        text = re.sub(r'<[^>]+>', '', text)  # Удаляем HTML теги
        text = re.sub(r'javascript:', '', text, flags=re.IGNORECASE)
        text = re.sub(r'http[s]?://[^\s]+', '', text)  # Удаляем URL
        text = re.sub(r'www\.[^\s]+', '', text)
        text = re.sub(r'[^\w\sа-яё]', ' ', text)  # Оставляем только буквы и пробелы
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Удаление стоп-слов и знаков препинания (расширенный список для русского)
        stop_words = {'и', 'в', 'на', 'не', 'с', 'а', 'о', 'для', 'по', 'из', 'к', 'у', 'от', 'но', 'как', 'что', 'это', 'то', 'или', 'за', 'при', 'meta', 'title', 'description', 'keywords', 'og', 'twitter', 'schema', 'json', 'ld'}
        words = re.findall(r'\w+', text)
        word_freq = {}
        for word in words:
            if word not in stop_words and len(word) > 3:
                word_freq[word] = word_freq.get(word, 0) + 1
        # Сортировка по частоте и выбор топ-10
        keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
        total_words = len(words) or 1  # Избежать деления на 0
        keyword_density = {word: freq / total_words for word, freq in keywords}

        # Функция для генерации склонений русского слова
        def generate_declensions(word):
            """Генерирует расширенные склонения для русского слова."""
            declensions = [word]
            
            # Специальные случаи для часто встречающихся слов
            special_declensions = {
                'доставка': ['доставки', 'доставку', 'доставкой', 'доставке', 'доставки', 'доставок', 'доставкам', 'доставками', 'доставках'],
                'цветы': ['цветов', 'цветам', 'цветы', 'цветами', 'цветах', 'цветок', 'цветка', 'цветку', 'цветком', 'цветке', 'цветки'],
                'цветок': ['цветка', 'цветку', 'цветок', 'цветком', 'цветке', 'цветки', 'цветов', 'цветам', 'цветами', 'цветах'],
                'купить': ['купить', 'купил', 'купила', 'купили', 'купим', 'купите', 'купишь', 'купит', 'купят', 'купил', 'купила', 'купили'],
                'заказать': ['заказать', 'заказал', 'заказала', 'заказали', 'закажем', 'закажете', 'закажет', 'закажут', 'заказал', 'заказала', 'заказали'],
                'цена': ['цены', 'цену', 'ценой', 'цене', 'цены', 'ценами', 'ценах'],
                'стоимость': ['стоимости', 'стоимость', 'стоимостью', 'стоимости', 'стоимостями', 'стоимостях'],
                'магазин': ['магазина', 'магазину', 'магазин', 'магазином', 'магазине', 'магазины', 'магазинов', 'магазинам', 'магазинами', 'магазинах'],
                'сайт': ['сайта', 'сайту', 'сайт', 'сайтом', 'сайте', 'сайты', 'сайтов', 'сайтам', 'сайтами', 'сайтах'],
                'услуга': ['услуги', 'услугу', 'услугой', 'услуге', 'услуги', 'услугами', 'услугах'],
                'товар': ['товара', 'товару', 'товар', 'товаром', 'товаре', 'товары', 'товаров', 'товарам', 'товарами', 'товарах'],
                'компания': ['компании', 'компанию', 'компанией', 'компании', 'компании', 'компаний', 'компаниям', 'компаниями', 'компаниях'],
                'фирма': ['фирмы', 'фирму', 'фирмой', 'фирме', 'фирмы', 'фирм', 'фирмам', 'фирмами', 'фирмах'],
                'интернет': ['интернета', 'интернету', 'интернет', 'интернетом', 'интернете'],
                'онлайн': ['онлайн', 'онлайном'],
                'быстро': ['быстро', 'быстрее', 'быстрей'],
                'качественно': ['качественно', 'качественнее'],
                'дешево': ['дешево', 'дешевле', 'дешевей'],
                'дорого': ['дорого', 'дороже'],
                'москва': ['москвы', 'москве', 'москвой', 'москве'],
                'санкт-петербург': ['санкт-петербурга', 'санкт-петербургу', 'санкт-петербургом', 'санкт-петербурге'],
                'россия': ['россии', 'россию', 'россией', 'россии'],
                'город': ['города', 'городу', 'город', 'городом', 'городе', 'города', 'городов', 'городам', 'городами', 'городах'],
                'регион': ['региона', 'региону', 'регион', 'регионом', 'регионе', 'регионы', 'регионов', 'регионам', 'регионами', 'регионах'],
                # Добавляем специальные случаи для цветов
                'розы': ['роз', 'розам', 'розы', 'розами', 'розах', 'роза', 'розы', 'розе', 'розой', 'розе'],
                'тюльпаны': ['тюльпанов', 'тюльпанам', 'тюльпаны', 'тюльпанами', 'тюльпанах', 'тюльпан', 'тюльпана', 'тюльпану', 'тюльпаном', 'тюльпане'],
                'лилии': ['лилий', 'лилиям', 'лилии', 'лилиями', 'лилиях', 'лилия', 'лилии', 'лилии', 'лилией', 'лилии'],
                'орхидеи': ['орхидей', 'орхидеям', 'орхидеи', 'орхидеями', 'орхидеях', 'орхидея', 'орхидеи', 'орхидее', 'орхидеей', 'орхидее'],
                'хризантемы': ['хризантем', 'хризантемам', 'хризантемы', 'хризантемами', 'хризантемах', 'хризантема', 'хризантемы', 'хризантеме', 'хризантемой', 'хризантеме']
            }
            
            # Проверяем специальные случаи
            if word in special_declensions:
                declensions.extend(special_declensions[word])
            else:
                # Универсальные правила склонения для любых русских слов
                word_len = len(word)
                
                # Существительные женского рода на -а
                if word.endswith('а') and word_len > 2:
                    base = word[:-1]
                    declensions.extend([base + 'ы', base + 'у', base + 'ой', base + 'е', base + 'ами', base + 'ах'])
                
                # Существительные женского рода на -я
                elif word.endswith('я') and word_len > 2:
                    base = word[:-1]
                    declensions.extend([base + 'и', base + 'ю', base + 'ей', base + 'е', base + 'ями', base + 'ях'])
                
                # Существительные женского рода на -ь
                elif word.endswith('ь') and word_len > 2:
                    base = word[:-1]
                    declensions.extend([base + 'я', base + 'ю', base + 'ью', base + 'и', base + 'ями', base + 'ях'])
                
                # Существительные мужского рода на -й
                elif word.endswith('й') and word_len > 2:
                    base = word[:-1]
                    declensions.extend([base + 'я', base + 'ю', base + 'ем', base + 'е', base + 'ями', base + 'ях'])
                
                # Существительные среднего рода на -о
                elif word.endswith('о') and word_len > 2:
                    base = word[:-1]
                    declensions.extend([base + 'а', base + 'у', base + 'ом', base + 'е', base + 'ами', base + 'ах'])
                
                # Существительные среднего рода на -е
                elif word.endswith('е') and word_len > 2:
                    base = word[:-1]
                    declensions.extend([base + 'я', base + 'ю', base + 'ем', base + 'е', base + 'ями', base + 'ях'])
                
                # Существительные мужского рода на согласную
                elif word_len > 2 and not word.endswith(('а', 'я', 'ь', 'й', 'о', 'е')):
                    # Добавляем окончания для мужского рода
                    declensions.extend([word + 'а', word + 'у', word + 'ом', word + 'е', word + 'ы', word + 'ов', word + 'ам', word + 'ами', word + 'ах'])
                
                # Глаголы (инфинитив на -ть)
                elif word.endswith('ть') and word_len > 3:
                    base = word[:-2]  # убираем 'ть'
                    declensions.extend([
                        base + 'л', base + 'ла', base + 'ли',  # прошедшее время
                        base + 'ю', base + 'ешь', base + 'ет',  # настоящее время
                        base + 'ем', base + 'ете', base + 'ют',
                        base + 'й', base + 'йте'  # повелительное наклонение
                    ])
                
                # Прилагательные на -ый, -ой, -ий
                elif word.endswith(('ый', 'ой', 'ий')) and word_len > 3:
                    base = word[:-2]
                    declensions.extend([
                        base + 'ого', base + 'ому', base + 'ым', base + 'ом',  # мужской род
                        base + 'ая', base + 'ой', base + 'ую', base + 'ой',    # женский род
                        base + 'ое', base + 'ого', base + 'ому', base + 'ым',  # средний род
                        base + 'ые', base + 'ых', base + 'ым', base + 'ыми'    # множественное число
                    ])
                
                # Прилагательные на -ая, -яя
                elif word.endswith(('ая', 'яя')) and word_len > 3:
                    base = word[:-2]
                    declensions.extend([
                        base + 'ой', base + 'ую', base + 'ой',  # женский род
                        base + 'ое', base + 'ого', base + 'ому', base + 'ым',  # средний род
                        base + 'ые', base + 'ых', base + 'ым', base + 'ыми'    # множественное число
                    ])
                
                # Прилагательные на -ое, -ее
                elif word.endswith(('ое', 'ее')) and word_len > 3:
                    base = word[:-2]
                    declensions.extend([
                        base + 'ого', base + 'ому', base + 'ым', base + 'ом',  # мужской род
                        base + 'ая', base + 'ой', base + 'ую', base + 'ой',    # женский род
                        base + 'ые', base + 'ых', base + 'ым', base + 'ыми'    # множественное число
                    ])
            
            # Убираем дубликаты и возвращаем уникальные склонения
            return list(set(declensions))

        # Анализ целевых ключевых слов с учетом склонений
        target_analysis = {}
        if target_keywords:
            target_list = [kw.strip().lower() for kw in target_keywords.split(',')]
            
            for keyword_phrase in target_list:
                # Разбиваем фразу на слова
                words_in_phrase = keyword_phrase.split()
                
                if len(words_in_phrase) == 1:
                    # Одно слово - генерируем склонения
                    word = words_in_phrase[0]
                    declensions = generate_declensions(word)
                    
                    # Подсчитываем вхождения каждого склонения (только уникальные)
                    total_count = 0
                    declension_counts = {}
                    
                    for declension in declensions:
                        # Используем только точный поиск с границами слов
                        pattern = r'\b' + re.escape(declension) + r'\b'
                        matches = re.findall(pattern, text, re.IGNORECASE)
                        
                        # Подсчитываем уникальные вхождения
                        count = len(matches)
                        
                        if count > 0:
                            # Проверяем, что это не дубликат (например, "доставки" и "доставки" - это одно и то же)
                            is_duplicate = False
                            for existing_declension in declension_counts:
                                if declension.lower() == existing_declension.lower():
                                    is_duplicate = True
                                    break
                            
                            if not is_duplicate:
                                declension_counts[declension] = count
                                total_count += count
                    
                    # Вычисляем плотность
                    density = (total_count / total_words) * 100 if total_words > 0 else 0
                    
                    target_analysis[keyword_phrase] = {
                        "freq": total_count,
                        "density": density,
                        "declensions_found": declension_counts
                    }
                
                else:
                    # Фраза из нескольких слов - генерируем склонения для каждого слова
                    word_declensions = []
                    for word in words_in_phrase:
                        declensions = generate_declensions(word)
                        word_declensions.append(declensions)
                    
                    # Генерируем все возможные комбинации фраз
                    phrase_variations = []
                    
                    def generate_combinations(current_phrase, word_index):
                        if word_index == len(words_in_phrase):
                            phrase_variations.append(current_phrase)
                            return
                        
                        for declension in word_declensions[word_index]:
                            new_phrase = current_phrase + " " + declension if current_phrase else declension
                            generate_combinations(new_phrase, word_index + 1)
                    
                    generate_combinations("", 0)
                    
                    # Подсчитываем вхождения каждой вариации фразы
                    total_count = 0
                    phrase_counts = {}
                    
                    for phrase_variation in phrase_variations:
                        # Используем только точный поиск с границами слов для фраз
                        pattern = r'\b' + re.escape(phrase_variation) + r'\b'
                        matches = re.findall(pattern, text, re.IGNORECASE)
                        
                        # Подсчитываем уникальные вхождения
                        count = len(matches)
                        
                        if count > 0:
                            # Проверяем, что это не дубликат
                            is_duplicate = False
                            for existing_phrase in phrase_counts:
                                if phrase_variation.lower() == existing_phrase.lower():
                                    is_duplicate = True
                                    break
                            
                            if not is_duplicate:
                                phrase_counts[phrase_variation] = count
                                total_count += count
                    
                    # Вычисляем плотность
                    density = (total_count / total_words) * 100 if total_words > 0 else 0
                    
                    target_analysis[keyword_phrase] = {
                        "freq": total_count,
                        "density": density,
                        "declensions_found": phrase_counts
                    }

        return [{"word": word, "freq": freq} for word, freq in keywords], keyword_density, target_analysis
    except Exception as e:
        return [], {}, f"Ошибка анализа ключевых слов: {str(e)}"

def check_open_graph(driver):
    """Проверяет теги Open Graph."""
    og_tags = {
        "og:title": "Не найден",
        "og:description": "Не найден",
        "og:image": "Не найден",
        "og:url": "Не найден",
        "og:type": "Не найден"
    }
    for meta in driver.find_elements(By.TAG_NAME, "meta"):
        property_attr = meta.get_attribute("property")
        if property_attr and property_attr.startswith("og:"):
            content = meta.get_attribute("content") or "Без содержимого"
            og_tags[property_attr] = content
    return og_tags

def check_schema_markup(driver):
    """Проверяет наличие микроразметки Schema.org."""
    schema_tags = driver.find_elements(By.XPATH, "//*[@itemscope]")
    if not schema_tags:
        return False, "Микроразметка Schema.org не найдена"
    valid_types = ["Article", "Product", "Organization", "Person", "WebPage"]
    for tag in schema_tags:
        itemtype = tag.get_attribute("itemtype")
        if itemtype and any(valid_type in itemtype for valid_type in valid_types):
            return True, f"Найдена микроразметка: {itemtype}"
    return False, "Микроразметка присутствует, но не соответствует ожидаемым типам"

def check_noindex_nofollow_noarchive(driver):
    """Проверяет наличие meta тегов noindex, nofollow, noarchive и тега <noindex>."""
    # Проверка meta robots
    meta_robots = driver.find_elements(By.XPATH, "//meta[@name='robots']")
    noindex_meta = False
    nofollow_meta = False
    noarchive_meta = False
    robots_content = "Не найден"
    
    if meta_robots:
        content = meta_robots[0].get_attribute("content").lower()
        noindex_meta = "noindex" in content
        nofollow_meta = "nofollow" in content
        noarchive_meta = "noarchive" in content
        robots_content = content
    
    # Проверка тега <noindex>
    noindex_tag = driver.find_elements(By.XPATH, "//noindex")
    has_noindex_tag = len(noindex_tag) > 0
    
    return noindex_meta, nofollow_meta, noarchive_meta, has_noindex_tag, robots_content

def check_hidden_blocks(driver):
    """Проверяет наличие элементов с display: none."""
    hidden_elements = driver.execute_script("""
        let elements = document.querySelectorAll('[style*="display: none"]');
        let results = [];
        elements.forEach(el => {
            if (el.textContent.trim().length > 0) {
                results.push(el.tagName + ': ' + el.textContent.trim().substring(0, 50));
            }
        });
        return results;
    """)
    return hidden_elements

def check_canonical(driver):
    """Проверяет canonical тег."""
    canonical = driver.find_elements(By.XPATH, "//link[@rel='canonical']")
    if canonical:
        href = canonical[0].get_attribute("href")
        return True, href
    return False, "Не найден"

def check_pagination_links(driver):
    """Проверяет rel=next/prev."""
    next_link = driver.find_elements(By.XPATH, "//link[@rel='next']")
    prev_link = driver.find_elements(By.XPATH, "//link[@rel='prev']")
    return bool(next_link), bool(prev_link)

def check_external_links(driver, site_url):
    """Проверяет внешние ссылки на nofollow и broken."""
    domain = re.sub(r'^https?://(www\.)?', '', site_url).rstrip('/')
    links = driver.find_elements(By.TAG_NAME, "a")
    external_links = [link for link in links if link.get_attribute("href") and domain not in link.get_attribute("href")]
    nofollow_count = sum(1 for link in external_links if "nofollow" in link.get_attribute("rel").lower())
    broken_externals = []
    for link in external_links:
        href = link.get_attribute("href")
        if href and "javascript:void" not in href:
            _, status, _, _ = check_resource(href, True)
            if not isinstance(status, int) or status != 200:
                broken_externals.append(href)
    return len(external_links), nofollow_count, broken_externals

def check_mirrors_and_redirects(site_url, ignore_ssl):
    """Проверяет зеркала и редиректы (www/non-www, http/https)."""
    variants = [
        site_url.replace("https://", "http://"),
        site_url.replace("https://www.", "https://") if "www." in site_url else site_url.replace("https://", "https://www."),
        site_url.replace("https://", "http://www.") if not "www." in site_url else site_url.replace("https://www.", "http://www."),
    ]
    issues = []
    for var in set(variants):  # Убрать дубли
        if var == site_url:
            continue
        _, status, _, history = check_resource(var, ignore_ssl)
        if not history or history[0].status_code != 301 or status != 200 or history[-1].url != site_url:
            issues.append(f"Проблема с зеркалом {var}: статус {status}, финальный URL {history[-1].url if history else 'нет редиректа'}")
    return issues

def check_redirect_chain(site_url, ignore_ssl):
    """Проверяет цепочки редиректов."""
    _, _, _, history = check_resource(site_url, ignore_ssl)
    if len(history) > 3:
        return True, len(history)
    return False, 0

def check_duplicates(driver):
    """Базовая проверка дубликатов (canonical на себя)."""
    has_canonical, href = check_canonical(driver)
    if has_canonical and href == driver.current_url:
        return False, "Canonical указывает на себя"
    elif has_canonical:
        return True, f"Canonical указывает на другой URL: {href}"
    return True, "Canonical отсутствует"

def check_ads(driver):
    """Базовая проверка рекламы (кол-во iframes/скриптов от ad сетей)."""
    ad_iframes = len(driver.find_elements(By.XPATH, "//iframe[contains(@src, 'googleads') or contains(@src, 'doubleclick')]"))
    ad_scripts = len(driver.find_elements(By.XPATH, "//script[contains(@src, 'ads') or contains(@src, 'doubleclick')]"))
    total_ads = ad_iframes + ad_scripts
    if total_ads > 5:
        return True, total_ads
    return False, total_ads

def check_security(driver, site_url):
    """Базовая проверка безопасности (HTTPS, mixed content)."""
    issues = []
    if not site_url.startswith("https"):
        issues.append("Сайт не использует HTTPS")
    html = driver.page_source
    mixed = re.findall(r'src="http://', html)
    if mixed:
        issues.append(f"Найдено {len(mixed)} mixed content (http на https странице)")
    return issues

def get_background_images(html):
    """Извлекает background-image из CSS в HTML."""
    soup = BeautifulSoup(html, 'html.parser')
    bg_images = []
    for element in soup.find_all(style=True):
        style = element['style']
        match = re.search(r'background-image:\s*url\(["\']?(.*?)["\']?\)', style)
        if match:
            bg_images.append(match.group(1))
    return bg_images

def generate_performance_chart(load_times, resource_times, js_css_times):
    """Генерирует график производительности загрузки."""
    plt.figure(figsize=(8, 5))
    resolutions = ["1920x1080", "768x1024", "375x667"]
    bar_width = 0.25
    index = range(len(resolutions))

    if len(load_times) != len(resource_times) or len(load_times) != len(js_css_times):
        raise ValueError("Размеры данных load_times, resource_times и js_css_times должны совпадать")

    plt.bar(index, load_times, bar_width, label="Общая загрузка", color='b')
    plt.bar([i + bar_width for i in index], resource_times, bar_width, label="Ресурсы", color='g')
    plt.bar([i + 2 * bar_width for i in index], js_css_times, bar_width, label="JS/CSS", color='r')

    plt.xlabel("Разрешение")
    plt.ylabel("Время (сек)")
    plt.title("Производительность загрузки")
    plt.xticks([i + bar_width for i in index], resolutions)
    plt.legend()
    buffer = BytesIO()
    plt.savefig(buffer, format="png")
    plt.close()
    buffer.seek(0)
    return base64.b64encode(buffer.getvalue()).decode()

def analyze_performance(site_url, ignore_ssl):
    """Анализирует производительность сайта для разных разрешений."""
    performance_data = {
        "load_times": [],
        "resource_times": [],
        "js_css_times": []
    }
    resolutions = [(1920, 1080), (768, 1024), (375, 667)]
    for width, height in resolutions:
        try:
            driver_local = create_webdriver(ignore_ssl=ignore_ssl, window_size=f"{width},{height}")
        except Exception as e:
            log_to_file(f"Ошибка создания WebDriver для анализа производительности: {str(e)}")
            performance_data["load_times"].append(0)
            performance_data["resource_times"].append(0)
            performance_data["js_css_times"].append(0)
            continue
        try:
            start_time = time.time()
            driver_local.get(site_url)
            load_time = time.time() - start_time
            performance_data["load_times"].append(load_time)

            # Время загрузки ресурсов
            resources = driver_local.execute_script("""
                return window.performance.getEntriesByType("resource");
            """)
            resource_time = sum(entry['duration'] / 1000 for entry in resources) / len(resources) if resources else 0
            performance_data["resource_times"].append(resource_time)

            # Время выполнения JS/CSS
            js_css_time = driver_local.execute_script("""
                let times = 0;
                performance.getEntriesByType("resource").forEach(entry => {
                    if (entry.name.endsWith('.js') || entry.name.endsWith('.css')) {
                        times += entry.duration / 1000;
                    }
                });
                return times;
            """)
            performance_data["js_css_times"].append(js_css_time)
        except Exception as e:
            log_to_file(f"Ошибка анализа производительности для {width}x{height}: {str(e)}")
            performance_data["load_times"].append(0)
            performance_data["resource_times"].append(0)
            performance_data["js_css_times"].append(0)
        finally:
            try:
                if 'driver_local' in locals() and driver_local:
                    driver_local.quit()
            except Exception as e:
                log_to_file(f"Ошибка закрытия WebDriver в analyze_performance: {str(e)}")
    return performance_data

def format_summary_section(positives, errors, recommendations, title):
    """Форматирует раздел сводки в читаемый вид."""
    output = f"### {title}\n\n"
    if positives:
        output += "**Хорошее:**\n" + "\n".join(f"✅ {p}" for p in positives) + "\n\n"
    if errors:
        output += "**Проблемы:**\n" + "\n".join(f"❌ {e}" for e in errors) + "\n\n"
    if recommendations:
        output += "**Рекомендации:**\n" + "\n".join(f"📝 {r}" for r in recommendations) + "\n\n"
    return output

def format_links_section(link_statuses):
    """Форматирует список ссылок с статусами и красивыми индикаторами."""
    output = f"### Ссылки ({len(link_statuses)} проверенных)\n\n"
    for url, status in link_statuses.items():
        if isinstance(status, int):
            if status == 200:
                status_icon = "🟢"
                status_text = f"200 OK"
            elif status == 301:
                status_icon = "🟡"
                status_text = f"301 Moved Permanently"
            elif status == 302:
                status_icon = "🟡"
                status_text = f"302 Found"
            elif status == 404:
                status_icon = "🔴"
                status_text = f"404 Not Found"
            elif status == 500:
                status_icon = "🔴"
                status_text = f"500 Internal Server Error"
            elif 300 <= status < 400:
                status_icon = "🟡"
                status_text = f"{status} Redirect"
            elif 400 <= status < 500:
                status_icon = "🔴"
                status_text = f"{status} Client Error"
            elif 500 <= status < 600:
                status_icon = "🔴"
                status_text = f"{status} Server Error"
            else:
                status_icon = "⚪"
                status_text = f"{status} Unknown"
        else:
            status_icon = "🔴"
            status_text = f"Error: {status}"
        
        output += f"{status_icon} **{status_text}** - {url}\n"
    return output

def check_links_summary(link_statuses):
    """Создает сводку ссылок по статусам с красивым форматированием."""
    categories = {
        "🟢 200 OK": [],
        "🟡 301 Moved Permanently": [],
        "🟡 302 Found": [],
        "🟡 Other Redirects (3xx)": [],
        "🔴 404 Not Found": [],
        "🔴 500 Internal Server Error": [],
        "🔴 Other Client Errors (4xx)": [],
        "🔴 Other Server Errors (5xx)": [],
        "⚪ Other/Errors": []
    }
    
    for url, status in link_statuses.items():
        if isinstance(status, int):
            if status == 200:
                categories["🟢 200 OK"].append(url)
            elif status == 301:
                categories["🟡 301 Moved Permanently"].append(url)
            elif status == 302:
                categories["🟡 302 Found"].append(url)
            elif 300 <= status < 400:
                categories["🟡 Other Redirects (3xx)"].append(url)
            elif status == 404:
                categories["🔴 404 Not Found"].append(url)
            elif status == 500:
                categories["🔴 500 Internal Server Error"].append(url)
            elif 400 <= status < 500:
                categories["🔴 Other Client Errors (4xx)"].append(url)
            elif 500 <= status < 600:
                categories["🔴 Other Server Errors (5xx)"].append(url)
            else:
                categories["⚪ Other/Errors"].append(url)
        else:
            categories["⚪ Other/Errors"].append(url)

    summary = f"### 📊 Сводка Ссылок ({len(link_statuses)} проверенных)\n\n"
    for cat, urls in categories.items():
        if urls:
            summary += f"**{cat}** ({len(urls)} ссылок):\n"
            for url in urls:
                summary += f"  • {url}\n"
            summary += "\n"
    
    return summary
    return summary

def run_test(site_url: str, summary_area: ft.TextField, page: ft.Page, progress_bar: ft.ProgressBar, ignore_ssl: bool, target_keywords: str):
    """Запускает тестирование сайта."""
    if not re.match(r'^https?://', site_url):
        summary_area.value = "❌ Неверный URL\n"
        page.update()
        return

    # Проверяем событие остановки
    stop_event = page.data.get('stop_event')
    if stop_event and stop_event.is_set():
        summary_area.value = "⏹ Тест остановлен пользователем"
        page.update()
        return

    driver = None
    general_positives = []
    general_errors = []
    general_recs = []
    seo_positives = []
    seo_errors = []
    seo_recs = []
    perf_positives = []
    perf_errors = []
    perf_recs = []
    broken_links = []
    link_statuses = {}  # Для вкладки Ссылки
    sitemap_errors = []  # Для проверки sitemap в сводке
    urls_in_sitemap = []
    site_links = []
    total_checks = 32  # Увеличено для расширенной проверки
    current_check = 0
    log_text = ""  # Внутренняя переменная для логов

    # --- Новый блок: Получаем sitemap-ы из robots.txt и обрабатываем рекурсивно ---
    sitemap_urls = []
    try:
        robots_url = site_url.rstrip('/') + '/robots.txt'
        r = requests.get(robots_url, timeout=10, verify=not ignore_ssl)
        if r.status_code == 200:
            for line in r.text.splitlines():
                if line.strip().lower().startswith('sitemap:'):
                    sitemap_url = line.split(':', 1)[1].strip()
                    if sitemap_url:
                        sitemap_urls.append(sitemap_url)
    except Exception as ex:
        pass  # robots.txt может отсутствовать — это не критично
    if not sitemap_urls:
        sitemap_urls = [site_url.rstrip('/') + '/sitemap.xml']
    
    # Собираем все ссылки из всех sitemap с поддержкой sitemap index
    all_sitemap_links = []
    sitemap_sources = {}
    sitemap_errors = []
    
    for sitemap_url in sitemap_urls:
        try:
            result = process_sitemap_recursively(sitemap_url, ignore_ssl)
            all_sitemap_links.extend(result['urls'])
            sitemap_sources.update(result['sources'])
            sitemap_errors.extend(result['errors'])
        except Exception as ex:
            sitemap_errors.append(f"Ошибка обработки {sitemap_url}: {str(ex)}")
    
    # Убираем дубликаты и пустые значения
    all_sitemap_links = list(set([link for link in all_sitemap_links if link]))
    urls_in_sitemap = all_sitemap_links
    
    if sitemap_errors:
        log_text += f"⚠️ Ошибки при обработке sitemap: {', '.join(sitemap_errors[:3])}{'...' if len(sitemap_errors) > 3 else ''}\n"
    # --- Конец нового блока ---

    def update_progress():
        nonlocal current_check
        current_check += 1
        progress_bar.value = min(current_check / total_checks, 1.0)
        page.update()
        
        # Проверяем событие остановки
        if stop_event and stop_event.is_set():
            return False  # Сигнал остановки
        return True  # Продолжаем

    try:
        log_text += f"\n=== 🧪 Тестирование сайта: {site_url} ===\n"
        general_positives.append(f"🔍 Тестирование сайта: {site_url} ({datetime.now().strftime('%Y-%m-%d %H:%M:%S %Z')})")
        if not update_progress():
            return

        # Проверка доступности сайта с обходом блокировок
        log_text += "\n🔍 Проверка доступности сайта с обходом блокировок...\n"
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        
        # Сначала пробуем обычный метод
        anti_bot_mode = False
        try:
            start_time = time.time()
            r = requests.get(site_url, timeout=10, verify=not ignore_ssl, allow_redirects=True)
            load_time = time.time() - start_time
            page_size = len(r.content) / 1024
            log_text += f"🔎 HTTP статус: {r.status_code}\n"
            log_text += f"⏱ Время загрузки (HTTP): {load_time:.2f} сек\n"
            log_text += f"📏 Размер страницы: {page_size:.2f} КБ\n"
            
            # Проверяем на блокировку
            if any(blocked_text in r.text.lower() for blocked_text in [
                'access denied', 'доступ запрещен', 'blocked', 'заблокирован',
                'cloudflare', 'captcha', 'recaptcha', 'bot', 'robot'
            ]):
                log_text += "⚠️ Обнаружена блокировка, пробуем обход...\n"
                general_errors.append("Сайт блокирует доступ (обнаружена защита от ботов)")
                anti_bot_mode = True
            else:
                general_positives.append(f"Сайт доступен (HTTP статус: {r.status_code}, время загрузки: {load_time:.2f} сек)")
                general_positives.append(f"Размер страницы: {page_size:.2f} КБ")
                anti_bot_mode = False
                
            if page_size > 2000:
                general_errors.append("Страница слишком тяжёлая")
                general_recs.append("Сожмите изображения, минимизируйте CSS/JS, используйте lazy loading.")
            if load_time > 3:
                general_errors.append("Время загрузки слишком большое")
                general_recs.append("Включите кэширование, используйте CDN, оптимизируйте сервер.")
            if r.status_code != 200:
                seo_errors.append("Проблема индексации: статус не 200")
                seo_recs.append("Исправьте статус страницы для индексации.")
                
        except Exception as e:
            log_text += f"❌ Ошибка подключения: {e}\n"
            log_text += "🔄 Пробуем обход блокировок...\n"
            anti_bot_mode = True
            general_errors.append(f"Сайт недоступен: {e}")
            seo_errors.append("Проблема несуществующей страницы или индексации")
        
        update_progress()

        # Создание основного драйвера для SEO проверок с обходом блокировок
        try:
            if anti_bot_mode:
                log_text += "🛡️ Создание WebDriver в режиме обхода блокировок...\n"
                driver = create_webdriver(ignore_ssl=ignore_ssl, anti_bot_mode=True)
            else:
                driver = create_webdriver(ignore_ssl=ignore_ssl)
            
            # Пробуем загрузить страницу
            try:
                driver.get(site_url)
                log_text += "✅ Страница успешно загружена\n"
            except Exception as e:
                log_text += f"⚠️ Ошибка загрузки страницы: {e}\n"
                if not anti_bot_mode:
                    log_text += "🔄 Пробуем в режиме обхода блокировок...\n"
                    driver.quit()
                    driver = create_webdriver(ignore_ssl=ignore_ssl, anti_bot_mode=True)
                    driver.get(site_url)
                    log_text += "✅ Страница загружена в режиме обхода\n"
                
        except Exception as e:
            log_to_file(f"Ошибка создания WebDriver: {str(e)}")
            summary_area.value = f"❌ Ошибка: Не удалось создать WebDriver: {str(e)}"
            page.update()
            return

        # Получение версии ChromeDriver
        driver_version = driver.capabilities['chrome']['chromedriverVersion'].split(' ')[0]
        log_text += f"📌 Версия ChromeDriver: {driver_version}\n"
        general_positives.append(f"Версия ChromeDriver: {driver_version}")

        # Проверка зеркал и стандарта URL
        log_text += "\n🔄 Проверка зеркал и стандарта URL\n"
        mirror_issues = check_mirrors_and_redirects(site_url, ignore_ssl)
        if mirror_issues:
            seo_errors.extend(mirror_issues)
            seo_recs.append("Настройте правильные 301 редиректы для зеркал (www/non-www, http/https).")
        else:
            seo_positives.append("Зеркала настроены правильно")
        update_progress()

        # Проверка цепочек редиректов
        log_text += "\n🔄 Проверка цепочек редиректов\n"
        has_long_chain, chain_len = check_redirect_chain(site_url, ignore_ssl)
        if has_long_chain:
            seo_errors.append(f"Длинная цепочка редиректов ({chain_len} шагов)")
            seo_recs.append("Сократите цепочку редиректов до 1-2 шагов.")
        else:
            seo_positives.append("Цепочки редиректов в норме")
        update_progress()

        # Проверка SEO-файлов
        log_text += "\n🤖 Проверка SEO-файлов\n"
        seo_files = check_seo_files(site_url, ignore_ssl)
        for file, status, content in seo_files:
            log_text += f"📋 Проверка {file}\n"
            if status:
                log_text += f"✅ {file} доступен\n"
                seo_positives.append(f"{file} найден")
                if file == "robots.txt":
                    errors, positives, found_directives, recommendations = analyze_robots_txt(content)
                    log_text += f"📋 Найденные директивы: {', '.join(found_directives) if found_directives else 'Нет директив'}\n"
                    seo_positives.extend(positives)
                    seo_errors.extend(errors)
                    seo_recs.extend(recommendations)
                else:
                    errors, positives, recommendations, _, page_details, sitemap_info, pages_not_in_sitemap, pages_in_sitemap_not_on_site = validate_sitemap(content, site_url, ignore_ssl)
                    seo_positives.extend(positives)
                    seo_errors.extend(errors)
                    sitemap_errors = errors  # Сохраняем для сводки
                    seo_recs.extend(recommendations)
                    log_text += f"📋 Содержимое sitemap.xml: {content[:200]}...\n"
                    seo_positives.append("sitemap.xml содержит данные")
            else:
                log_text += f"❌ {file} недоступен: {content}\n"
                seo_errors.append(f"{file} не найден")
                seo_recs.append(f"Создайте {file}.")
            update_progress()

        # Анализ производительности
        log_text += "\n⏱ Анализ производительности\n"
        performance_data = analyze_performance(site_url, ignore_ssl)
        load_times = performance_data["load_times"]
        resource_times = performance_data["resource_times"]
        js_css_times = performance_data["js_css_times"]

        log_text += f"⏱ Общее время загрузки: {', '.join(f'{t:.2f} сек' for t in load_times)}\n"
        log_text += f"⏱ Время загрузки ресурсов: {', '.join(f'{t:.2f} сек' for t in resource_times)}\n"
        log_text += f"⏱ Время выполнения JS/CSS: {', '.join(f'{t:.2f} сек' for t in js_css_times)}\n"
        perf_positives.append(f"Общее время загрузки: {', '.join(f'{t:.2f} сек' for t in load_times)}")
        perf_positives.append(f"Время загрузки ресурсов: {', '.join(f'{t:.2f} сек' for t in resource_times)}")
        perf_positives.append(f"Время выполнения JS/CSS: {', '.join(f'{t:.2f} сек' for t in js_css_times)}")
        for i, (lt, rt, jct) in enumerate(zip(load_times, resource_times, js_css_times)):
            res = ['1920x1080', '768x1024', '375x667'][i]
            if lt > 3:
                perf_errors.append(f"Слишком долгое время загрузки на {res}: {lt:.2f} сек")
                perf_recs.append("Оптимизируйте сервер и кэширование.")
            else:
                perf_positives.append(f"Нормальное время загрузки на {res}: {lt:.2f} сек")
            if rt > 1:
                perf_errors.append(f"Долгое время загрузки ресурсов на {res}: {rt:.2f} сек")
                perf_recs.append("Сожмите ресурсы, используйте CDN.")
            if jct > 1:
                perf_errors.append(f"Долгое время выполнения JS/CSS на {res}: {jct:.2f} сек")
                perf_recs.append("Минимизируйте и асинхронно загружайте скрипты.")
        update_progress()

        # Core Web Vitals
        log_text += "\n⚡ Core Web Vitals\n"
        try:
            lcp, fid, cls = get_core_web_vitals(driver)
            core_vitals_summary = f"LCP: {lcp:.2f} мс | FID: {fid:.2f} мс | CLS: {cls:.3f}\n"
            if lcp > 2500:
                perf_errors.append(f"LCP слишком большой: {lcp:.0f} мс")
                perf_recs.append("Сократите время Largest Contentful Paint до <2.5 сек.")
            if cls > 0.1:
                perf_errors.append(f"CLS слишком большой: {cls:.3f}")
                perf_recs.append("Снизьте Cumulative Layout Shift до <0.1.")
            log_text += core_vitals_summary
            perf_positives.append(core_vitals_summary)
        except Exception as e:
            log_to_file(f"Ошибка получения Core Web Vitals: {str(e)}")
            perf_errors.append("Не удалось получить Core Web Vitals")
        update_progress()

        # Микроразметка
        log_text += "\n🔎 Микроразметка (Schema.org, JSON-LD, OpenGraph, Twitter)\n"
        try:
            schema_items, jsonld_blocks, og_tags, twitter_tags = get_microdata(driver)
            if schema_items:
                seo_positives.append(f"Schema.org items: {', '.join(schema_items)}")
            else:
                seo_errors.append("Schema.org не найдена")
                seo_recs.append("Добавьте Schema.org микроразметку.")
            if jsonld_blocks:
                seo_positives.append(f"JSON-LD блоков: {len(jsonld_blocks)}")
            else:
                seo_errors.append("JSON-LD не найден")
                seo_recs.append("Добавьте JSON-LD для расширенного описания сайта.")
            if og_tags:
                seo_positives.append(f"OpenGraph: {', '.join(og_tags.keys())}")
            else:
                seo_errors.append("OpenGraph не найден")
                seo_recs.append("Добавьте OpenGraph для соцсетей.")
            if twitter_tags:
                seo_positives.append(f"Twitter Cards: {', '.join(twitter_tags.keys())}")
            else:
                seo_errors.append("Twitter Cards не найдены")
                seo_recs.append("Добавьте Twitter Cards для соцсетей.")
        except Exception as e:
            log_to_file(f"Ошибка получения микроразметки: {str(e)}")
            seo_errors.append("Не удалось получить микроразметку")
        update_progress()

        # Проверка SEO-элементов
        try:
            title_tag = driver.title if driver.title else "Не найден"
            log_text += f"📝 Тег title: {title_tag} (Длина: {len(title_tag)})\n"
            seo_positives.append(f"Тег title: {title_tag} (Длина: {len(title_tag)})")
            if len(title_tag) > 60:
                seo_errors.append("Длина title превышает 60 символов")
                seo_recs.append("Сократите title до 60 символов.")
            elif title_tag == "Не найден":
                seo_errors.append("Тег title отсутствует")
                seo_recs.append("Добавьте тег title.")
            else:
                seo_positives.append("Тег title в норме")
        except Exception as e:
            log_to_file(f"Ошибка получения title: {str(e)}")
            seo_errors.append("Не удалось получить title")
        update_progress()

        try:
            meta_desc = driver.find_elements(By.XPATH, "//meta[@name='description']")
            if meta_desc:
                desc_content = meta_desc[0].get_attribute("content")
                log_text += f"📝 Мета-описание: {desc_content} (Длина: {len(desc_content)})\n"
                seo_positives.append(f"Мета-описание: {desc_content} (Длина: {len(desc_content)})")
                if len(desc_content) > 160:
                    seo_errors.append("Длина мета-описания превышает 160 символов")
                    seo_recs.append("Сократите до 160 символов.")
            else:
                seo_errors.append("Мета-описание отсутствует")
                seo_recs.append("Добавьте мета-описание.")
        except Exception as e:
            log_to_file(f"Ошибка получения мета-описания: {str(e)}")
            seo_errors.append("Не удалось получить мета-описание")
        update_progress()

        # Проверка заголовков
        try:
            h1_tags = driver.find_elements(By.TAG_NAME, "h1")
            if len(h1_tags) == 1:
                seo_positives.append("Один H1 найден")
            elif len(h1_tags) > 1:
                seo_errors.append(f"Найдено {len(h1_tags)} тегов H1")
                seo_recs.append("Используйте только один H1.")
            else:
                seo_errors.append("Тег H1 отсутствует")
                seo_recs.append("Добавьте тег H1.")
        except Exception as e:
            log_to_file(f"Ошибка проверки заголовков: {str(e)}")
            seo_errors.append("Не удалось проверить заголовки")
        update_progress()

        # Проверка Open Graph
        log_text += "\n🌐 Проверка Open Graph\n"
        try:
            og_tags = check_open_graph(driver)
            for tag, content in og_tags.items():
                log_text += f"  - {tag}: {content}\n"
                seo_positives.append(f"{tag}: {content}")
                if "Не найден" in content or "Без содержимого" in content:
                    seo_errors.append(f"Отсутствует тег Open Graph: {tag}")
                    seo_recs.append(f"Добавьте {tag}.")
        except Exception as e:
            log_to_file(f"Ошибка проверки Open Graph: {str(e)}")
            seo_errors.append("Не удалось проверить Open Graph")
        update_progress()

        # Проверка микроразметки
        log_text += "\n📋 Проверка микроразметки Schema.org\n"
        try:
            has_schema, schema_result = check_schema_markup(driver)
            log_text += f"  - {schema_result}\n"
            if has_schema:
                seo_positives.append(schema_result)
            else:
                seo_errors.append(schema_result)
                seo_recs.append("Добавьте микроразметку Schema.org.")
        except Exception as e:
            log_to_file(f"Ошибка проверки микроразметки: {str(e)}")
            seo_errors.append("Не удалось проверить микроразметку")
        update_progress()

        # Проверка noindex, nofollow, noarchive и тега <noindex>
        log_text += "\n🤖 Проверка meta robots и тега <noindex>\n"
        try:
            noindex_meta, nofollow_meta, noarchive_meta, has_noindex_tag, robots_content = check_noindex_nofollow_noarchive(driver)
            if robots_content != "Не найден":
                log_text += f"📝 Meta robots content: {robots_content}\n"
                seo_positives.append(f"Meta robots: {robots_content}")
                if noindex_meta:
                    seo_errors.append("Страница имеет noindex в meta robots (не индексируется)")
                    seo_recs.append("Удалите noindex, если страница должна индексироваться.")
                if nofollow_meta:
                    seo_errors.append("Страница имеет nofollow в meta robots (ссылки не следуются)")
                    seo_recs.append("Удалите nofollow, если ссылки должны следовать.")
                if noarchive_meta:
                    seo_errors.append("Страница имеет noarchive в meta robots (нет сохраненной копии)")
                    seo_recs.append("Удалите noarchive, если нужна сохраненная копия.")
            else:
                seo_positives.append("Meta robots не найден (по умолчанию index, follow, archive)")
            
            # Проверка тега <noindex>
            if has_noindex_tag:
                seo_errors.append("Найден тег <noindex> (контент не индексируется)")
                seo_recs.append("Удалите тег <noindex>, если контент должен индексироваться.")
            else:
                seo_positives.append("Тег <noindex> не найден (контент может индексироваться)")
        except Exception as e:
            log_to_file(f"Ошибка проверки meta robots и тега <noindex>: {str(e)}")
            seo_errors.append("Не удалось проверить meta robots и тег <noindex>")
        update_progress()

        # Проверка скрытых блоков display: none
        log_text += "\n🕵️ Проверка скрытых блоков (display: none)\n"
        try:
            hidden_blocks = check_hidden_blocks(driver)
            if hidden_blocks:
                log_text += f"❌ Найдено {len(hidden_blocks)} скрытых элементов с текстом:\n"
                for block in hidden_blocks[:5]:  # Ограничим вывод первыми 5
                    log_text += f"  - {block}...\n"
                seo_errors.append(f"Найдено {len(hidden_blocks)} скрытых элементов (display: none) с текстом")
                seo_recs.append("Проверьте скрытый контент на спам; поисковики могут penalize за скрытый текст.")
            else:
                seo_positives.append("Нет скрытых элементов с display: none содержащих текст")
        except Exception as e:
            log_to_file(f"Ошибка проверки скрытых блоков: {str(e)}")
            seo_errors.append("Не удалось проверить скрытые блоки")
        update_progress()

        # Проверка canonical
        log_text += "\n📌 Проверка canonical\n"
        try:
            has_canonical, canonical_href = check_canonical(driver)
            if has_canonical:
                seo_positives.append(f"Canonical найден: {canonical_href}")
                if canonical_href != site_url:
                    seo_errors.append("Canonical указывает на другой URL (возможный дубликат)")
                    seo_recs.append("Убедитесь, что canonical указывает на предпочтительную версию.")
            else:
                seo_errors.append("Canonical отсутствует")
                seo_recs.append("Добавьте canonical для предотвращения дубликатов.")
        except Exception as e:
            log_to_file(f"Ошибка проверки canonical: {str(e)}")
            seo_errors.append("Не удалось проверить canonical")
        update_progress()

        # Проверка пагинации
        log_text += "\n📄 Проверка пагинации (rel=next/prev)\n"
        try:
            has_next, has_prev = check_pagination_links(driver)
            if has_next or has_prev:
                seo_positives.append(f"Пагинация найдена: next={has_next}, prev={has_prev}")
            else:
                seo_positives.append("Пагинация не найдена (возможно, не требуется)")
            # Для пагинации проверить canonical
            if has_next or has_prev and has_canonical and canonical_href == site_url:
                seo_positives.append("Canonical на пагинированной странице указывает на первую страницу")
            elif has_next or has_prev:
                seo_errors.append("Проблема пагинации: canonical не указывает на первую страницу")
                seo_recs.append("Для пагинированных страниц используйте self-canonical или на первую.")
        except Exception as e:
            log_to_file(f"Ошибка проверки пагинации: {str(e)}")
            seo_errors.append("Не удалось проверить пагинацию")
        update_progress()

        # Проверка внешних ссылок
        log_text += "\n🔗 Проверка внешних ссылок\n"
        try:
            ext_count, nofollow_count, broken_ext = check_external_links(driver, site_url)
            seo_positives.append(f"Внешних ссылок: {ext_count}, с nofollow: {nofollow_count}")
            if broken_ext:
                seo_errors.append(f"Битые внешние ссылки: {len(broken_ext)} ({', '.join(broken_ext[:5])}...)")
                seo_recs.append("Исправьте битые внешние ссылки.")
            if nofollow_count < ext_count / 2:
                seo_errors.append("Мало nofollow на внешних ссылках")
                seo_recs.append("Добавьте nofollow на внешние ссылки, если они не доверенные.")
        except Exception as e:
            log_to_file(f"Ошибка проверки внешних ссылок: {str(e)}")
            seo_errors.append("Не удалось проверить внешние ссылки")
        update_progress()

        # Проверка дубликатов (базовая)
        log_text += "\n🔍 Проверка дубликатов\n"
        try:
            is_duplicate, dup_msg = check_duplicates(driver)
            if is_duplicate:
                seo_errors.append(f"Проблема дубликатов: {dup_msg}")
                seo_recs.append("Используйте canonical для указания основной версии.")
            else:
                seo_positives.append(dup_msg)
        except Exception as e:
            log_to_file(f"Ошибка проверки дубликатов: {str(e)}")
            seo_errors.append("Не удалось проверить дубликаты")
        update_progress()

        # Проверка рекламы
        log_text += "\n📢 Проверка рекламы\n"
        try:
            has_ads_issue, ads_count = check_ads(driver)
            if has_ads_issue:
                seo_errors.append(f"Много рекламы ({ads_count} элементов)")
                seo_recs.append("Сократите количество рекламы, чтобы не ухудшать UX и SEO.")
            else:
                seo_positives.append("Реклама в норме")
        except Exception as e:
            log_to_file(f"Ошибка проверки рекламы: {str(e)}")
            seo_errors.append("Не удалось проверить рекламу")
        update_progress()

        # Проверка безопасности
        log_text += "\n🛡️ Проверка угроз безопасности\n"
        try:
            security_issues = check_security(driver, site_url)
            if security_issues:
                seo_errors.extend(security_issues)
                seo_recs.append("Перейдите на HTTPS, исправьте mixed content.")
            else:
                seo_positives.append("Безопасность в норме")
        except Exception as e:
            log_to_file(f"Ошибка проверки безопасности: {str(e)}")
            seo_errors.append("Не удалось проверить безопасность")
        update_progress()

        # Проверка изображений
        log_text += "\n🖼 Проверка изображений\n"
        try:
            images = driver.find_elements(By.TAG_NAME, "img")
            html = driver.page_source
            bg_images = get_background_images(html)
            total_images = len(images) + len(bg_images)
            log_text += f"🖼 Количество изображений (img + bg): {total_images}\n"
            images_list = []
            if total_images == 0:
                seo_positives.append("На странице нет изображений")
            else:
                seo_positives.append(f"Изображений: {total_images}")
                images_no_alt = []
                images_no_title = []
                large_images = []
                for img in images:
                    alt = img.get_attribute("alt")
                    title = img.get_attribute("title")
                    src = img.get_attribute("src") or "No src"
                    size_kb = get_image_size(src, ignore_ssl)
                    images_list.append({'src': src, 'alt': alt, 'title': title, 'size': size_kb})
                    if not alt:
                        images_no_alt.append(src)
                    if not title:
                        images_no_title.append(src)
                    if size_kb > 300:
                        large_images.append(src)
                for bg in bg_images:
                    alt = ""  # bg images usually don't have alt/title
                    title = ""
                    size_kb = get_image_size(bg, ignore_ssl)
                    images_list.append({'src': bg, 'alt': alt, 'title': title, 'size': size_kb})
                    images_no_alt.append(bg)  # Since no alt
                    images_no_title.append(bg)  # Since no title
                    if size_kb > 300:
                        large_images.append(bg)
                if images_no_alt:
                    log_text += f"❌ Изображения без alt ({len(images_no_alt)}):\n" + "\n".join(f"- {s}" for s in images_no_alt) + "\n"
                    seo_errors.append(f"{len(images_no_alt)} изображений без alt")
                    seo_recs.append("Добавьте alt атрибуты к изображениям для улучшения SEO и доступности.")
                else:
                    seo_positives.append("Все изображения имеют alt")
                if images_no_title:
                    log_text += f"⚠️ Изображения без title ({len(images_no_title)}):\n" + "\n".join(f"- {s}" for s in images_no_title) + "\n"
                    seo_errors.append(f"{len(images_no_title)} изображений без title")
                    seo_recs.append("Рассмотрите добавление title атрибутов для подсказок.")
                else:
                    seo_positives.append("Все изображения имеют title")
                if large_images:
                    log_text += f"❌ Крупные изображения (>300 КБ) ({len(large_images)}):\n" + "\n".join(f"- {s}" for s in large_images) + "\n"
                    seo_errors.append(f"{len(large_images)} изображений >300 КБ")
                    seo_recs.append("Оптимизируйте изображения для уменьшения размера.")
        except Exception as e:
            log_to_file(f"Ошибка проверки изображений: {str(e)}")
            seo_errors.append("Не удалось проверить изображения")
        update_progress()

        # Проверка ключевых слов
        if target_keywords and target_keywords.strip():
            log_text += "\n🔍 АНАЛИЗ ЦЕЛЕВЫХ КЛЮЧЕВЫХ СЛОВ\n"
            try:
                keywords, density, target_analysis = analyze_keywords(driver, site_url, target_keywords)
                
                if isinstance(target_analysis, dict) and target_analysis:
                    log_text += "=" * 60 + "\n"
                    for tkw, data in target_analysis.items():
                        log_text += f"🎯 ЦЕЛЕВОЕ КЛЮЧЕВОЕ СЛОВО: '{tkw}'\n"
                        log_text += f"📊 ОБЩАЯ ЧАСТОТА (со склонениями): {data['freq']} раз\n"
                        # Безопасное форматирование плотности
                        density_value = data['density']
                        if isinstance(density_value, (int, float)):
                            log_text += f"📈 ПЛОТНОСТЬ: {density_value:.2%}\n"
                        else:
                            log_text += f"📈 ПЛОТНОСТЬ: {density_value}\n"
                        
                        # Показываем найденные склонения
                        if 'declensions_found' in data and data['declensions_found']:
                            log_text += "📝 НАЙДЕННЫЕ СКЛОНЕНИЯ:\n"
                            for declension, count in data['declensions_found'].items():
                                if count > 0:
                                    log_text += f"  ✅ '{declension}': {count} раз\n"
                            
                            # Показываем полный текст с подсветкой
                            log_text += "\n📄 ПОЛНЫЙ ТЕКСТ С ПОДСВЕТКОЙ:\n"
                            html = driver.page_source
                            soup = BeautifulSoup(html, 'html.parser')
                            text = soup.get_text(separator=' ', strip=True)
                            
                            # Подсвечиваем найденные склонения
                            highlighted_text = text
                            for declension, count in data['declensions_found'].items():
                                if count > 0:
                                    # Заменяем найденные склонения на подсвеченные версии
                                    pattern = r'\b' + re.escape(declension) + r'\b'
                                    highlighted_text = re.sub(pattern, f"【{declension}】", highlighted_text, flags=re.IGNORECASE)
                            
                            # Показываем первые 500 символов с подсветкой
                            preview = highlighted_text[:500] + "..." if len(highlighted_text) > 500 else highlighted_text
                            log_text += f"{preview}\n"
                            
                        else:
                            log_text += "❌ Склонения не найдены\n"
                        
                        log_text += "=" * 60 + "\n\n"
                        
                        # Оценка плотности
                        if isinstance(density_value, (int, float)):
                            if density_value < 0.01:
                                seo_errors.append(f"Низкая плотность для целевого '{tkw}' ({density_value:.2%})")
                                seo_recs.append(f"Увеличьте использование '{tkw}' до 1-2%.")
                            elif density_value > 0.03:
                                seo_errors.append(f"Высокая плотность для целевого '{tkw}' ({density_value:.2%})")
                                seo_recs.append(f"Снизьте использование '{tkw}' до 1-2%.")
                            else:
                                seo_positives.append(f"Нормальная плотность для целевого '{tkw}' ({density_value:.2%})")
                        else:
                            seo_errors.append(f"Ошибка расчета плотности для '{tkw}': {density_value}")
                else:
                    log_text += "❌ Целевые ключевые слова не найдены на странице\n"
                    seo_errors.append("Целевые ключевые слова не найдены")
            except Exception as e:
                log_to_file(f"Ошибка анализа ключевых слов: {str(e)}")
                seo_errors.append("Не удалось проанализировать ключевые слова")
        update_progress()

        # Проверка ссылок (битые) и сбор статусов
        try:
            links = driver.find_elements(By.TAG_NAME, "a")
            log_text += f"🔗 Ссылок: {len(links)}\n"
            general_positives.append(f"Ссылок: {len(links)}")
            site_links = [link.get_attribute("href") for link in links if link.get_attribute("href") and site_url in link.get_attribute("href")]
            for i, link in enumerate(links, 1):
                href = link.get_attribute("href") or "Нет href"
                if href and "javascript:void" not in href and not href.startswith("#"):
                    result = check_resource(href, ignore_ssl)
                    href, status, _, _ = result
                    link_statuses[href] = status
                    if not isinstance(status, int) or status != 200:
                        broken_links.append(href)
        except Exception as e:
            log_to_file(f"Ошибка проверки ссылок: {str(e)}")
            general_errors.append("Не удалось проверить ссылки")
        if broken_links:
            general_errors.append(f"Найдены битые ссылки: {len(broken_links)} ({', '.join(broken_links)})")
            general_recs.append("Исправьте битые ссылки.")
        else:
            general_positives.append("Нет битых ссылок")
        # Сравнение с sitemap (используем уже полученные данные)
        if urls_in_sitemap:
            # Получаем страницы сайта для сравнения
            site_pages = get_site_pages(site_url, ignore_ssl)
            
            # Нормализуем URL для корректного сравнения
            def normalize_url(url):
                """Нормализует URL для сравнения"""
                if not url:
                    return url
                url = url.rstrip('/')
                if '#' in url:
                    url = url.split('#')[0]
                if '?' in url:
                    url = url.split('?')[0]
                return url
            
            # Нормализуем все URL
            normalized_site_pages = set(normalize_url(url) for url in site_pages)
            normalized_sitemap_urls = set(normalize_url(url) for url in urls_in_sitemap)
            
            # Страницы на сайте, но не в sitemap
            pages_not_in_sitemap = [url for url in site_pages if normalize_url(url) not in normalized_sitemap_urls]
            
            # Страницы в sitemap, но не на сайте (недоступные)
            # В run_test мы не имеем page_details, поэтому просто сравниваем URL
            pages_in_sitemap_not_on_site = [url for url in urls_in_sitemap if normalize_url(url) not in normalized_site_pages]
            
            if pages_not_in_sitemap:
                seo_errors.append(f"Страницы на сайте не в sitemap: {len(pages_not_in_sitemap)} ({', '.join(pages_not_in_sitemap[:5])}...)")
                seo_recs.append("Добавьте недостающие страницы в sitemap.")
            if pages_in_sitemap_not_on_site:
                seo_errors.append(f"Страницы в sitemap не на сайте: {len(pages_in_sitemap_not_on_site)} ({', '.join(pages_in_sitemap_not_on_site[:5])}...)")
                seo_recs.append("Удалите лишние страницы из sitemap или проверьте сайт.")
        update_progress()

        # Генерация графика производительности
        try:
            chart_base64 = generate_performance_chart(load_times, resource_times, js_css_times)
            page.add(ft.Image(src_base64=chart_base64, width=800, height=500))
        except Exception as e:
            log_to_file(f"Ошибка генерации графика: {str(e)}")
        update_progress()

        log_to_file(f"{site_url} - Успешно протестирован")

        # Форматирование сводок
        seo_area = ft.TextField()
        perf_area = ft.TextField()
        links_area = ft.TextField()
        seo_area.value = format_summary_section(seo_positives, seo_errors, seo_recs, "SEO Анализ")
        perf_area.value = format_summary_section(perf_positives, perf_errors, perf_recs, "Производительность")
        if not sitemap_errors:
            general_positives.append("✅ Sitemap OK")
        full_summary = format_summary_section(general_positives + seo_positives + perf_positives,
                                                    general_errors + seo_errors + perf_errors,
                                                    general_recs + seo_recs + perf_recs,
                                                    "Общая Сводка")
        summary_area.value = full_summary  # Только итоговая сводка!
        page.data['full_summary'] = full_summary
        page.data['seo_summary'] = seo_area.value  # Сохраняем SEO сводку отдельно
        links_area.value = format_links_section(link_statuses)
        links_summary = check_links_summary(link_statuses)
        page.data['links_summary'] = links_summary

        # Форматирование сводки изображений
        images_summary = "### Изображения\n\n"
        if total_images == 0:
            images_summary += "На странице нет изображений\n"
        else:
            for img in images_list:
                alt_emoji = "🟢" if img['alt'] else "🔴"
                title_emoji = "🟢" if img['title'] else "🔴"
                size_emoji = "🟢" if img['size'] <= 300 else "🔴"
                images_summary += f"Ссылка: {img['src']}\n"
                images_summary += f"Alt: {alt_emoji} {img['alt'] or 'Нет'}\n"
                images_summary += f"Title: {title_emoji} {img['title'] or 'Нет'}\n"
                images_summary += f"Размер: {size_emoji} {img['size']:.2f} КБ\n\n"
        page.data['images_summary'] = images_summary

        page.data['robots_summary'] = check_robots_summary(site_url, ignore_ssl)
        sitemap_summary = check_sitemap_summary(site_url, ignore_ssl)
        page.data['sitemap_summary'] = sitemap_summary

        # Сохранение результатов
        save_results(site_url, summary_area.value, full_summary)

    except Exception as e:
        summary_area.value = f"❌ Ошибка: {str(e)}\n"
        general_errors.append(f"Ошибка тестирования: {str(e)}")
        log_to_file(f"{site_url} - Ошибка: {str(e)}")
    finally:
        try:
            if 'driver' in locals() and driver:
                driver.quit()
        except Exception as e:
            log_to_file(f"Ошибка закрытия WebDriver: {str(e)}")
        
        # Проверяем, была ли остановка
        if stop_event and stop_event.is_set():
            summary_area.value = "⏹ Тест остановлен пользователем"
        else:
            progress_bar.value = 1.0
        
        # Скрываем кнопку остановки и показываем кнопку запуска
        page.data['stop_btn_visible'] = False
        page.data['run_btn_visible'] = True
        page.update()







def run_links_test(site_url: str, summary_area: ft.TextField, page: ft.Page, progress_bar: ft.ProgressBar, ignore_ssl: bool, target_keywords: str, max_links: int = 15000):
    """Запускает тестирование ссылок без проверки robots и sitemap."""
    if not re.match(r'^https?://', site_url):
        summary_area.value = "❌ Неверный URL\n"
        page.update()
        return

    # Проверяем событие остановки
    stop_event = page.data.get('stop_event')
    if stop_event and stop_event.is_set():
        summary_area.value = "⏹ Проверка ссылок остановлена пользователем"
        page.update()
        return

    driver = None
    general_positives = []
    general_errors = []
    general_recs = []
    seo_positives = []
    seo_errors = []
    seo_recs = []
    perf_positives = []
    perf_errors = []
    perf_recs = []
    broken_links = []
    link_statuses = {}  # Для вкладки Ссылки
    site_links = []
    total_checks = 25  # Уменьшено без robots и sitemap
    current_check = 0
    log_text = ""  # Внутренняя переменная для логов
    total_images = 0  # Инициализация переменной для изображений
    images_list = []  # Инициализация списка изображений

    def update_progress():
        nonlocal current_check
        current_check += 1
        progress_bar.value = min(current_check / total_checks, 1.0)
        page.update()
        
        # Проверяем событие остановки
        if stop_event and stop_event.is_set():
            return False  # Сигнал остановки
        return True  # Продолжаем

    try:
        log_text += f"\n=== 🧪 Тестирование ссылок: {site_url} ===\n"
        general_positives.append(f"🔍 Тестирование ссылок: {site_url} ({datetime.now().strftime('%Y-%m-%d %H:%M:%S %Z')})")
        if not update_progress():
            return

        # Проверка доступности сайта с обходом блокировок
        log_text += "\n🔍 Проверка доступности сайта с обходом блокировок...\n"
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        
        # Сначала пробуем обычный метод
        anti_bot_mode = False
        try:
            start_time = time.time()
            r = requests.get(site_url, timeout=5, verify=not ignore_ssl, allow_redirects=True)
            load_time = time.time() - start_time
            page_size = len(r.content) / 1024
            log_text += f"🔎 HTTP статус: {r.status_code}\n"
            log_text += f"⏱ Время загрузки (HTTP): {load_time:.2f} сек\n"
            log_text += f"📏 Размер страницы: {page_size:.2f} КБ\n"
            
            # Проверяем на блокировку
            if any(blocked_text in r.text.lower() for blocked_text in [
                'access denied', 'доступ запрещен', 'blocked', 'заблокирован',
                'cloudflare', 'captcha', 'recaptcha', 'bot', 'robot'
            ]):
                log_text += "⚠️ Обнаружена блокировка, пробуем обход...\n"
                general_errors.append("Сайт блокирует доступ (обнаружена защита от ботов)")
                anti_bot_mode = True
            else:
                general_positives.append(f"Сайт доступен (HTTP статус: {r.status_code}, время загрузки: {load_time:.2f} сек)")
                general_positives.append(f"Размер страницы: {page_size:.2f} КБ")
                anti_bot_mode = False
                
            if page_size > 2000:
                general_errors.append("Страница слишком тяжёлая")
                general_recs.append("Сожмите изображения, минимизируйте CSS/JS, используйте lazy loading.")
            if load_time > 3:
                general_errors.append("Время загрузки слишком большое")
                general_recs.append("Включите кэширование, используйте CDN, оптимизируйте сервер.")
            if r.status_code != 200:
                seo_errors.append("Проблема индексации: статус не 200")
                seo_recs.append("Исправьте статус страницы для индексации.")
                
        except Exception as e:
            log_text += f"❌ Ошибка подключения: {e}\n"
            log_text += "🔄 Пробуем обход блокировок...\n"
            anti_bot_mode = True
            general_errors.append(f"Сайт недоступен: {e}")
            seo_errors.append("Проблема несуществующей страницы или индексации")
        
        update_progress()

        # Создание основного драйвера для SEO проверок с обходом блокировок
        try:
            if anti_bot_mode:
                log_text += "🛡️ Создание WebDriver в режиме обхода блокировок...\n"
                driver = create_webdriver(ignore_ssl=ignore_ssl, anti_bot_mode=True)
            else:
                driver = create_webdriver(ignore_ssl=ignore_ssl)
            
            # Пробуем загрузить страницу
            try:
                driver.get(site_url)
                log_text += "✅ Страница успешно загружена\n"
            except Exception as e:
                log_text += f"⚠️ Ошибка загрузки страницы: {e}\n"
                if not anti_bot_mode:
                    log_text += "🔄 Пробуем в режиме обхода блокировок...\n"
                    driver.quit()
                    driver = create_webdriver(ignore_ssl=ignore_ssl, anti_bot_mode=True)
                    driver.get(site_url)
                    log_text += "✅ Страница загружена в режиме обхода\n"
                
        except Exception as e:
            log_to_file(f"Ошибка создания WebDriver: {str(e)}")
            summary_area.value = f"❌ Ошибка: Не удалось создать WebDriver: {str(e)}"
            page.update()
            return

        # Получение версии ChromeDriver
        driver_version = driver.capabilities['chrome']['chromedriverVersion'].split(' ')[0]
        log_text += f"📌 Версия ChromeDriver: {driver_version}\n"
        general_positives.append(f"Версия ChromeDriver: {driver_version}")

        # Проверка зеркал и стандарта URL
        log_text += "\n🔄 Проверка зеркал и стандарта URL\n"
        mirror_issues = check_mirrors_and_redirects(site_url, ignore_ssl)
        if mirror_issues:
            seo_errors.extend(mirror_issues)
            seo_recs.append("Настройте правильные 301 редиректы для зеркал (www/non-www, http/https).")
        else:
            seo_positives.append("Зеркала настроены правильно")
        update_progress()

        # Проверка цепочек редиректов
        log_text += "\n🔄 Проверка цепочек редиректов\n"
        has_long_chain, chain_len = check_redirect_chain(site_url, ignore_ssl)
        if has_long_chain:
            seo_errors.append(f"Длинная цепочка редиректов ({chain_len} шагов)")
            seo_recs.append("Сократите цепочку редиректов до 1-2 шагов.")
        else:
            seo_positives.append("Цепочки редиректов в норме")
        update_progress()

        # Анализ производительности
        log_text += "\n⏱ Анализ производительности\n"
        performance_data = analyze_performance(site_url, ignore_ssl)
        load_times = performance_data["load_times"]
        resource_times = performance_data["resource_times"]
        js_css_times = performance_data["js_css_times"]

        log_text += f"⏱ Общее время загрузки: {', '.join(f'{t:.2f} сек' for t in load_times)}\n"
        log_text += f"⏱ Время загрузки ресурсов: {', '.join(f'{t:.2f} сек' for t in resource_times)}\n"
        log_text += f"⏱ Время выполнения JS/CSS: {', '.join(f'{t:.2f} сек' for t in js_css_times)}\n"
        perf_positives.append(f"Общее время загрузки: {', '.join(f'{t:.2f} сек' for t in load_times)}")
        perf_positives.append(f"Время загрузки ресурсов: {', '.join(f'{t:.2f} сек' for t in resource_times)}")
        perf_positives.append(f"Время выполнения JS/CSS: {', '.join(f'{t:.2f} сек' for t in js_css_times)}")
        for i, (lt, rt, jct) in enumerate(zip(load_times, resource_times, js_css_times)):
            res = ['1920x1080', '768x1024', '375x667'][i]
            if lt > 3:
                perf_errors.append(f"Слишком долгое время загрузки на {res}: {lt:.2f} сек")
                perf_recs.append("Оптимизируйте сервер и кэширование.")
            else:
                perf_positives.append(f"Нормальное время загрузки на {res}: {lt:.2f} сек")
            if rt > 1:
                perf_errors.append(f"Долгое время загрузки ресурсов на {res}: {rt:.2f} сек")
                perf_recs.append("Сожмите ресурсы, используйте CDN.")
            if jct > 1:
                perf_errors.append(f"Долгое время выполнения JS/CSS на {res}: {jct:.2f} сек")
                perf_recs.append("Минимизируйте и асинхронно загружайте скрипты.")
        update_progress()

        # Core Web Vitals
        log_text += "\n⚡ Core Web Vitals\n"
        try:
            lcp, fid, cls = get_core_web_vitals(driver)
            core_vitals_summary = f"LCP: {lcp:.2f} мс | FID: {fid:.2f} мс | CLS: {cls:.3f}\n"
            if lcp > 2500:
                perf_errors.append(f"LCP слишком большой: {lcp:.0f} мс")
                perf_recs.append("Сократите время Largest Contentful Paint до <2.5 сек.")
            if cls > 0.1:
                perf_errors.append(f"CLS слишком большой: {cls:.3f}")
                perf_recs.append("Снизьте Cumulative Layout Shift до <0.1.")
            log_text += core_vitals_summary
            perf_positives.append(core_vitals_summary)
        except Exception as e:
            log_to_file(f"Ошибка получения Core Web Vitals: {str(e)}")
            perf_errors.append("Не удалось получить Core Web Vitals")
        update_progress()

        # Микроразметка
        log_text += "\n🔎 Микроразметка (Schema.org, JSON-LD, OpenGraph, Twitter)\n"
        try:
            schema_items, jsonld_blocks, og_tags, twitter_tags = get_microdata(driver)
            if schema_items:
                seo_positives.append(f"Schema.org items: {', '.join(schema_items)}")
            else:
                seo_errors.append("Schema.org не найдена")
                seo_recs.append("Добавьте Schema.org микроразметку.")
            if jsonld_blocks:
                seo_positives.append(f"JSON-LD блоков: {len(jsonld_blocks)}")
            else:
                seo_errors.append("JSON-LD не найден")
                seo_recs.append("Добавьте JSON-LD для расширенного описания сайта.")
            if og_tags:
                seo_positives.append(f"OpenGraph: {', '.join(og_tags.keys())}")
            else:
                seo_errors.append("OpenGraph не найден")
                seo_recs.append("Добавьте OpenGraph для соцсетей.")
            if twitter_tags:
                seo_positives.append(f"Twitter Cards: {', '.join(twitter_tags.keys())}")
            else:
                seo_errors.append("Twitter Cards не найдены")
                seo_recs.append("Добавьте Twitter Cards для соцсетей.")
        except Exception as e:
            log_to_file(f"Ошибка получения микроразметки: {str(e)}")
            seo_errors.append("Не удалось получить микроразметку")
        update_progress()

        # Проверка SEO-элементов
        try:
            title_tag = driver.title if driver.title else "Не найден"
            log_text += f"📝 Тег title: {title_tag} (Длина: {len(title_tag)})\n"
            seo_positives.append(f"Тег title: {title_tag} (Длина: {len(title_tag)})")
            if len(title_tag) > 60:
                seo_errors.append("Длина title превышает 60 символов")
                seo_recs.append("Сократите title до 60 символов.")
            elif title_tag == "Не найден":
                seo_errors.append("Тег title отсутствует")
                seo_recs.append("Добавьте тег title.")
            else:
                seo_positives.append("Тег title в норме")
        except Exception as e:
            log_to_file(f"Ошибка получения title: {str(e)}")
            seo_errors.append("Не удалось получить title")
        update_progress()

        try:
            meta_desc = driver.find_elements(By.XPATH, "//meta[@name='description']")
            if meta_desc:
                desc_content = meta_desc[0].get_attribute("content")
                log_text += f"📝 Мета-описание: {desc_content} (Длина: {len(desc_content)})\n"
                seo_positives.append(f"Мета-описание: {desc_content} (Длина: {len(desc_content)})")
                if len(desc_content) > 160:
                    seo_errors.append("Длина мета-описания превышает 160 символов")
                    seo_recs.append("Сократите до 160 символов.")
            else:
                seo_errors.append("Мета-описание отсутствует")
                seo_recs.append("Добавьте мета-описание.")
        except Exception as e:
            log_to_file(f"Ошибка получения мета-описания: {str(e)}")
            seo_errors.append("Не удалось получить мета-описание")
        update_progress()

        # Проверка заголовков
        try:
            h1_tags = driver.find_elements(By.TAG_NAME, "h1")
            if len(h1_tags) == 1:
                seo_positives.append("Один H1 найден")
            elif len(h1_tags) > 1:
                seo_errors.append(f"Найдено {len(h1_tags)} тегов H1")
                seo_recs.append("Используйте только один H1.")
            else:
                seo_errors.append("Тег H1 отсутствует")
                seo_recs.append("Добавьте тег H1.")
        except Exception as e:
            log_to_file(f"Ошибка проверки заголовков: {str(e)}")
            seo_errors.append("Не удалось проверить заголовки")
        update_progress()

        # Проверка Open Graph
        log_text += "\n🌐 Проверка Open Graph\n"
        try:
            og_tags = check_open_graph(driver)
            for tag, content in og_tags.items():
                log_text += f"  - {tag}: {content}\n"
                seo_positives.append(f"{tag}: {content}")
                if "Не найден" in content or "Без содержимого" in content:
                    seo_errors.append(f"Отсутствует тег Open Graph: {tag}")
                    seo_recs.append(f"Добавьте {tag}.")
        except Exception as e:
            log_to_file(f"Ошибка проверки Open Graph: {str(e)}")
            seo_errors.append("Не удалось проверить Open Graph")
        update_progress()

        # Проверка микроразметки
        log_text += "\n📋 Проверка микроразметки Schema.org\n"
        try:
            has_schema, schema_result = check_schema_markup(driver)
            log_text += f"  - {schema_result}\n"
            if has_schema:
                seo_positives.append(schema_result)
            else:
                seo_errors.append(schema_result)
                seo_recs.append("Добавьте микроразметку Schema.org.")
        except Exception as e:
            log_to_file(f"Ошибка проверки микроразметки: {str(e)}")
            seo_errors.append("Не удалось проверить микроразметку")
        update_progress()

        # Проверка noindex, nofollow, noarchive и тега <noindex>
        log_text += "\n🤖 Проверка meta robots и тега <noindex>\n"
        try:
            noindex_meta, nofollow_meta, noarchive_meta, has_noindex_tag, robots_content = check_noindex_nofollow_noarchive(driver)
            if robots_content != "Не найден":
                log_text += f"📝 Meta robots content: {robots_content}\n"
                seo_positives.append(f"Meta robots: {robots_content}")
                if noindex_meta:
                    seo_errors.append("Страница имеет noindex в meta robots (не индексируется)")
                    seo_recs.append("Удалите noindex, если страница должна индексироваться.")
                if nofollow_meta:
                    seo_errors.append("Страница имеет nofollow в meta robots (ссылки не следуются)")
                    seo_recs.append("Удалите nofollow, если ссылки должны следовать.")
                if noarchive_meta:
                    seo_errors.append("Страница имеет noarchive в meta robots (нет сохраненной копии)")
                    seo_recs.append("Удалите noarchive, если нужна сохраненная копия.")
            else:
                seo_positives.append("Meta robots не найден (по умолчанию index, follow, archive)")
            
            # Проверка тега <noindex>
            if has_noindex_tag:
                seo_errors.append("Найден тег <noindex> (контент не индексируется)")
                seo_recs.append("Удалите тег <noindex>, если контент должен индексироваться.")
            else:
                seo_positives.append("Тег <noindex> не найден (контент может индексироваться)")
        except Exception as e:
            log_to_file(f"Ошибка проверки meta robots и тега <noindex>: {str(e)}")
            seo_errors.append("Не удалось проверить meta robots и тег <noindex>")
        update_progress()

        # Проверка скрытых блоков display: none
        log_text += "\n🕵️ Проверка скрытых блоков (display: none)\n"
        try:
            hidden_blocks = check_hidden_blocks(driver)
            if hidden_blocks:
                log_text += f"❌ Найдено {len(hidden_blocks)} скрытых элементов с текстом:\n"
                for block in hidden_blocks[:5]:  # Ограничим вывод первыми 5
                    log_text += f"  - {block}...\n"
                seo_errors.append(f"Найдено {len(hidden_blocks)} скрытых элементов (display: none) с текстом")
                seo_recs.append("Проверьте скрытый контент на спам; поисковики могут penalize за скрытый текст.")
            else:
                seo_positives.append("Нет скрытых элементов с display: none содержащих текст")
        except Exception as e:
            log_to_file(f"Ошибка проверки скрытых блоков: {str(e)}")
            seo_errors.append("Не удалось проверить скрытые блоки")
        update_progress()

        # Проверка canonical
        log_text += "\n📌 Проверка canonical\n"
        try:
            has_canonical, canonical_href = check_canonical(driver)
            if has_canonical:
                seo_positives.append(f"Canonical найден: {canonical_href}")
                if canonical_href != site_url:
                    seo_errors.append("Canonical указывает на другой URL (возможный дубликат)")
                    seo_recs.append("Убедитесь, что canonical указывает на предпочтительную версию.")
            else:
                seo_errors.append("Canonical отсутствует")
                seo_recs.append("Добавьте canonical для предотвращения дубликатов.")
        except Exception as e:
            log_to_file(f"Ошибка проверки canonical: {str(e)}")
            seo_errors.append("Не удалось проверить canonical")
        update_progress()

        # Проверка пагинации
        log_text += "\n📄 Проверка пагинации (rel=next/prev)\n"
        try:
            has_next, has_prev = check_pagination_links(driver)
            if has_next or has_prev:
                seo_positives.append(f"Пагинация найдена: next={has_next}, prev={has_prev}")
            else:
                seo_positives.append("Пагинация не найдена (возможно, не требуется)")
            # Для пагинации проверить canonical
            if has_next or has_prev and has_canonical and canonical_href == site_url:
                seo_positives.append("Canonical на пагинированной странице указывает на первую страницу")
            elif has_next or has_prev:
                seo_errors.append("Проблема пагинации: canonical не указывает на первую страницу")
                seo_recs.append("Для пагинированных страниц используйте self-canonical или на первую.")
        except Exception as e:
            log_to_file(f"Ошибка проверки пагинации: {str(e)}")
            seo_errors.append("Не удалось проверить пагинацию")
        update_progress()

        # Проверка внешних ссылок
        log_text += "\n🔗 Проверка внешних ссылок\n"
        try:
            ext_count, nofollow_count, broken_ext = check_external_links(driver, site_url)
            seo_positives.append(f"Внешних ссылок: {ext_count}, с nofollow: {nofollow_count}")
            if broken_ext:
                seo_errors.append(f"Битые внешние ссылки: {len(broken_ext)} ({', '.join(broken_ext[:5])}...)")
                seo_recs.append("Исправьте битые внешние ссылки.")
            if nofollow_count < ext_count / 2:
                seo_errors.append("Мало nofollow на внешних ссылках")
                seo_recs.append("Добавьте nofollow на внешние ссылки, если они не доверенные.")
        except Exception as e:
            log_to_file(f"Ошибка проверки внешних ссылок: {str(e)}")
            seo_errors.append("Не удалось проверить внешние ссылки")
        update_progress()

        # Проверка дубликатов
        log_text += "\n🔄 Проверка дубликатов\n"
        try:
            duplicates = check_duplicates(driver)
            if duplicates:
                seo_errors.append(f"Найдено {len(duplicates)} дубликатов")
                seo_recs.append("Исправьте дубликаты контента.")
            else:
                seo_positives.append("Дубликаты не найдены")
        except Exception as e:
            log_to_file(f"Ошибка проверки дубликатов: {str(e)}")
            seo_errors.append("Не удалось проверить дубликаты")
        update_progress()

        # Проверка рекламы
        log_text += "\n📢 Проверка рекламы\n"
        try:
            ads = check_ads(driver)
            if ads:
                seo_errors.append(f"Найдено {len(ads)} рекламных блоков")
                seo_recs.append("Проверьте рекламные блоки на соответствие правилам.")
            else:
                seo_positives.append("Рекламные блоки не найдены")
        except Exception as e:
            log_to_file(f"Ошибка проверки рекламы: {str(e)}")
            seo_errors.append("Не удалось проверить рекламу")
        update_progress()

        # Проверка безопасности
        log_text += "\n🔒 Проверка безопасности\n"
        try:
            security_issues = check_security(driver, site_url)
            if security_issues:
                seo_errors.extend(security_issues)
                seo_recs.append("Исправьте проблемы безопасности.")
            else:
                seo_positives.append("Проблемы безопасности не найдены")
        except Exception as e:
            log_to_file(f"Ошибка проверки безопасности: {str(e)}")
            seo_errors.append("Не удалось проверить безопасность")
        update_progress()

        # Проверка изображений
        log_text += "\n🖼 Проверка изображений\n"
        try:
            images = driver.find_elements(By.TAG_NAME, "img")
            total_images = len(images)
            images_list = []
            for img in images:
                src = img.get_attribute("src") or "Нет src"
                alt = img.get_attribute("alt") or ""
                title = img.get_attribute("title") or ""
                try:
                    size = get_image_size(src, ignore_ssl)
                except:
                    size = 0
                images_list.append({"src": src, "alt": alt, "title": title, "size": size})
            log_text += f"🖼 Изображений: {total_images}\n"
            general_positives.append(f"Изображений: {total_images}")
            if total_images == 0:
                seo_errors.append("На странице нет изображений")
                seo_recs.append("Добавьте изображения для улучшения SEO.")
        except Exception as e:
            log_to_file(f"Ошибка проверки изображений: {str(e)}")
            seo_errors.append("Не удалось проверить изображения")
        update_progress()

        # Проверка ключевых слов
        if target_keywords and target_keywords.strip():
            log_text += "\n🔍 АНАЛИЗ ЦЕЛЕВЫХ КЛЮЧЕВЫХ СЛОВ\n"
            try:
                keywords, density, target_analysis = analyze_keywords(driver, site_url, target_keywords)
                
                if isinstance(target_analysis, dict) and target_analysis:
                    log_text += "=" * 60 + "\n"
                    for tkw, data in target_analysis.items():
                        log_text += f"🎯 ЦЕЛЕВОЕ КЛЮЧЕВОЕ СЛОВО: '{tkw}'\n"
                        log_text += f"📊 ОБЩАЯ ЧАСТОТА (со склонениями): {data['freq']} раз\n"
                        # Безопасное форматирование плотности
                        density_value = data['density']
                        if isinstance(density_value, (int, float)):
                            log_text += f"📈 ПЛОТНОСТЬ: {density_value:.2%}\n"
                        else:
                            log_text += f"📈 ПЛОТНОСТЬ: {density_value}\n"
                        
                        # Показываем найденные склонения
                        if 'declensions_found' in data and data['declensions_found']:
                            log_text += "📝 НАЙДЕННЫЕ СКЛОНЕНИЯ:\n"
                            for declension, count in data['declensions_found'].items():
                                if count > 0:
                                    log_text += f"  ✅ '{declension}': {count} раз\n"
                            
                            # Показываем полный текст с подсветкой
                            log_text += "\n📄 ПОЛНЫЙ ТЕКСТ С ПОДСВЕТКОЙ:\n"
                            html = driver.page_source
                            soup = BeautifulSoup(html, 'html.parser')
                            text = soup.get_text(separator=' ', strip=True)
                            
                            # Подсвечиваем найденные склонения
                            highlighted_text = text
                            for declension, count in data['declensions_found'].items():
                                if count > 0:
                                    # Заменяем найденные склонения на подсвеченные версии
                                    pattern = r'\b' + re.escape(declension) + r'\b'
                                    highlighted_text = re.sub(pattern, f"【{declension}】", highlighted_text, flags=re.IGNORECASE)
                            
                            # Показываем первые 500 символов с подсветкой
                            preview = highlighted_text[:500] + "..." if len(highlighted_text) > 500 else highlighted_text
                            log_text += f"{preview}\n"
                        else:
                            log_text += "❌ Склонения не найдены\n"
                        
                        log_text += "=" * 60 + "\n\n"
                        
                        # Оценка плотности
                        if isinstance(density_value, (int, float)):
                            if density_value < 0.01:
                                seo_errors.append(f"Низкая плотность для целевого '{tkw}' ({density_value:.2%})")
                                seo_recs.append(f"Увеличьте использование '{tkw}' до 1-2%.")
                            elif density_value > 0.03:
                                seo_errors.append(f"Высокая плотность для целевого '{tkw}' ({density_value:.2%})")
                                seo_recs.append(f"Снизьте использование '{tkw}' до 1-2%.")
                            else:
                                seo_positives.append(f"Нормальная плотность для целевого '{tkw}' ({density_value:.2%})")
                        else:
                            seo_errors.append(f"Ошибка расчета плотности для '{tkw}': {density_value}")
                else:
                    log_text += "❌ Целевые ключевые слова не найдены на странице\n"
                    seo_errors.append("Целевые ключевые слова не найдены")
            except Exception as e:
                log_to_file(f"Ошибка анализа ключевых слов: {str(e)}")
                seo_errors.append("Не удалось проанализировать ключевые слова")
        update_progress()

        # Проверка ссылок (битые) и сбор статусов
        try:
            links = driver.find_elements(By.TAG_NAME, "a")
            log_text += f"🔗 Найдено тегов <a>: {len(links)}\n"
            
            # Проверяем все найденные ссылки
            links_to_check = links
            checked_links_count = 0
            
            for i, link in enumerate(links_to_check, 1):
                href = link.get_attribute("href") or "Нет href"
                if href and "javascript:void" not in href and not href.startswith("#"):
                    result = check_resource(href, ignore_ssl)
                    href, status, _, _ = result
                    link_statuses[href] = status
                    checked_links_count += 1
                    if not isinstance(status, int) or status != 200:
                        broken_links.append(href)
            
            # Добавляем информацию о проверенных ссылках
            log_text += f"🔗 Проверено ссылок: {checked_links_count}\n"
            general_positives.append(f"Проверено ссылок: {checked_links_count}")
            site_links = [link.get_attribute("href") for link in links if link.get_attribute("href") and site_url in link.get_attribute("href")]
        except Exception as e:
            log_to_file(f"Ошибка проверки ссылок: {str(e)}")
            general_errors.append("Не удалось проверить ссылки")
        if broken_links:
            general_errors.append(f"Найдены битые ссылки: {len(broken_links)} ({', '.join(broken_links)})")
            general_recs.append("Исправьте битые ссылки.")
        else:
            general_positives.append("Нет битых ссылок")
        update_progress()

        # Генерация графика производительности
        try:
            chart_base64 = generate_performance_chart(load_times, resource_times, js_css_times)
            page.add(ft.Image(src_base64=chart_base64, width=800, height=500))
        except Exception as e:
            log_to_file(f"Ошибка генерации графика: {str(e)}")
        update_progress()

        log_to_file(f"{site_url} - Успешно протестирован (только ссылки)")

        # Форматирование сводок
        seo_area = ft.TextField()
        perf_area = ft.TextField()
        links_area = ft.TextField()
        seo_area.value = format_summary_section(seo_positives, seo_errors, seo_recs, "SEO Анализ")
        perf_area.value = format_summary_section(perf_positives, perf_errors, perf_recs, "Производительность")
        full_summary = format_summary_section(general_positives + seo_positives + perf_positives,
                                                    general_errors + seo_errors + perf_errors,
                                                    general_recs + seo_recs + perf_recs,
                                                    "Общая Сводка")
        
        # Создаем сводку ссылок с кнопками для детального просмотра
        links_summary_with_buttons = "### Ссылки\n\n"
        for url, status in link_statuses.items():
            status_emoji = "🟢" if isinstance(status, int) and status == 200 else "🔴"
            links_summary_with_buttons += f"{status_emoji} {url} (Статус: {status})\n"
            links_summary_with_buttons += f"   [Детали] - кнопка для просмотра деталей\n\n"
        
        # Создаем сводку ссылок с информацией о деталях
        links_summary_simple = f"### Ссылки ({len(link_statuses)} проверенных)\n\n"
        for url, status in link_statuses.items():
            status_emoji = "🟢" if isinstance(status, int) and status == 200 else "🔴"
            links_summary_simple += f"{status_emoji} {url} (Статус: {status})\n"
        links_summary_simple += f"\n💡 Проверено {len(link_statuses)} ссылок из {len(links)} найденных тегов <a>"
        
        # Сохраняем данные ссылок для детального просмотра
        page.data['link_statuses'] = link_statuses
        
        summary_area.value = full_summary  # Только итоговая сводка!
        page.data['full_summary'] = full_summary
        page.data['seo_summary'] = seo_area.value  # Сохраняем SEO сводку отдельно
        links_area.value = format_links_section(link_statuses)
        links_summary = check_links_summary(link_statuses)
        page.data['links_summary'] = links_summary
        page.data['links_summary_with_buttons'] = links_summary_with_buttons

        # Форматирование сводки изображений
        images_summary = "### Изображения\n\n"
        if total_images == 0:
            images_summary += "На странице нет изображений\n"
        else:
            for img in images_list:
                alt_emoji = "🟢" if img['alt'] else "🔴"
                title_emoji = "🟢" if img['title'] else "🔴"
                size_emoji = "🟢" if img['size'] <= 300 else "🔴"
                images_summary += f"Ссылка: {img['src']}\n"
                images_summary += f"Alt: {alt_emoji} {img['alt'] or 'Нет'}\n"
                images_summary += f"Title: {title_emoji} {img['title'] or 'Нет'}\n"
                images_summary += f"Размер: {size_emoji} {img['size']:.2f} КБ\n\n"
        page.data['images_summary'] = images_summary

        # Сохранение результатов
        save_results(site_url, summary_area.value, full_summary)

    except Exception as e:
        summary_area.value = f"❌ Ошибка: {str(e)}\n"
        general_errors.append(f"Ошибка тестирования: {str(e)}")
        log_to_file(f"{site_url} - Ошибка: {str(e)}")
    finally:
        try:
            if 'driver' in locals() and driver:
                driver.quit()
        except Exception as e:
            log_to_file(f"Ошибка закрытия WebDriver: {str(e)}")
        
        # Проверяем, была ли остановка
        if stop_event and stop_event.is_set():
            summary_area.value = "⏹ Проверка ссылок остановлена пользователем"
        else:
            progress_bar.value = 1.0
        
        # Скрываем кнопку остановки и показываем кнопку запуска
        page.data['links_stop_btn_visible'] = False
        page.data['links_run_btn_visible'] = True
        
        # Показываем кнопки экспорта через page.data
        page.data['links_export_btn_visible'] = True
        page.data['links_export_word_btn_visible'] = True
        
        # Обновляем интерфейс
        page.update()

def run_multiple_links_test(urls: list, summary_area: ft.TextField, page: ft.Page, progress_bar: ft.ProgressBar, ignore_ssl: bool, target_keywords: str):
    """Запускает тестирование множественных ссылок без проверки robots и sitemap."""
    if not urls:
        summary_area.value = "❌ Нет ссылок для проверки\n"
        page.update()
        return

    all_results = []
    all_seo_summaries = []
    all_links_summaries = []
    all_images_summaries = []
    all_full_summaries = []
    total_urls = len(urls)
    current_url = 0
    
    summary_area.value = f"🔄 Начинаем проверку {total_urls} ссылок...\n"
    page.update()

    # Ограничиваем количество ссылок для оптимизации
    max_urls_per_batch = 1000  # Максимум 1000 ссылок за раз
    if total_urls > max_urls_per_batch:
        summary_area.value += f"⚠️ Ограничиваем проверку до {max_urls_per_batch} ссылок для оптимизации\n"
        urls = urls[:max_urls_per_batch]
        total_urls = len(urls)

    for url in urls:
        url = url.strip()
        if not url:
            continue
            
        current_url += 1
        summary_area.value += f"\n📋 Проверяем {current_url}/{total_urls}: {url}\n"
        progress_bar.value = current_url / total_urls
        page.update()
        
        try:
            # Создаем временную область для результатов одной ссылки
            temp_summary = ft.TextField()
            
            # Запускаем проверку одной ссылки с увеличенным лимитом
            run_links_test(url, temp_summary, page, progress_bar, ignore_ssl, target_keywords, 15000)
            
            # Сохраняем результаты
            if hasattr(temp_summary, 'value') and temp_summary.value:
                all_results.append({
                    'url': url,
                    'summary': temp_summary.value
                })
                
                # Сохраняем данные для кнопок
                if 'seo_summary' in page.data:
                    all_seo_summaries.append(f"## 🔗 {url}\n\n{page.data['seo_summary']}\n\n{'='*60}\n\n")
                if 'links_summary' in page.data:
                    all_links_summaries.append(f"## 🔗 {url}\n\n{page.data['links_summary']}\n\n{'='*60}\n\n")
                if 'images_summary' in page.data:
                    all_images_summaries.append(f"## 🔗 {url}\n\n{page.data['images_summary']}\n\n{'='*60}\n\n")
                if 'full_summary' in page.data:
                    all_full_summaries.append(f"## 🔗 {url}\n\n{page.data['full_summary']}\n\n{'='*60}\n\n")
                
        except Exception as e:
            all_results.append({
                'url': url,
                'summary': f"❌ Ошибка проверки: {str(e)}"
            })
    
    # Сохраняем объединенные данные для кнопок
    if all_seo_summaries:
        page.data['multiple_seo_summary'] = "# 📊 SEO Анализ всех сайтов\n\n" + "".join(all_seo_summaries)
    if all_links_summaries:
        page.data['multiple_links_summary'] = "# 📊 Анализ ссылок всех сайтов\n\n" + "".join(all_links_summaries)
    if all_images_summaries:
        page.data['multiple_images_summary'] = "# 📊 Анализ изображений всех сайтов\n\n" + "".join(all_images_summaries)
    if all_full_summaries:
        page.data['multiple_full_summary'] = "# 📊 Общая сводка всех сайтов\n\n" + "".join(all_full_summaries)
    
    # Формируем общую сводку
    combined_summary = f"# 📊 Сводка проверки {total_urls} ссылок\n\n"
    combined_summary += f"**Время проверки:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S %Z')}\n\n"
    
    for i, result in enumerate(all_results, 1):
        combined_summary += f"## 🔗 {i}. {result['url']}\n\n"
        combined_summary += result['summary']
        combined_summary += "\n" + "="*80 + "\n\n"
    
    summary_area.value = combined_summary
    page.data['multiple_results'] = all_results
    progress_bar.value = 1.0
    
    # Показываем кнопки экспорта через page.data
    page.data['links_export_btn_visible'] = True
    page.data['links_export_word_btn_visible'] = True
    
    # Обновляем интерфейс
    page.update()

def run_robots_check(site_url: str, ignore_ssl: bool, page: ft.Page, robots_area: ft.TextField, summary_area: ft.TextField):
    """Запускает отдельную проверку robots.txt."""
    summary = check_robots_summary(site_url, ignore_ssl)
    robots_area.value = summary
    summary_area.value = summary
    page.update()

def run_sitemap_check(site_url: str, ignore_ssl: bool, page: ft.Page, sitemap_area: ft.TextField, summary_area: ft.TextField):
    """Запускает отдельную проверку sitemap.xml."""
    summary = check_sitemap_summary(site_url, ignore_ssl)
    sitemap_area.value = summary
    summary_area.value = summary
    page.update()

def generate_report(summary, site_url, report_type='full', format='txt'):
    """Генерирует отчет в TXT или Excel."""
    report_path = save_results(site_url, '', summary, report_type, format)
    return report_path

def generate_sitemap_excel_report(site_url, ignore_ssl):
    """Генерирует Excel отчет со всеми URL из sitemap."""
    global sitemap_export_data
    
    if not sitemap_export_data:
        # Если данных нет, запускаем проверку sitemap
        check_sitemap_summary(site_url, ignore_ssl)
    
    if not sitemap_export_data or not sitemap_export_data.get('urls'):
        return None
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    report_path = f"{REPORT_DIR}/sitemap_full_report_{timestamp}.xlsx"
    
    # Создаем данные для Excel
    data = {
        'URL': [],
        'Статус': [],
        'Источник Sitemap': [],
        'Last Modified': [],
        'Priority': [],
        'Change Frequency': []
    }
    
    # Получаем все URL и их метаданные
    urls = sitemap_export_data['urls']
    page_details = sitemap_export_data['page_details']
    
    # Создаем словарь для быстрого поиска метаданных
    details_dict = {detail['url']: detail for detail in page_details}
    
    # Добавляем все URL в отчет
    for url in urls:
        detail = details_dict.get(url, {})
        data['URL'].append(url)
        data['Статус'].append(detail.get('status', 'неизвестно'))
        data['Источник Sitemap'].append(detail.get('source_sitemap', 'неизвестно'))
        data['Last Modified'].append(detail.get('lastmod', '-'))
        data['Priority'].append(detail.get('priority', '-'))
        data['Change Frequency'].append(detail.get('changefreq', '-'))
    
    # Создаем DataFrame и сохраняем в Excel
    df = pd.DataFrame(data)
    
    # Создаем Excel файл с несколькими листами
    with pd.ExcelWriter(report_path, engine='openpyxl') as writer:
        # Основной лист со всеми URL
        df.to_excel(writer, sheet_name='Все URL', index=False)
        
        # Лист только с проблемными URL
        broken_df = df[df['Статус'] != 'ОК'].copy()
        if not broken_df.empty:
            broken_df.to_excel(writer, sheet_name='Проблемные URL', index=False)
        
        # Лист только с доступными URL
        working_df = df[df['Статус'] == 'ОК'].copy()
        if not working_df.empty:
            working_df.to_excel(writer, sheet_name='Доступные URL', index=False)
        
        # Лист со страницами на сайте, но не в sitemap
        pages_not_in_sitemap = sitemap_export_data.get('pages_not_in_sitemap', [])
        if pages_not_in_sitemap:
            not_in_sitemap_df = pd.DataFrame({
                'URL': pages_not_in_sitemap,
                'Статус': ['На сайте, но не в sitemap'] * len(pages_not_in_sitemap)
            })
            not_in_sitemap_df.to_excel(writer, sheet_name='Страницы не в sitemap', index=False)
        
        # Лист со страницами в sitemap, но не на сайте
        pages_in_sitemap_not_on_site = sitemap_export_data.get('pages_in_sitemap_not_on_site', [])
        if pages_in_sitemap_not_on_site:
            not_on_site_df = pd.DataFrame({
                'URL': pages_in_sitemap_not_on_site,
                'Статус': ['В sitemap, но не на сайте'] * len(pages_in_sitemap_not_on_site)
            })
            not_on_site_df.to_excel(writer, sheet_name='Страницы не на сайте', index=False)
    
    return report_path

def generate_sitemap_word_report(site_url, ignore_ssl):
    """Генерирует Word отчет со всеми URL из sitemap."""
    global sitemap_export_data
    
    if not Document:
        return None
    
    if not sitemap_export_data:
        # Если данных нет, запускаем проверку sitemap
        check_sitemap_summary(site_url, ignore_ssl)
    
    if not sitemap_export_data or not sitemap_export_data.get('urls'):
        return None
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    report_path = f"{REPORT_DIR}/sitemap_full_report_{timestamp}.docx"
    
    doc = Document()
    
    # Заголовок
    title = doc.add_heading(f'Полный отчет Sitemap для {site_url}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Информация о дате
    doc.add_paragraph(f"Дата создания: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Всего найдено URL: {len(sitemap_export_data['urls'])}")
    doc.add_paragraph("="*50)
    
    # Информация о sitemap
    sitemap_info = sitemap_export_data.get('sitemap_info', {})
    if sitemap_info.get('type') == 'sitemapindex':
        doc.add_heading('Тип Sitemap: Sitemap Index', level=1)
        if sitemap_info.get('sub_sitemaps'):
            doc.add_paragraph(f"Количество подчиненных sitemap: {len(sitemap_info['sub_sitemaps'])}")
    else:
        doc.add_heading('Тип Sitemap: Обычный Sitemap', level=1)
    
    # Статистика
    doc.add_heading('Статистика', level=1)
    doc.add_paragraph(f"Всего URL: {len(sitemap_export_data['urls'])}")
    
    # Подсчитываем статусы
    status_counts = {}
    for detail in sitemap_export_data['page_details']:
        status = detail.get('status', 'неизвестно')
        status_counts[status] = status_counts.get(status, 0) + 1
    
    for status, count in status_counts.items():
        doc.add_paragraph(f"Статус '{status}': {count} URL")
    
    # Список проблемных URL
    broken_urls = [detail for detail in sitemap_export_data['page_details'] if detail.get('status') != 'ОК']
    if broken_urls:
        doc.add_heading('Проблемные URL', level=1)
        doc.add_paragraph(f"Найдено {len(broken_urls)} проблемных URL:")
        
        for i, detail in enumerate(broken_urls, 1):
            doc.add_paragraph(f"{i}. {detail['url']} - {detail.get('status', 'неизвестно')}")
    
    # Список доступных URL
    working_urls = [detail for detail in sitemap_export_data['page_details'] if detail.get('status') == 'ОК']
    if working_urls:
        doc.add_heading('Доступные URL', level=1)
        doc.add_paragraph(f"Найдено {len(working_urls)} доступных URL:")
        
        for i, detail in enumerate(working_urls, 1):
            doc.add_paragraph(f"{i}. {detail['url']}")
    
    # Страницы на сайте, но не в sitemap
    pages_not_in_sitemap = sitemap_export_data.get('pages_not_in_sitemap', [])
    if pages_not_in_sitemap:
        doc.add_heading('Страницы на сайте, но не в sitemap', level=1)
        doc.add_paragraph(f"Найдено {len(pages_not_in_sitemap)} страниц:")
        
        for i, url in enumerate(pages_not_in_sitemap, 1):
            doc.add_paragraph(f"{i}. {url}")
    
    # Страницы в sitemap, но не на сайте
    pages_in_sitemap_not_on_site = sitemap_export_data.get('pages_in_sitemap_not_on_site', [])
    if pages_in_sitemap_not_on_site:
        doc.add_heading('Страницы в sitemap, но не на сайте', level=1)
        doc.add_paragraph(f"Найдено {len(pages_in_sitemap_not_on_site)} страниц:")
        
        for i, url in enumerate(pages_in_sitemap_not_on_site, 1):
            doc.add_paragraph(f"{i}. {url}")

    # Полный список URL
    doc.add_heading('Полный список URL', level=1)
    
    urls = sitemap_export_data['urls']
    page_details = sitemap_export_data['page_details']
    details_dict = {detail['url']: detail for detail in page_details}
    
    for i, url in enumerate(urls, 1):
        detail = details_dict.get(url, {})
        
        # Заголовок URL
        doc.add_heading(f'{i}. {url}', level=2)
        
        # Таблица с информацией
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Параметр'
        hdr_cells[1].text = 'Значение'
        
        # Данные
        data_rows = [
            ('Статус', detail.get('status', 'неизвестно')),
            ('Источник Sitemap', detail.get('source_sitemap', 'неизвестно')),
            ('Last Modified', detail.get('lastmod', '-')),
            ('Priority', detail.get('priority', '-')),
            ('Change Frequency', detail.get('changefreq', '-'))
        ]
        
        for param, value in data_rows:
            row_cells = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = str(value)
        
        doc.add_paragraph()  # Пустая строка между URL
    
    # Сохраняем документ
    doc.save(report_path)
    return report_path

def generate_word_report(data, site_url, report_type='parser'):
    """Генерирует отчет в Word формате."""
    if not Document:
        return None
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    report_path = f"reports/{report_type}_report_{timestamp}.docx"
    
    doc = Document()
    
    # Заголовок
    title = doc.add_heading(f'SEO Отчет для {site_url}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Информация о дате
    doc.add_paragraph(f"Дата создания: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Тип отчета: {report_type}")
    doc.add_paragraph("="*50)
    
    if report_type == 'parser':
        if isinstance(data, list):
            # Отчет парсера
            doc.add_heading('Результаты парсинга сайта', level=1)
            
            # Проверяем, что есть данные
            if not data:
                doc.add_paragraph("Нет данных для экспорта. Проверьте настройки парсера.")
                doc.save(report_path)
                return report_path
            
            doc.add_paragraph(f"Найдено страниц: {len(data)}")
            doc.add_paragraph()
            
            for i, item in enumerate(data, 1):
                # Заголовок страницы
                doc.add_heading(f'Страница {i}: {item["Ссылка"]}', level=2)
                
                # Таблица с информацией
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Заголовки таблицы
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Параметр'
                hdr_cells[1].text = 'Значение'
                
                # Данные
                data_rows = [
                    ('HTTP Статус', str(item['HTTP'])),
                    ('Редирект', item['Редирект'] if item['Редирект'] else 'Нет'),
                    ('SEO Статус', item['SEO']),
                    ('Title', item['Title'] if item['Title'] else 'Отсутствует'),
                    ('Meta Description', item['Meta_Description'] if item['Meta_Description'] else 'Отсутствует'),
                    ('H1', item['H1'] if item['H1'] else 'Отсутствует'),
                ]
                
                for param, value in data_rows:
                    row_cells = table.add_row().cells
                    row_cells[0].text = param
                    row_cells[1].text = value
                
                doc.add_paragraph()  # Пустая строка между страницами
        else:
            # Если данные приходят в виде строки (summary_content)
            doc.add_heading('Результаты парсинга сайта', level=1)
            doc.add_paragraph("Данные в текстовом формате:")
            doc.add_paragraph(data)
    else:
        # Текстовый отчет (для других типов)
        doc.add_heading('SEO Анализ', level=1)
        doc.add_paragraph(data)
    
    # Сохраняем документ
    doc.save(report_path)
    return report_path

def crawl_site_without_sitemap(start_url, ignore_ssl, update_callback, done_callback, stop_event, max_threads=20, max_pages=15000):
    """Рекурсивно обходит все внутренние страницы сайта без использования sitemap."""
    visited = set()
    results = []
    lock = threading.Lock()
    domain = urlparse(start_url).netloc
    queue = [start_url]
    threads = []

    def check_redirect(url, ignore_ssl):
        """Проверяет редиректы на странице."""
        try:
            # Сначала делаем запрос без редиректов, чтобы увидеть исходный статус
            r = requests.get(url, timeout=5, verify=not ignore_ssl, allow_redirects=False)
            if r.status_code in [301, 302, 303, 307, 308]:
                redirect_url = r.headers.get('Location', '')
                return f"{r.status_code} → {redirect_url}"
            return ""
        except Exception:
            return ""

    def analyze_seo_basic(soup, url):
        """Базовый SEO анализ страницы."""
        title = soup.title.string.strip() if soup.title and soup.title.string else ''
        h1 = soup.find('h1').text.strip() if soup.find('h1') else ''
        meta_desc = ''
        meta_desc_tag = soup.find('meta', attrs={'name': 'description'})
        if meta_desc_tag:
            meta_desc = meta_desc_tag.get('content', '')
        
        seo_score = 0
        seo_issues = []
        
        if title:
            seo_score += 1
        else:
            seo_issues.append("Нет title")
            
        if h1:
            seo_score += 1
        else:
            seo_issues.append("Нет H1")
            
        if meta_desc:
            seo_score += 1
        else:
            seo_issues.append("Нет meta description")
        
        if seo_score == 3:
            return "✅ Отлично", "Все основные SEO элементы присутствуют"
        elif seo_score == 2:
            return "⚠️ Хорошо", f"Проблемы: {', '.join(seo_issues)}"
        elif seo_score == 1:
            return "❌ Плохо", f"Проблемы: {', '.join(seo_issues)}"
        else:
            return "❌ Критично", f"Проблемы: {', '.join(seo_issues)}"

    def worker():
        while True:
            # Проверяем сигнал остановки
            if stop_event.is_set():
                return
                
            with lock:
                if not queue or len(visited) >= max_pages:
                    return
                url = queue.pop(0)
                if url in visited:
                    continue
                visited.add(url)
            try:
                # Проверяем сигнал остановки перед запросом
                if stop_event.is_set():
                    return
                    
                # Уменьшенный таймаут для ускорения
                r = requests.get(url, timeout=5, verify=not ignore_ssl, allow_redirects=True)
                status = r.status_code
                soup = BeautifulSoup(r.text, 'html.parser')
                
                # Проверяем редиректы
                redirect_info = check_redirect(url, ignore_ssl)
                
                # Анализируем SEO
                seo_status, seo_details = analyze_seo_basic(soup, url)
                
                results.append({
                    'Ссылка': url,
                    'HTTP': status,
                    'Редирект': redirect_info,
                    'SEO': seo_status,
                    'SEO_Details': seo_details,
                    'Title': soup.title.string.strip() if soup.title and soup.title.string else '',
                    'H1': soup.find('h1').text.strip() if soup.find('h1') else '',
                    'Meta_Description': soup.find('meta', attrs={'name': 'description'}).get('content', '') if soup.find('meta', attrs={'name': 'description'}) else ''
                })
                
                # Собираем новые ссылки
                for a in soup.find_all('a', href=True):
                    link = urljoin(url, a['href'])
                    parsed = urlparse(link)
                    if parsed.netloc == domain and link not in visited and link.startswith('http'):
                        with lock:
                            if link not in queue:
                                queue.append(link)
                
                # Небольшая задержка для снижения нагрузки на сервер
                time.sleep(0.05)  # 50ms задержка
                
                update_callback(len(visited), len(results))
            except Exception as e:
                results.append({
                    'Ссылка': url,
                    'HTTP': f"Ошибка: {str(e)}",
                    'Редирект': '',
                    'SEO': '❌ Ошибка',
                    'SEO_Details': f'Ошибка при анализе: {str(e)}',
                    'Title': '',
                    'H1': '',
                    'Meta_Description': ''
                })
                update_callback(len(visited), len(results))

    for _ in range(max_threads):
        t = threading.Thread(target=worker)
        t.start()
        threads.append(t)
    for t in threads:
        t.join()
    done_callback(results)

def analyze_text_content(html_content, url):
    """Анализирует текстовое содержимое страницы для SEO с использованием Selenium для получения всего контента."""
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Создаем копию для анализа заголовков (до удаления элементов)
    soup_for_headers = BeautifulSoup(html_content, 'html.parser')
    
    # Анализ заголовков (до удаления элементов)
    h1_tags = soup_for_headers.find_all('h1')
    h2_tags = soup_for_headers.find_all('h2')
    h3_tags = soup_for_headers.find_all('h3')
    h4_tags = soup_for_headers.find_all('h4')
    h5_tags = soup_for_headers.find_all('h5')
    h6_tags = soup_for_headers.find_all('h6')
    
    h1_texts = [h.get_text(strip=True) for h in h1_tags]
    h2_texts = [h.get_text(strip=True) for h in h2_tags]
    h3_texts = [h.get_text(strip=True) for h in h3_tags]
    
    # Анализ мета-тегов
    title = soup.title.string.strip() if soup.title and soup.title.string else ''
    meta_desc = ''
    meta_keywords = ''
    
    meta_desc_tag = soup.find('meta', attrs={'name': 'description'})
    if meta_desc_tag:
        meta_desc = meta_desc_tag.get('content', '')
    
    meta_keywords_tag = soup.find('meta', attrs={'name': 'keywords'})
    if meta_keywords_tag:
        meta_keywords = meta_keywords_tag.get('content', '')
    
    # Получаем весь текст через Selenium (как в функции анализа склонений)
    try:
        from selenium import webdriver
        from selenium.webdriver.chrome.options import Options
        
        chrome_options = Options()
        chrome_options.add_argument('--headless')
        chrome_options.add_argument('--no-sandbox')
        chrome_options.add_argument('--disable-dev-shm-usage')
        
        driver = webdriver.Chrome(options=chrome_options)
        driver.get(url)
        
        # Получаем весь видимый текст через улучшенный JavaScript с точным подсчетом
        full_text = driver.execute_script("""
            function getAllVisibleText() {
                // Получаем весь текст со страницы с максимальной точностью
                let allText = '';
                let wordCount = 0;
                let processedNodes = new Set();
                
                // Функция для рекурсивного обхода всех элементов с точным подсчетом
                function walkTextNodes(node) {
                    // Избегаем повторной обработки узлов
                    if (processedNodes.has(node)) {
                        return;
                    }
                    processedNodes.add(node);
                    
                    if (node.nodeType === Node.TEXT_NODE) {
                        // Получаем текст из текстовых узлов с точной обработкой
                        let text = node.textContent.trim();
                        if (text && text.length > 0) {
                            // Очищаем текст от лишних символов
                            text = text.replace(/\\s+/g, ' ').trim();
                            if (text) {
                                allText += text + ' ';
                                // Подсчитываем слова точно
                                let words = text.split(/\\s+/).filter(word => word.length >= 3);
                                wordCount += words.length;
                            }
                        }
                    } else if (node.nodeType === Node.ELEMENT_NODE) {
                        // Пропускаем header и footer
                        if (node.tagName === 'HEADER' || node.tagName === 'FOOTER') {
                            return;
                        }
                        
                        // Пропускаем скрытые элементы
                        const style = window.getComputedStyle(node);
                        if (style.display === 'none' || style.visibility === 'hidden' || style.opacity === '0') {
                            return;
                        }
                        
                        // Пропускаем скрипты и стили
                        if (node.tagName === 'SCRIPT' || node.tagName === 'STYLE' || node.tagName === 'NOSCRIPT') {
                            return;
                        }
                        
                        // НЕ получаем текст из атрибутов (alt, title, placeholder, aria-label)
                        // Исключаем эти атрибуты по требованию
                        
                        // Рекурсивно обходим дочерние элементы
                        for (let child of node.childNodes) {
                            walkTextNodes(child);
                        }
                    }
                }
                
                // Ждем загрузки динамического контента с повторными попытками
                function waitForDynamicContent() {
                    return new Promise((resolve) => {
                        let attempts = 0;
                        const maxAttempts = 3;
                        
                        function tryExtract() {
                            attempts++;
                            
                            // Очищаем предыдущие данные
                            allText = '';
                            wordCount = 0;
                            processedNodes.clear();
                            
                            // Начинаем обход с body
                            walkTextNodes(document.body);
                            
                            // Проверяем, что получили достаточно контента
                            if (allText.trim().length > 100 || attempts >= maxAttempts) {
                                resolve({
                                    text: allText,
                                    wordCount: wordCount,
                                    attempts: attempts
                                });
                            } else {
                                // Ждем еще и пробуем снова
                                setTimeout(tryExtract, 500);
                            }
                        }
                        
                        setTimeout(tryExtract, 1000);
                    });
                }
                
                return waitForDynamicContent();
            }
            return getAllVisibleText();
        """)
        
        # Извлекаем текст и информацию о подсчете
        if isinstance(full_text, dict):
            text_data = full_text
            full_text = text_data.get('text', '')
            selenium_word_count = text_data.get('wordCount', 0)
            attempts = text_data.get('attempts', 1)
        else:
            selenium_word_count = 0
            attempts = 1
        
        driver.quit()
        
        # Очищаем полученный текст
        text = re.sub(r'\s+', ' ', full_text).strip()
        
    except Exception as e:
        # Если Selenium не работает, используем старый метод
        # Более точное удаление элементов навигации и служебных блоков
        # Удаляем только явные элементы навигации и служебные блоки
        elements_to_remove = []
        
        # Находим элементы для удаления
        for element in soup.find_all():
            # Удаляем скрипты и стили
            if element.name in ['script', 'style', 'noscript']:
                elements_to_remove.append(element)
                continue
                
            # Удаляем элементы навигации
            if element.name in ['nav']:
                elements_to_remove.append(element)
                continue
                
            # Удаляем футеры и хедеры (но только если они явно помечены)
            if element.name in ['footer', 'header']:
                elements_to_remove.append(element)
                continue
                
            # Удаляем элементы с классами, указывающими на навигацию/служебные блоки
            classes = element.get('class', [])
            if isinstance(classes, str):
                classes = [classes]
            
            navigation_classes = [
                'nav', 'navigation', 'menu', 'header', 'footer', 'sidebar', 
                'breadcrumb', 'breadcrumbs', 'pagination', 'pager',
                'social', 'socials', 'share', 'sharing', 'widget', 'widgets',
                'advertisement', 'ad', 'ads', 'banner', 'banners',
                'cookie', 'cookies', 'popup', 'modal', 'overlay',
                'search', 'search-form', 'searchbox', 'search-box',
                'login', 'signin', 'register', 'signup', 'auth',
                'cart', 'basket', 'checkout', 'order',
                'newsletter', 'subscribe', 'subscription',
                'toolbar', 'toolbar-top', 'toolbar-bottom',
                'announcement', 'notice', 'alert', 'notification'
            ]
            
            # Проверяем классы на наличие навигационных элементов
            if any(nav_class in ' '.join(classes).lower() for nav_class in navigation_classes):
                elements_to_remove.append(element)
                continue
                
            # Проверяем ID на наличие навигационных элементов
            element_id = element.get('id', '').lower()
            if any(nav_class in element_id for nav_class in navigation_classes):
                elements_to_remove.append(element)
                continue
        
        # Удаляем найденные элементы
        for element in elements_to_remove:
            if element.parent:
                element.decompose()
        
        # Получаем основной текст после очистки
        text = soup.get_text(separator=' ', strip=True)
        
        # Очищаем текст от лишних пробелов и переносов строк
        text = re.sub(r'\s+', ' ', text).strip()
    
    # Анализ ключевых слов (максимально точный подсчет)
    # Очищаем текст от HTML тегов и лишних символов
    clean_text = re.sub(r'<[^>]+>', '', text)  # Удаляем HTML теги
    clean_text = re.sub(r'javascript:', '', clean_text, flags=re.IGNORECASE)
    clean_text = re.sub(r'http[s]?://[^\s]+', '', clean_text)  # Удаляем URL
    clean_text = re.sub(r'www\.[^\s]+', '', clean_text)
    clean_text = re.sub(r'[^\w\sа-яё]', ' ', clean_text)  # Оставляем только буквы и пробелы
    clean_text = re.sub(r'\s+', ' ', clean_text).strip()
    
    # Находим все слова (русские и английские) с максимально точным подсчетом
    # Используем улучшенный regex для избежания погрешности
    words = re.findall(r'\b[а-яёa-z]{3,}\b', clean_text.lower())
    
    # Дополнительная очистка для максимальной точности
    words = [word.strip() for word in words if word.strip() and len(word.strip()) >= 3]
    
    # Проверяем точность подсчета с Selenium
    if 'selenium_word_count' in locals() and selenium_word_count > 0:
        # Сравниваем подсчеты для валидации
        python_word_count = len(words)
        if abs(python_word_count - selenium_word_count) > 2:
            # Если разница большая, используем более точный метод
            print(f"⚠️ Разница в подсчете: Python={python_word_count}, Selenium={selenium_word_count}")
            # Дополнительная очистка для точности
            words = [word for word in words if re.match(r'^[а-яёa-z]{3,}$', word)]
    
    # Расширенный список русских стоп-слов
    russian_stop_words = {
        'это', 'как', 'так', 'и', 'в', 'над', 'к', 'до', 'не', 'на', 'но', 'за', 'то', 'с', 'ли',
        'а', 'во', 'от', 'со', 'для', 'о', 'же', 'ну', 'вы', 'бы', 'что', 'кто', 'он', 'она',
        'и', 'в', 'на', 'не', 'с', 'а', 'о', 'для', 'по', 'из', 'к', 'у', 'от', 'но', 'как',
        'что', 'это', 'то', 'или', 'за', 'при', 'да', 'но', 'же', 'бы', 'ли', 'быть', 'был',
        'была', 'были', 'было', 'есть', 'быть', 'мой', 'моя', 'мои', 'твой', 'твоя', 'твои',
        'наш', 'наша', 'наши', 'ваш', 'ваша', 'ваши', 'его', 'ее', 'их', 'себя', 'себе',
        'себя', 'мне', 'тебе', 'ему', 'ей', 'нам', 'вам', 'им', 'меня', 'тебя', 'его', 'ее',
        'нас', 'вас', 'них', 'мной', 'тобой', 'им', 'ей', 'нами', 'вами', 'ими', 'мой', 'твой',
        'свой', 'наш', 'ваш', 'его', 'ее', 'их', 'этот', 'тот', 'такой', 'такая', 'такое',
        'такие', 'столько', 'сколько', 'который', 'которая', 'которое', 'которые', 'кто',
        'что', 'какой', 'какая', 'какое', 'какие', 'чей', 'чья', 'чье', 'чьи', 'где', 'куда',
        'откуда', 'когда', 'почему', 'зачем', 'как', 'сколько', 'насколько', 'столько',
        'такой', 'такая', 'такое', 'такие', 'этот', 'эта', 'это', 'эти', 'тот', 'та', 'то',
        'те', 'сам', 'сама', 'само', 'сами', 'самый', 'самая', 'самое', 'самые', 'весь',
        'вся', 'все', 'все', 'каждый', 'каждая', 'каждое', 'каждые', 'любой', 'любая',
        'любое', 'любые', 'никакой', 'никакая', 'никакое', 'никакие', 'некоторый',
        'некоторая', 'некоторое', 'некоторые', 'всякий', 'всякая', 'всякое', 'всякие',
        # Добавляем английские стоп-слова
        'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by',
        'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did',
        'will', 'would', 'could', 'should', 'may', 'might', 'can', 'must', 'shall'
    }
    
    # Фильтруем стоп-слова и короткие слова
    filtered_words = [word for word in words if word not in russian_stop_words and len(word) > 2]
    
    # Подсчитываем частоту слов с максимально точным подсчетом
    word_freq = Counter(filtered_words)
    top_keywords = word_freq.most_common(20)
    
    # Дополнительная валидация точности подсчета
    total_words_counted = len(filtered_words)
    
    # Проверяем консистентность подсчета с Selenium
    if 'selenium_word_count' in locals() and selenium_word_count > 0:
        selenium_diff = abs(total_words_counted - selenium_word_count)
        if selenium_diff > 5:
            print(f"⚠️ Большая разница в подсчете: Python={total_words_counted}, Selenium={selenium_word_count}")
            # Используем более строгую фильтрацию
            filtered_words = [word for word in filtered_words if re.match(r'^[а-яёa-z]{3,}$', word) and word not in russian_stop_words]
            word_freq = Counter(filtered_words)
            top_keywords = word_freq.most_common(20)
            total_words_counted = len(filtered_words)
    
    # Анализ предложений
    sentences = text.split('.')
    sentences = [s.strip() for s in sentences if s.strip()]
    avg_sentence_length = sum(len(s.split()) for s in sentences) / len(sentences) if sentences else 0
    
    # Анализ абзацев (из очищенного soup)
    paragraphs = [p.get_text(strip=True) for p in soup.find_all('p') if p.get_text(strip=True)]
    avg_paragraph_length = sum(len(p.split()) for p in paragraphs) / len(paragraphs) if paragraphs else 0
    
    # Анализ изображений (из оригинального soup)
    images = soup_for_headers.find_all('img')
    images_with_alt = [img for img in images if img.get('alt')]
    images_without_alt = [img for img in images if not img.get('alt')]
    
    # Анализ ссылок (из оригинального soup)
    links = soup_for_headers.find_all('a', href=True)
    internal_links = [link for link in links if link['href'].startswith('/') or url in link['href']]
    external_links = [link for link in links if link['href'].startswith('http') and url not in link['href']]
    
    # Анализ структуры
    structure_score = 0
    if len(h1_tags) == 1:
        structure_score += 20
    elif len(h1_tags) > 1:
        structure_score -= 10 * (len(h1_tags) - 1)
    
    if len(h2_tags) >= 2:
        structure_score += 15
    if len(h3_tags) >= 3:
        structure_score += 10
    
    # Анализ плотности ключевых слов (максимально точный)
    total_words = total_words_counted if 'total_words_counted' in locals() else len(filtered_words)
    keyword_density = {}
    for keyword, count in top_keywords[:20]:  # Увеличиваем до 20 для лучшего анализа
        density = (count / total_words) * 100 if total_words > 0 else 0
        keyword_density[keyword] = round(density, 2)
    
    # SEO рекомендации
    recommendations = []
    
    if len(h1_tags) == 0:
        recommendations.append("❌ Отсутствует H1 заголовок")
    elif len(h1_tags) > 1:
        recommendations.append(f"⚠️ Найдено {len(h1_tags)} H1 заголовков (должен быть только один)")
    
    if len(h2_tags) < 2:
        recommendations.append("⚠️ Мало H2 заголовков для структурирования контента")
    
    if not meta_desc:
        recommendations.append("❌ Отсутствует meta description")
    elif len(meta_desc) < 120:
        recommendations.append("⚠️ Meta description слишком короткий (менее 120 символов)")
    elif len(meta_desc) > 160:
        recommendations.append("⚠️ Meta description слишком длинный (более 160 символов)")
    
    if not title:
        recommendations.append("❌ Отсутствует title")
    elif len(title) < 30:
        recommendations.append("⚠️ Title слишком короткий (менее 30 символов)")
    elif len(title) > 60:
        recommendations.append("⚠️ Title слишком длинный (более 60 символов)")
    
    if len(images_without_alt) > 0:
        recommendations.append(f"⚠️ {len(images_without_alt)} изображений без alt атрибута")
    
    if avg_sentence_length > 25:
        recommendations.append("⚠️ Предложения слишком длинные (в среднем более 25 слов)")
    
    if avg_paragraph_length > 150:
        recommendations.append("⚠️ Абзацы слишком длинные (в среднем более 150 слов)")
    
    # Положительные моменты
    positives = []
    
    if len(h1_tags) == 1:
        positives.append("✅ Правильная структура H1 заголовков")
    
    if len(h2_tags) >= 2:
        positives.append(f"✅ Хорошая структура с {len(h2_tags)} H2 заголовками")
    
    if meta_desc and 120 <= len(meta_desc) <= 160:
        positives.append("✅ Оптимальная длина meta description")
    
    if title and 30 <= len(title) <= 60:
        positives.append("✅ Оптимальная длина title")
    
    if len(images_with_alt) > 0:
        positives.append(f"✅ {len(images_with_alt)} изображений с alt атрибутом")
    
    if len(internal_links) >= 3:
        positives.append(f"✅ Хорошая внутренняя перелинковка ({len(internal_links)} ссылок)")
    
    # Сохраняем полный текст для вывода
    full_text = text
    
    return {
        'url': url,
        'title': title,
        'meta_description': meta_desc,
        'meta_keywords': meta_keywords,
        'h1_count': len(h1_tags),
        'h2_count': len(h2_tags),
        'h3_count': len(h3_tags),
        'h1_texts': h1_texts,
        'h2_texts': h2_texts,
        'h3_texts': h3_texts,
        'total_words': total_words,
        'avg_sentence_length': round(avg_sentence_length, 1),
        'avg_paragraph_length': round(avg_paragraph_length, 1),
        'top_keywords': top_keywords,
        'keyword_density': keyword_density,
        'images_total': len(images),
        'images_with_alt': len(images_with_alt),
        'images_without_alt': len(images_without_alt),
        'internal_links': len(internal_links),
        'external_links': len(external_links),
        'structure_score': structure_score,
        'recommendations': recommendations,
        'positives': positives,
        'text_preview': text[:500] + '...' if len(text) > 500 else text,
        'full_text': full_text  # Добавляем полный текст
    }

def analyze_code_content(html_content, url):
    """Анализирует код страницы на наличие ошибок в HTML, CSS, JS, PHP с детальной статистикой."""
    soup = BeautifulSoup(html_content, 'html.parser')
    
    errors = []
    warnings = []
    positives = []
    
    # === HTML АНАЛИЗ ===
    html_errors = []
    html_warnings = []
    html_stats = {}
    
    # Проверка DOCTYPE
    doctype_match = re.search(r'<!DOCTYPE[^>]*>', html_content, re.IGNORECASE)
    if not doctype_match:
        html_errors.append("❌ Отсутствует DOCTYPE")
    else:
        doctype = doctype_match.group(0)
        positives.append(f"✅ DOCTYPE: {doctype}")
        html_stats['doctype'] = doctype
    
    # Проверка обязательных тегов
    if not soup.html:
        html_errors.append("❌ Отсутствует тег <html>")
    else:
        html_stats['html_lang'] = soup.html.get('lang', 'не указан')
        if soup.html.get('lang'):
            positives.append(f"✅ HTML lang: {soup.html.get('lang')}")
        else:
            html_warnings.append("⚠️ Не указан атрибут lang в теге <html>")
    
    if not soup.head:
        html_errors.append("❌ Отсутствует тег <head>")
    else:
        positives.append("✅ Тег <head> присутствует")
    
    if not soup.body:
        html_errors.append("❌ Отсутствует тег <body>")
    else:
        positives.append("✅ Тег <body> присутствует")
    
    if not soup.title:
        html_errors.append("❌ Отсутствует тег <title>")
    else:
        title_text = soup.title.string.strip() if soup.title.string else ""
        if title_text:
            positives.append(f"✅ Title: {title_text[:50]}{'...' if len(title_text) > 50 else ''}")
            html_stats['title_length'] = len(title_text)
        else:
            html_warnings.append("⚠️ Тег <title> пустой")
    
    # Детальный анализ тегов
    all_tags = soup.find_all()
    tag_counts = {}
    for tag in all_tags:
        tag_name = tag.name
        tag_counts[tag_name] = tag_counts.get(tag_name, 0) + 1
    
    # Топ-10 самых используемых тегов
    top_tags = sorted(tag_counts.items(), key=lambda x: x[1], reverse=True)[:10]
    html_stats['total_tags'] = len(all_tags)
    html_stats['unique_tags'] = len(tag_counts)
    html_stats['top_tags'] = top_tags
    
    # Проверка закрытых тегов
    unclosed_tags = []
    for tag in soup.find_all():
        if tag.name in ['img', 'br', 'hr', 'input', 'meta', 'link', 'area', 'base', 'col', 'embed', 'source', 'track', 'wbr']:
            continue  # Одиночные теги
        if not tag.string and not tag.find_all():
            unclosed_tags.append(tag.name)
    
    if unclosed_tags:
        html_warnings.append(f"⚠️ Возможно незакрытые теги: {', '.join(set(unclosed_tags))}")
        html_stats['unclosed_tags'] = list(set(unclosed_tags))
    
    # Проверка атрибутов alt для изображений
    images_without_alt = soup.find_all('img', alt=lambda x: not x or x.strip() == '')
    total_images = len(soup.find_all('img'))
    if images_without_alt:
        alt_coverage = ((total_images - len(images_without_alt)) / total_images * 100) if total_images > 0 else 0
        html_warnings.append(f"⚠️ {len(images_without_alt)} из {total_images} изображений без alt атрибута ({alt_coverage:.1f}% покрытие)")
        html_stats['images_without_alt'] = len(images_without_alt)
        html_stats['total_images'] = total_images
        html_stats['alt_coverage'] = alt_coverage
    else:
        if total_images > 0:
            positives.append(f"✅ Все {total_images} изображений имеют alt атрибут")
    
    # Проверка валидности ссылок
    all_links = soup.find_all('a', href=True)
    invalid_links = []
    external_links = []
    internal_links = []
    broken_links = []
    
    for link in all_links:
        href = link['href']
        if href.startswith('javascript:') or href.startswith('mailto:') or href.startswith('tel:'):
            continue
        elif href.startswith('http'):
            if not href.startswith('http://') and not href.startswith('https://'):
                invalid_links.append(href)
            else:
                external_links.append(href)
        elif href.startswith('/') or href.startswith('./') or href.startswith('../'):
            internal_links.append(href)
        elif href.startswith('#'):
            # Проверка якорных ссылок
            anchor_id = href[1:]
            if anchor_id and not soup.find(id=anchor_id):
                broken_links.append(href)
    
    html_stats['total_links'] = len(all_links)
    html_stats['external_links'] = len(external_links)
    html_stats['internal_links'] = len(internal_links)
    html_stats['broken_links'] = len(broken_links)
    
    if invalid_links:
        html_warnings.append(f"⚠️ Подозрительные ссылки: {', '.join(invalid_links[:3])}")
    if broken_links:
        html_warnings.append(f"⚠️ Сломанные якорные ссылки: {', '.join(broken_links[:3])}")
    
    positives.append(f"✅ Ссылок: {len(all_links)} (внешних: {len(external_links)}, внутренних: {len(internal_links)})")
    
    # === CSS АНАЛИЗ ===
    css_errors = []
    css_warnings = []
    css_stats = {}
    
    # Извлечение CSS из тегов style
    style_tags = soup.find_all('style')
    inline_styles = soup.find_all(style=True)
    external_css = soup.find_all('link', rel='stylesheet')
    
    css_code = '\n'.join([tag.string for tag in style_tags if tag.string])
    css_code += '\n'.join([tag['style'] for tag in inline_styles if tag.get('style')])
    
    css_stats['style_blocks'] = len(style_tags)
    css_stats['inline_styles'] = len(inline_styles)
    css_stats['external_css'] = len(external_css)
    
    if css_code:
        # Детальный анализ CSS
        css_lines = css_code.split('\n')
        css_stats['total_css_lines'] = len(css_lines)
        
        # Подсчет CSS правил
        css_rules = re.findall(r'[^{}]+{', css_code)
        css_stats['css_rules'] = len(css_rules)
        
        # Поиск медиа-запросов
        media_queries = re.findall(r'@media[^{]+{', css_code, re.IGNORECASE)
        css_stats['media_queries'] = len(media_queries)
        
        # Поиск анимаций и переходов
        animations = re.findall(r'@keyframes[^{]+{', css_code, re.IGNORECASE)
        transitions = re.findall(r'transition:', css_code, re.IGNORECASE)
        css_stats['animations'] = len(animations)
        css_stats['transitions'] = len(transitions)
        
        # Проверка синтаксиса CSS
        bracket_errors = 0
        semicolon_errors = 0
        
        for i, line in enumerate(css_lines, 1):
            line = line.strip()
            if not line or line.startswith('/*') or line.startswith('//'):
                continue
            
            # Проверка незакрытых скобок
            if line.count('{') != line.count('}'):
                css_warnings.append(f"⚠️ Строка {i}: Несбалансированные скобки в CSS")
                bracket_errors += 1
            
            # Проверка отсутствующих точек с запятой
            if ':' in line and not line.strip().endswith(';') and not line.strip().endswith('{') and not line.strip().endswith('}'):
                css_warnings.append(f"⚠️ Строка {i}: Возможно отсутствует точка с запятой")
                semicolon_errors += 1
        
        css_stats['bracket_errors'] = bracket_errors
        css_stats['semicolon_errors'] = semicolon_errors
        
        positives.append(f"✅ CSS: {len(style_tags)} блоков, {len(css_rules)} правил, {len(media_queries)} медиа-запросов")
        
        if animations:
            positives.append(f"✅ Найдено {len(animations)} анимаций")
        if transitions:
            positives.append(f"✅ Найдено {len(transitions)} переходов")
    else:
        css_warnings.append("⚠️ Не найдено встроенных стилей")
    
    if external_css:
        positives.append(f"✅ Подключено {len(external_css)} внешних CSS файлов")
    
    # === JAVASCRIPT АНАЛИЗ ===
    js_errors = []
    js_warnings = []
    js_stats = {}
    
    # Извлечение JavaScript
    script_tags = soup.find_all('script')
    inline_scripts = [tag.string for tag in script_tags if tag.string]
    external_scripts = [tag.get('src') for tag in script_tags if tag.get('src')]
    
    js_stats['script_blocks'] = len(script_tags)
    js_stats['inline_scripts'] = len(inline_scripts)
    js_stats['external_scripts'] = len(external_scripts)
    
    if inline_scripts:
        js_code = '\n'.join(inline_scripts)
        js_stats['total_js_lines'] = len(js_code.split('\n'))
        
        # Поиск функций
        functions = re.findall(r'function\s+\w+\s*\(', js_code)
        arrow_functions = re.findall(r'const\s+\w+\s*=\s*\([^)]*\)\s*=>', js_code)
        js_stats['functions'] = len(functions)
        js_stats['arrow_functions'] = len(arrow_functions)
        
        # Поиск переменных
        vars = re.findall(r'\bvar\s+\w+', js_code)
        lets = re.findall(r'\blet\s+\w+', js_code)
        consts = re.findall(r'\bconst\s+\w+', js_code)
        js_stats['var_declarations'] = len(vars)
        js_stats['let_declarations'] = len(lets)
        js_stats['const_declarations'] = len(consts)
        
        # Поиск console.log
        console_logs = re.findall(r'console\.log', js_code)
        js_stats['console_logs'] = len(console_logs)
        
        # Проверка синтаксиса JavaScript
        bracket_errors = 0
        semicolon_errors = 0
        
        js_lines = js_code.split('\n')
        for i, line in enumerate(js_lines, 1):
            line = line.strip()
            if not line or line.startswith('//') or line.startswith('/*'):
                continue
            
            # Проверка незакрытых скобок
            if line.count('(') != line.count(')'):
                js_warnings.append(f"⚠️ Строка {i}: Несбалансированные скобки в JS")
                bracket_errors += 1
            
            if line.count('{') != line.count('}'):
                js_warnings.append(f"⚠️ Строка {i}: Несбалансированные фигурные скобки в JS")
                bracket_errors += 1
            
            # Проверка отсутствующих точек с запятой
            if line and not line.endswith(';') and not line.endswith('{') and not line.endswith('}') and not line.endswith('('):
                if any(keyword in line for keyword in ['var ', 'let ', 'const ', 'return', 'console.log']):
                    js_warnings.append(f"⚠️ Строка {i}: Возможно отсутствует точка с запятой")
                    semicolon_errors += 1
        
        js_stats['bracket_errors'] = bracket_errors
        js_stats['semicolon_errors'] = semicolon_errors
        
        positives.append(f"✅ JS: {len(inline_scripts)} блоков, {len(functions)} функций, {len(vars + lets + consts)} переменных")
        
        if console_logs:
            js_warnings.append(f"⚠️ Найдено {len(console_logs)} console.log (рекомендуется убрать в продакшене)")
    
    if external_scripts:
        positives.append(f"✅ Подключено {len(external_scripts)} внешних JS файлов")
    
    # === PHP АНАЛИЗ ===
    php_errors = []
    php_warnings = []
    php_stats = {}
    
    # Поиск PHP кода в HTML
    php_patterns = [
        r'<\?php.*?\?>',
        r'<\?=.*?\?>',
        r'<\?.*?\?>'
    ]
    
    php_code_found = False
    php_blocks = []
    
    for pattern in php_patterns:
        matches = re.findall(pattern, html_content, re.DOTALL | re.IGNORECASE)
        if matches:
            php_code_found = True
            php_blocks.extend(matches)
            
            for match in matches:
                # Проверка синтаксиса PHP
                php_lines = match.split('\n')
                for i, line in enumerate(php_lines, 1):
                    line = line.strip()
                    if not line or line.startswith('//') or line.startswith('#'):
                        continue
                    
                    # Проверка незакрытых скобок
                    if line.count('(') != line.count(')'):
                        php_warnings.append(f"⚠️ PHP: Несбалансированные скобки")
                    
                    if line.count('{') != line.count('}'):
                        php_warnings.append(f"⚠️ PHP: Несбалансированные фигурные скобки")
                    
                    # Проверка отсутствующих точек с запятой
                    if line and not line.endswith(';') and not line.endswith('{') and not line.endswith('}') and not line.endswith('('):
                        if any(keyword in line for keyword in ['echo', 'return', '$', 'function']):
                            php_warnings.append(f"⚠️ PHP: Возможно отсутствует точка с запятой")
    
    php_stats['php_blocks'] = len(php_blocks)
    php_stats['total_php_lines'] = sum(len(block.split('\n')) for block in php_blocks)
    
    if php_code_found:
        positives.append(f"✅ PHP: {len(php_blocks)} блоков, {php_stats['total_php_lines']} строк")
    else:
        php_warnings.append("ℹ️ PHP код не обнаружен")
    
    # === SEO И МЕТА-ТЕГИ АНАЛИЗ ===
    seo_errors = []
    seo_warnings = []
    seo_stats = {}
    
    # Проверка кодировки
    meta_charset = soup.find('meta', charset=True)
    meta_content_type = soup.find('meta', attrs={'http-equiv': 'Content-Type'})
    
    if not meta_charset and not meta_content_type:
        seo_warnings.append("⚠️ Не указана кодировка страницы")
    else:
        charset = meta_charset.get('charset') if meta_charset else meta_content_type.get('content', '').split('charset=')[-1]
        positives.append(f"✅ Кодировка: {charset}")
        seo_stats['charset'] = charset
    
    # Проверка viewport
    viewport = soup.find('meta', attrs={'name': 'viewport'})
    if not viewport:
        seo_warnings.append("⚠️ Отсутствует viewport meta тег")
    else:
        positives.append("✅ Viewport meta тег присутствует")
        seo_stats['viewport'] = viewport.get('content', '')
    
    # Проверка favicon
    favicon = soup.find('link', rel='icon') or soup.find('link', rel='shortcut icon')
    if not favicon:
        seo_warnings.append("⚠️ Отсутствует favicon")
    else:
        positives.append("✅ Favicon подключен")
        seo_stats['favicon'] = favicon.get('href', '')
    
    # Проверка meta description
    meta_description = soup.find('meta', attrs={'name': 'description'})
    if not meta_description:
        seo_warnings.append("⚠️ Отсутствует meta description")
    else:
        desc_text = meta_description.get('content', '')
        if desc_text:
            positives.append(f"✅ Meta description: {desc_text[:50]}{'...' if len(desc_text) > 50 else ''}")
            seo_stats['description_length'] = len(desc_text)
        else:
            seo_warnings.append("⚠️ Meta description пустой")
    
    # Проверка meta keywords
    meta_keywords = soup.find('meta', attrs={'name': 'keywords'})
    if meta_keywords:
        keywords_text = meta_keywords.get('content', '')
        seo_stats['keywords'] = keywords_text
        positives.append("✅ Meta keywords присутствуют")
    
    # Проверка Open Graph
    og_tags = soup.find_all('meta', attrs={'property': re.compile(r'^og:')})
    if og_tags:
        positives.append(f"✅ Open Graph теги: {len(og_tags)}")
        seo_stats['og_tags'] = len(og_tags)
    else:
        seo_warnings.append("⚠️ Отсутствуют Open Graph теги")
    
    # Проверка Twitter Cards
    twitter_tags = soup.find_all('meta', attrs={'name': re.compile(r'^twitter:')})
    if twitter_tags:
        positives.append(f"✅ Twitter Cards: {len(twitter_tags)}")
        seo_stats['twitter_tags'] = len(twitter_tags)
    
    # Проверка canonical
    canonical = soup.find('link', attrs={'rel': 'canonical'})
    if canonical:
        positives.append("✅ Canonical URL указан")
        seo_stats['canonical'] = canonical.get('href', '')
    else:
        seo_warnings.append("⚠️ Отсутствует canonical URL")
    
    # Проверка robots
    robots = soup.find('meta', attrs={'name': 'robots'})
    if robots:
        positives.append("✅ Robots meta тег присутствует")
        seo_stats['robots'] = robots.get('content', '')
    
    # === ПРОИЗВОДИТЕЛЬНОСТЬ ===
    performance_stats = {}
    
    # Подсчет размеров
    performance_stats['html_size_kb'] = len(html_content) / 1024
    performance_stats['css_size_kb'] = len(css_code) / 1024 if css_code else 0
    performance_stats['js_size_kb'] = len(js_code) / 1024 if 'js_code' in locals() else 0
    
    # Рекомендации по производительности
    if performance_stats['html_size_kb'] > 100:
        seo_warnings.append("⚠️ HTML размер превышает 100KB")
    if performance_stats['css_size_kb'] > 50:
        seo_warnings.append("⚠️ CSS размер превышает 50KB")
    if performance_stats['js_size_kb'] > 100:
        seo_warnings.append("⚠️ JavaScript размер превышает 100KB")
    
    # === ОБЩИЕ ПРОВЕРКИ ===
    
    # Подсчет ошибок и предупреждений
    total_errors = len(html_errors) + len(css_errors) + len(js_errors) + len(php_errors) + len(seo_errors)
    total_warnings = len(html_warnings) + len(css_warnings) + len(js_warnings) + len(php_warnings) + len(seo_warnings)
    
    # Оценка качества кода
    quality_score = 100
    quality_score -= total_errors * 10
    quality_score -= total_warnings * 2
    
    # Бонусы за хорошие практики
    if html_stats.get('alt_coverage', 0) == 100:
        quality_score += 5
    if seo_stats.get('description_length', 0) > 0:
        quality_score += 3
    if seo_stats.get('og_tags', 0) > 0:
        quality_score += 2
    if canonical:
        quality_score += 2
    
    quality_score = max(0, min(100, quality_score))
    
    return {
        'url': url,
        'html_errors': html_errors,
        'html_warnings': html_warnings,
        'css_errors': css_errors,
        'css_warnings': css_warnings,
        'js_errors': js_errors,
        'js_warnings': js_warnings,
        'php_errors': php_errors,
        'php_warnings': php_warnings,
        'seo_errors': seo_errors,
        'seo_warnings': seo_warnings,
        'positives': positives,
        'total_errors': total_errors,
        'total_warnings': total_warnings,
        'quality_score': quality_score,
        'html_stats': html_stats,
        'css_stats': css_stats,
        'js_stats': js_stats,
        'php_stats': php_stats,
        'seo_stats': seo_stats,
        'performance_stats': performance_stats,
        'style_blocks': len(style_tags),
        'script_blocks': len(script_tags),
        'external_scripts': len(external_scripts),
        'inline_styles': len(inline_styles),
        'images_without_alt': len(images_without_alt),
        'unclosed_tags': len(set(unclosed_tags)) if unclosed_tags else 0
    }

def check_redirects(urls, ignore_ssl, update_callback, done_callback):
    """Проверяет редиректы для списка URL."""
    import requests
    import threading
    from concurrent.futures import ThreadPoolExecutor, as_completed
    
    results = []
    lock = threading.Lock()
    
    def check_single_url(url):
        try:
            r = requests.get(url, timeout=5, verify=not ignore_ssl, allow_redirects=True)
            chain = [resp.url for resp in r.history] + [r.url] if r.history else [r.url]
            status = r.status_code
            redirected = len(chain) > 1
            return {
                'Исходный URL': url,
                'Редирект': '✅' if redirected else '❌',
                'Конечный URL': chain[-1],
                'HTTP': status,
                'OK': status in (200, 301, 302)
            }
        except Exception as e:
            return {
                'Исходный URL': url,
                'Редирект': '❌',
                'Конечный URL': str(e),
                'HTTP': '-',
                'OK': False
            }
    
    # Используем ThreadPoolExecutor для параллельной обработки
    max_workers = min(20, len(urls))  # Максимум 20 потоков
    completed = 0
    
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_url = {executor.submit(check_single_url, url): url for url in urls}
        
        for future in as_completed(future_to_url):
            result = future.result()
            with lock:
                results.append(result)
                completed += 1
                update_callback(completed, len(urls))
    
    done_callback(results)

def main(page: ft.Page):
    """Основная функция интерфейса Flet с боковым меню."""
    page.title = "🔍 SEO Автотестер"
    page.window_favicon = "assets/favicon.png"
    page.scroll = ft.ScrollMode.AUTO
    page.window.width = 1200
    page.window.height = 800
    page.theme_mode = ft.ThemeMode.DARK
    page.theme = ft.Theme(
        color_scheme_seed="#394459",
        use_material3=True,
    )
    page.data = {}

    # --- Функции для получения цветов в зависимости от темы ---
    def get_text_color():
        """Возвращает цвет текста для заголовков и описаний."""
        return "#FFFFFF"  # Белый цвет для заголовков и описаний
    
    def get_label_color():
        """Возвращает цвет лейблов для полей ввода."""
        return "#394459"  # Темный цвет для лейблов полей ввода
    
    def get_secondary_text_color():
        """Возвращает цвет вторичного текста."""
        return "#FFFFFF"  # Белый цвет для описаний
    
    def get_input_text_color():
        """Возвращает цвет текста в полях ввода."""
        return "#394459"  # Темный цвет для текста в полях ввода

    # --- Top Navigation Menu ---
    def nav_home(e):
        switch_page(0)
    
    def nav_links_check(e):
        switch_page(1)
        # Обновляем видимость кнопок экспорта при переключении на вкладку
        update_links_export_buttons()
    
    def nav_parser(e):
        switch_page(2)
    
    def nav_text_analysis(e):
        switch_page(3)
    
    def nav_code_analysis(e):
        switch_page(4)
    
    def nav_redirects(e):
        switch_page(5)
    
    def nav_competitors(e):
        switch_page(6)
    
    def nav_exports(e):
        switch_page(7)
    
    def nav_serp_tracker(e):
        switch_page(8)
    
    nav_buttons = [
        ft.ElevatedButton(
            "🏠 Главная", 
            on_click=nav_home,
            style=ft.ButtonStyle(
                bgcolor="#F2F2F2",
                color="#394459",
                shape=ft.RoundedRectangleBorder(radius=10),
                elevation=3
            )
        ),
        ft.ElevatedButton(
            "🔗 Проверка ссылок", 
            on_click=nav_links_check,
            style=ft.ButtonStyle(
                bgcolor="#F2F2F2",
                color="#394459",
                shape=ft.RoundedRectangleBorder(radius=10),
                elevation=3
            )
        ),
        ft.ElevatedButton(
            "📋 Парсер всех страниц", 
            on_click=nav_parser,
            style=ft.ButtonStyle(
                bgcolor="#F2F2F2",
                color="#394459",
                shape=ft.RoundedRectangleBorder(radius=10),
                elevation=3
            )
        ),
        ft.ElevatedButton(
            "📝 Анализ текста", 
            on_click=nav_text_analysis,
            style=ft.ButtonStyle(
                bgcolor="#F2F2F2",
                color="#394459",
                shape=ft.RoundedRectangleBorder(radius=10),
                elevation=3
            )
        ),
        ft.ElevatedButton(
            "💻 Анализ кода", 
            on_click=nav_code_analysis,
            style=ft.ButtonStyle(
                bgcolor="#F2F2F2",
                color="#394459",
                shape=ft.RoundedRectangleBorder(radius=10),
                elevation=3
            )
        ),
        ft.ElevatedButton(
            "🔄 Редиректы", 
            on_click=nav_redirects,
            style=ft.ButtonStyle(
                bgcolor="#F2F2F2",
                color="#394459",
                shape=ft.RoundedRectangleBorder(radius=10),
                elevation=3
            )
        ),
        ft.ElevatedButton(
            "👥 Анализ конкурентов", 
            on_click=nav_competitors,
            style=ft.ButtonStyle(
                bgcolor="#F2F2F2",
                color="#394459",
                shape=ft.RoundedRectangleBorder(radius=10),
                elevation=3
            )
        ),
        ft.ElevatedButton(
            "📁 Экспорт", 
            on_click=nav_exports,
            style=ft.ButtonStyle(
                bgcolor="#F2F2F2",
                color="#394459",
                shape=ft.RoundedRectangleBorder(radius=10),
                elevation=3
            )
        ),
        ft.ElevatedButton(
            "📊 SERP Tracker", 
            on_click=nav_serp_tracker,
            style=ft.ButtonStyle(
                bgcolor="#F2F2F2",
                color="#394459",
                shape=ft.RoundedRectangleBorder(radius=10),
                elevation=3
            )
        ),
    ]
    
    # Создаем контейнер для навигации
    navigation_bar = ft.Container(
        content=ft.Row(
            nav_buttons,
            alignment=ft.MainAxisAlignment.CENTER,
            spacing=10
        ),
        padding=15,
        bgcolor="#394459",
        border_radius=15,
        margin=ft.Margin(0, 0, 0, 20),
        border=ft.border.all(2, "#F2E307")
    )

    # --- Контейнеры для страниц ---
    main_content = ft.Container(visible=True)
    links_check_content = ft.Container(visible=False)
    parser_content = ft.Container(visible=False)
    text_analysis_content = ft.Container(visible=False)
    code_analysis_content = ft.Container(visible=False)
    redirects_content = ft.Container(visible=False)
    competitors_content = ft.Container(visible=False)
    exports_content = ft.Container(visible=False)
    serp_tracker_content = ft.Container(visible=False)

    def switch_page(idx):
        # Скрываем все страницы
        main_content.visible = False
        links_check_content.visible = False
        parser_content.visible = False
        text_analysis_content.visible = False
        code_analysis_content.visible = False
        redirects_content.visible = False
        competitors_content.visible = False
        exports_content.visible = False
        serp_tracker_content.visible = False
        
        # Сбрасываем стили всех кнопок
        for btn in nav_buttons:
            btn.style = None
            btn.bgcolor = None
        
        # Подсвечиваем активную кнопку
        if 0 <= idx < len(nav_buttons):
            nav_buttons[idx].bgcolor = "#F2E307"
            nav_buttons[idx].color = "#394459"
            nav_buttons[idx].style = ft.ButtonStyle(
                shape=ft.RoundedRectangleBorder(radius=10),
                elevation=5
            )
        
        # Показываем нужную страницу
        if idx == 0:
            main_content.visible = True
            # При переключении на главную страницу проверяем, есть ли данные
            if not page.data:
                summary_area.value = ""
                export_btn.visible = False
                export_word_btn.visible = False
                progress_bar.value = 0.0
        elif idx == 1:
            links_check_content.visible = True
            # Обновляем видимость кнопок экспорта при переключении на вкладку проверки ссылок
            update_links_export_buttons()
        elif idx == 2:
            parser_content.visible = True
            parser_status.value = "Нажмите 'Запустить парсер', чтобы начать обход сайта"
            parser_status.visible = True
            parser_table.rows = []
            parser_export_btn.visible = False
        elif idx == 3:
            text_analysis_content.visible = True
        elif idx == 4:
            code_analysis_content.visible = True
        elif idx == 5:
            redirects_content.visible = True
        elif idx == 6:
            competitors_content.visible = True
        elif idx == 7:
            exports_content.visible = True
            refresh_exports_list()
        elif idx == 8:
            serp_tracker_content.visible = True
        
        page.update()

    # --- Главная страница (минималистичный функционал с одной сводкой и кнопками) ---
    url_input = ft.TextField(
        label="URL сайта", 
        width=400, 
        filled=True, 
        border_radius=10,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color="#394459",
        label_style=ft.TextStyle(color="#394459")
    )
    ssl_checkbox = ft.Checkbox(label="Игнорировать SSL", value=True)
    run_btn = ft.ElevatedButton(
        "Запустить тест", 
        icon=ft.Icons.PLAY_ARROW,
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=10),
            elevation=5
        )
    )
    stop_btn = ft.ElevatedButton(
        "Остановить", 
        icon=ft.Icons.STOP,
        visible=False,
        style=ft.ButtonStyle(
            bgcolor="#FF5722",
            color="white",
            shape=ft.RoundedRectangleBorder(radius=10),
            elevation=5
        )
    )
    progress_bar = ft.ProgressBar(width=600, color="#F2E307", bgcolor="#394459", value=0.0, height=10, border_radius=20)
    summary_area = ft.TextField(
        label="Сводка", 
        multiline=True, 
        min_lines=16, 
        max_lines=30, 
        width=1000, 
        filled=True, 
        border_radius=10, 
        expand=True,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color="#394459",
        label_style=ft.TextStyle(color="#394459")
    )

    # Кнопки для разных сводок
    seo_btn = ft.ElevatedButton(
        "SEO", 
        icon=ft.Icons.SEARCH,
        style=ft.ButtonStyle(
            bgcolor="#F2CC0C",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    robots_btn = ft.ElevatedButton(
        "Robots", 
        icon=ft.Icons.ANDROID,
        style=ft.ButtonStyle(
            bgcolor="#F2CC0C",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    sitemap_btn = ft.ElevatedButton(
        "Sitemap", 
        icon=ft.Icons.MAP,
        style=ft.ButtonStyle(
            bgcolor="#F2CC0C",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    links_btn = ft.ElevatedButton(
        "Ссылки", 
        icon=ft.Icons.LINK,
        style=ft.ButtonStyle(
            bgcolor="#F2CC0C",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    images_btn = ft.ElevatedButton(
        "Изображения", 
        icon=ft.Icons.IMAGE,
        style=ft.ButtonStyle(
            bgcolor="#F2CC0C",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    full_btn = ft.ElevatedButton(
        "Общая сводка", 
        icon=ft.Icons.ASSIGNMENT,
        style=ft.ButtonStyle(
            bgcolor="#F2CC0C",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    clear_btn = ft.ElevatedButton(
        "Очистить сводку", 
        icon=ft.Icons.CLEAR, 
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    export_btn = ft.ElevatedButton(
        "Экспорт в Excel", 
        icon=ft.Icons.DOWNLOAD, 
        visible=False,
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    export_word_btn = ft.ElevatedButton(
        "Экспорт в Word", 
        icon=ft.Icons.DESCRIPTION, 
        visible=False,
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    
    # Кнопки для экспорта полного sitemap
    sitemap_full_excel_btn = ft.ElevatedButton(
        "📊 Полный Sitemap Excel", 
        icon=ft.Icons.DOWNLOAD, 
        visible=False,
        style=ft.ButtonStyle(
            bgcolor="#4CAF50",
            color="white",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    sitemap_full_word_btn = ft.ElevatedButton(
        "📄 Полный Sitemap Word", 
        icon=ft.Icons.DESCRIPTION, 
        visible=False,
        style=ft.ButtonStyle(
            bgcolor="#2196F3",
            color="white",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )

    # --- Страница проверки ссылок ---
    links_ssl_checkbox = ft.Checkbox(label="Игнорировать SSL", value=True)
    links_multiple_input = ft.TextField(
        label="Ссылки для проверки (по одной на строку)", 
        width=800, 
        filled=True, 
        border_radius=10,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color="#394459",
        label_style=ft.TextStyle(color="#394459"),
        multiline=True,
        min_lines=5,
        max_lines=12,
        hint_text="https://example.com\nhttps://example.com/page1\nhttps://example.com/page2\n\nВведите ссылки для проверки (по одной на строку)"
    )
    links_run_btn = ft.ElevatedButton(
        "Запустить проверку ссылок", 
        icon=ft.Icons.PLAY_ARROW,
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=10),
            elevation=5
        )
    )
    links_stop_btn = ft.ElevatedButton(
        "Остановить", 
        icon=ft.Icons.STOP,
        visible=False,
        style=ft.ButtonStyle(
            bgcolor="#FF5722",
            color="white",
            shape=ft.RoundedRectangleBorder(radius=10),
            elevation=5
        )
    )
    links_progress_bar = ft.ProgressBar(width=600, color="#F2E307", bgcolor="#394459", value=0.0, height=10, border_radius=20)
    links_summary_area = ft.TextField(
        label="Сводка", 
        multiline=True, 
        min_lines=16, 
        max_lines=30, 
        width=1000, 
        filled=True, 
        border_radius=10, 
        expand=True,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color="#394459",
        label_style=ft.TextStyle(color="#394459")
    )

    # Кнопки для разных сводок на странице ссылок
    links_seo_btn = ft.ElevatedButton(
        "SEO", 
        icon=ft.Icons.SEARCH,
        style=ft.ButtonStyle(
            bgcolor="#F2CC0C",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    links_links_btn = ft.ElevatedButton(
        "Ссылки", 
        icon=ft.Icons.LINK,
        style=ft.ButtonStyle(
            bgcolor="#F2CC0C",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    links_images_btn = ft.ElevatedButton(
        "Изображения", 
        icon=ft.Icons.IMAGE,
        style=ft.ButtonStyle(
            bgcolor="#F2CC0C",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    links_full_btn = ft.ElevatedButton(
        "Общая сводка", 
        icon=ft.Icons.ASSIGNMENT,
        style=ft.ButtonStyle(
            bgcolor="#F2CC0C",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    links_clear_btn = ft.ElevatedButton(
        "Очистить сводку", 
        icon=ft.Icons.CLEAR, 
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    links_export_btn = ft.ElevatedButton(
        "Экспорт в Excel", 
        icon=ft.Icons.DOWNLOAD, 
        visible=False,
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    links_export_word_btn = ft.ElevatedButton(
        "Экспорт в Word", 
        icon=ft.Icons.DESCRIPTION, 
        visible=False,
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )



    main_content.content = ft.Column([
        ft.Text("🔍 SEO Автотестер", size=24, weight=ft.FontWeight.BOLD),
        ft.Text("Введите адрес сайта для анализа и нажмите 'Запустить тест'", size=16),
        ft.Row([url_input, ssl_checkbox, run_btn, stop_btn], spacing=10),
        progress_bar,
        ft.Row([seo_btn, robots_btn, sitemap_btn, links_btn, images_btn, full_btn, clear_btn], spacing=10),
        summary_area,
        ft.Row([export_btn, export_word_btn]),
        ft.Row([sitemap_full_excel_btn, sitemap_full_word_btn])
    ], expand=True)

    links_check_content.content = ft.Column([
        ft.Text("🔗 Проверка ссылок", size=24, weight=ft.FontWeight.BOLD),
        ft.Text("Проверка ссылок без robots и sitemap. Введите ссылки для проверки (по одной на строку).", size=16),
        ft.Row([links_ssl_checkbox, links_run_btn, links_stop_btn], spacing=10),
        links_multiple_input,
        links_progress_bar,
        ft.Row([links_seo_btn, links_links_btn, links_images_btn, links_full_btn, links_clear_btn], spacing=10),
        links_summary_area,
        ft.Row([links_export_btn, links_export_word_btn])
    ], expand=True)

    # --- SERP Tracker UI элементы ---
    serp_domain_input = ft.TextField(
        label="Домен сайта", 
        width=400, 
        filled=True, 
        border_radius=10,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color="#394459",
        label_style=ft.TextStyle(color="#394459"),
        hint_text="example.com"
    )
    
    serp_keywords_input = ft.TextField(
        label="Ключевые слова (необязательно)", 
        width=400, 
        filled=True, 
        border_radius=10,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color="#394459",
        label_style=ft.TextStyle(color="#394459"),
        multiline=True,
        min_lines=3,
        max_lines=5,
        hint_text="купить цветы\nдоставка цветов\nцветы москва\n\n💡 Оставьте пустым для автоматического поиска ключевых слов и детального анализа позиций"
    )
    
    serp_engines_dropdown = ft.Dropdown(
        label="Поисковые системы",
        width=200,
        options=[
            ft.dropdown.Option("google", "Google"),
            ft.dropdown.Option("yandex", "Яндекс"),
            ft.dropdown.Option("both", "Google + Яндекс")
        ],
        value="both"
    )
    
    serp_run_btn = ft.ElevatedButton(
        "🚀 Запустить анализ", 
        icon=ft.Icons.PLAY_ARROW,
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=10),
            elevation=5
        )
    )
    
    serp_stop_btn = ft.ElevatedButton(
        "⏹ Остановить", 
        icon=ft.Icons.STOP,
        visible=False,
        style=ft.ButtonStyle(
            bgcolor="#FF5722",
            color="white",
            shape=ft.RoundedRectangleBorder(radius=10),
            elevation=5
        )
    )
    
    serp_progress_bar = ft.ProgressBar(width=600, color="#F2E307", bgcolor="#394459", value=0.0, height=10, border_radius=20)
    
    serp_results_area = ft.TextField(
        label="Результаты трекинга", 
        multiline=True, 
        min_lines=16, 
        max_lines=25,
        read_only=True,
        border_color="#394459",
        bgcolor="#F2F2F2",
        color="#394459"
    )
    
    serp_export_btn = ft.ElevatedButton(
        "📊 Экспорт в Excel",
        visible=False,
        style=ft.ButtonStyle(
            bgcolor="#4CAF50",
            color="white",
            shape=ft.RoundedRectangleBorder(radius=10),
            elevation=5
        )
    )
    
    serp_chart_container = ft.Container(
        content=ft.Text("График появится после проверки позиций", 
                       color=ft.Colors.GREY_500, size=14),
        alignment=ft.alignment.center,
        padding=20
    )

    serp_tracker_content.content = ft.Column([
        ft.Text("📊 SERP Tracker - Подробный анализ позиций", size=24, weight=ft.FontWeight.BOLD),
        ft.Text("Автоматический поиск ключевых слов и детальная проверка позиций в Google и Яндекс", size=16),
        ft.Row([serp_domain_input, serp_engines_dropdown], spacing=10),
        serp_keywords_input,
        ft.Row([serp_run_btn, serp_stop_btn], spacing=10),
        serp_progress_bar,
        serp_results_area,
        ft.Row([serp_export_btn]),
        serp_chart_container
    ], expand=True)

    # --- Обработчики кнопок ---
    def show_seo(e):
        summary = page.data.get('seo_summary', 'Нет данных по SEO')
        summary_area.value = summary
        export_btn.visible = True
        export_word_btn.visible = True
        sitemap_full_excel_btn.visible = False
        sitemap_full_word_btn.visible = False
        export_btn.data = ('seo', summary)
        page.update()
    seo_btn.on_click = show_seo

    def show_robots(e):
        summary = page.data.get('robots_summary', 'Нет данных по robots.txt')
        summary_area.value = summary
        export_btn.visible = True
        export_word_btn.visible = True
        sitemap_full_excel_btn.visible = False
        sitemap_full_word_btn.visible = False
        export_btn.data = ('robots', summary)
        page.update()
    robots_btn.on_click = show_robots

    def show_sitemap(e):
        summary = page.data.get('sitemap_summary', 'Нет данных по sitemap.xml')
        summary_area.value = summary
        export_btn.visible = True
        export_word_btn.visible = True
        sitemap_full_excel_btn.visible = True
        sitemap_full_word_btn.visible = True
        export_btn.data = ('sitemap', summary)
        page.update()
    sitemap_btn.on_click = show_sitemap

    def show_links(e):
        summary = page.data.get('links_summary', 'Нет данных по ссылкам')
        summary_area.value = summary
        export_btn.visible = True
        export_word_btn.visible = True
        sitemap_full_excel_btn.visible = False
        sitemap_full_word_btn.visible = False
        export_btn.data = ('links', summary)
        page.update()
    links_btn.on_click = show_links

    def show_images(e):
        summary = page.data.get('images_summary', 'Нет данных по изображениям')
        summary_area.value = summary
        export_btn.visible = True
        export_word_btn.visible = True
        sitemap_full_excel_btn.visible = False
        sitemap_full_word_btn.visible = False
        export_btn.data = ('images', summary)
        page.update()
    images_btn.on_click = show_images

    def show_full(e):
        summary = page.data.get('full_summary', 'Нет общей сводки')
        summary_area.value = summary
        export_btn.visible = True
        export_word_btn.visible = True
        sitemap_full_excel_btn.visible = False
        sitemap_full_word_btn.visible = False
        export_btn.data = ('full', summary)
        page.update()
    full_btn.on_click = show_full

    # --- Обработчики кнопок для страницы проверки ссылок ---
    def links_show_seo(e):
        # Проверяем, есть ли данные для множественных сайтов
        if 'multiple_seo_summary' in page.data:
            summary = page.data.get('multiple_seo_summary', 'Нет данных по SEO')
        else:
            summary = page.data.get('seo_summary', 'Нет данных по SEO')
        links_summary_area.value = summary
        page.data['links_export_btn_visible'] = True
        page.data['links_export_word_btn_visible'] = True
        links_export_btn.data = ('seo', summary)
        update_links_export_buttons()
    links_seo_btn.on_click = links_show_seo

    def links_show_links(e):
        # Проверяем, есть ли данные для множественных сайтов
        if 'multiple_links_summary' in page.data:
            detailed_summary = page.data.get('multiple_links_summary', 'Нет данных по ссылкам')
        else:
            link_statuses = page.data.get('link_statuses', {})
            if link_statuses:
                # Создаем детальную сводку ссылок
                detailed_summary = "### Детальная информация о ссылках\n\n"
                for url, status in link_statuses.items():
                    status_emoji = "🟢" if isinstance(status, int) and status == 200 else "🔴"
                    detailed_summary += f"{status_emoji} **{url}**\n"
                    detailed_summary += f"   Статус: {status}\n"
                    
                    # Добавляем дополнительную информацию о ссылке
                    try:
                        result = check_resource(url, links_ssl_checkbox.value)
                        href, status_code, headers, size = result
                        detailed_summary += f"   Размер: {size:.2f} КБ\n"
                        
                        if headers:
                            detailed_summary += f"   Content-Type: {headers.get('content-type', 'Не указан')}\n"
                            detailed_summary += f"   Server: {headers.get('server', 'Не указан')}\n"
                        
                        # Проверяем редирект
                        try:
                            response = requests.head(url, timeout=5, verify=not links_ssl_checkbox.value, allow_redirects=False)
                            if response.status_code in [301, 302, 303, 307, 308]:
                                redirect_url = response.headers.get('Location', 'Не указан')
                                detailed_summary += f"   🔄 Редирект на: {redirect_url}\n"
                        except:
                            pass
                            
                    except Exception as e:
                        detailed_summary += f"   ❌ Ошибка проверки: {str(e)}\n"
                    
                    detailed_summary += "\n"
            else:
                detailed_summary = "Нет данных по ссылкам"
        
        links_summary_area.value = detailed_summary
        page.data['links_export_btn_visible'] = True
        page.data['links_export_word_btn_visible'] = True
        links_export_btn.data = ('links', detailed_summary)
        update_links_export_buttons()
    links_links_btn.on_click = links_show_links

    def links_show_images(e):
        # Проверяем, есть ли данные для множественных сайтов
        if 'multiple_images_summary' in page.data:
            summary = page.data.get('multiple_images_summary', 'Нет данных по изображениям')
        else:
            summary = page.data.get('images_summary', 'Нет данных по изображениям')
        links_summary_area.value = summary
        page.data['links_export_btn_visible'] = True
        page.data['links_export_word_btn_visible'] = True
        links_export_btn.data = ('images', summary)
        update_links_export_buttons()
    links_images_btn.on_click = links_show_images

    def links_show_full(e):
        # Проверяем, есть ли данные для множественных сайтов
        if 'multiple_full_summary' in page.data:
            summary = page.data.get('multiple_full_summary', 'Нет общей сводки')
        else:
            summary = page.data.get('full_summary', 'Нет общей сводки')
        links_summary_area.value = summary
        page.data['links_export_btn_visible'] = True
        page.data['links_export_word_btn_visible'] = True
        links_export_btn.data = ('full', summary)
        update_links_export_buttons()
    links_full_btn.on_click = links_show_full

    def links_clear_summary(e):
        # Очищаем все данные в page.data
        page.data.clear()
        
        # Очищаем область сводки
        links_summary_area.value = ""
        
        # Скрываем кнопки экспорта
        page.data['links_export_btn_visible'] = False
        page.data['links_export_word_btn_visible'] = False
        
        # Скрываем кнопку остановки и показываем кнопку запуска
        links_stop_btn.visible = False
        links_run_btn.visible = True
        
        # Сбрасываем прогресс-бар
        links_progress_bar.value = 0.0
        
        # Обновляем интерфейс
        update_links_export_buttons()
        
        # Показываем уведомление
        page.snack_bar = ft.SnackBar(content=ft.Text("Сводка очищена"))
        page.snack_bar.open = True
    links_clear_btn.on_click = links_clear_summary

    def links_export_summary(e):
        # Безопасно получаем данные для экспорта
        if hasattr(links_export_btn, 'data') and links_export_btn.data is not None:
            try:
                report_type, summary = links_export_btn.data
            except (ValueError, TypeError):
                report_type, summary = 'full', links_summary_area.value
        else:
            report_type, summary = 'full', links_summary_area.value
        
        # Проверяем, что есть данные для экспорта
        if not summary or summary.strip() == '':
            page.snack_bar = ft.SnackBar(content=ft.Text("❌ Нет данных для экспорта. Сначала запустите анализ."))
            page.snack_bar.open = True
            page.update()
            return
        
        try:
            report_path = generate_report(summary, "multiple_links", report_type=report_type)
            if report_path:
                page.snack_bar = ft.SnackBar(content=ft.Text(f"✅ Отчет сохранен: {os.path.basename(report_path)}"))
                page.snack_bar.open = True
                # Обновляем список экспортов если открыта страница экспорта
                if exports_content.visible:
                    refresh_exports_list()
        except Exception as e:
            page.snack_bar = ft.SnackBar(content=ft.Text(f"❌ Ошибка экспорта: {str(e)}"))
            page.snack_bar.open = True
        page.update()
    links_export_btn.on_click = links_export_summary

    def links_export_summary_word(e):
        # Безопасно получаем данные для экспорта
        if hasattr(links_export_word_btn, 'data') and links_export_word_btn.data is not None:
            try:
                report_type, summary = links_export_word_btn.data
            except (ValueError, TypeError):
                report_type, summary = 'full', links_summary_area.value
        else:
            report_type, summary = 'full', links_summary_area.value
        
        # Проверяем, что есть данные для экспорта
        if not summary or summary.strip() == '':
            page.snack_bar = ft.SnackBar(content=ft.Text("❌ Нет данных для экспорта. Сначала запустите анализ."))
            page.snack_bar.open = True
            page.update()
            return
        
        try:
            report_path = generate_word_report(summary, "multiple_links", report_type=report_type)
            if report_path:
                page.snack_bar = ft.SnackBar(content=ft.Text(f"✅ Отчет сохранен: {os.path.basename(report_path)}"))
                page.snack_bar.open = True
                # Обновляем список экспортов если открыта страница экспорта
                if exports_content.visible:
                    refresh_exports_list()
            else:
                page.snack_bar = ft.SnackBar(content=ft.Text("❌ Ошибка: модуль python-docx не установлен"))
                page.snack_bar.open = True
        except Exception as e:
            page.snack_bar = ft.SnackBar(content=ft.Text(f"❌ Ошибка экспорта: {str(e)}"))
            page.snack_bar.open = True
        page.update()
    links_export_word_btn.on_click = links_export_summary_word

    def show_link_detail(e):
        """Показывает детальную информацию о ссылке."""
        url = e.control.data
        if url:
            # Создаем детальную информацию о ссылке
            detail_info = f"🔗 Детали ссылки: {url}\n\n"
            
            # Проверяем ссылку
            try:
                result = check_resource(url, links_ssl_checkbox.value)
                href, status, headers, size = result
                
                detail_info += f"📊 Статус: {status}\n"
                detail_info += f"📏 Размер: {size} КБ\n"
                detail_info += f"📋 Заголовки:\n"
                
                if headers:
                    for header, value in headers.items():
                        detail_info += f"   {header}: {value}\n"
                else:
                    detail_info += "   Заголовки не получены\n"
                
                # Проверяем редирект
                try:
                    response = requests.head(url, timeout=10, verify=not links_ssl_checkbox.value, allow_redirects=False)
                    if response.status_code in [301, 302, 303, 307, 308]:
                        redirect_url = response.headers.get('Location', 'Не указан')
                        detail_info += f"🔄 Редирект на: {redirect_url}\n"
                except:
                    detail_info += "🔄 Информация о редиректе недоступна\n"
                
            except Exception as e:
                detail_info += f"❌ Ошибка проверки: {str(e)}\n"
            
            # Показываем детали в сводке
            links_summary_area.value = detail_info
            page.update()
            
            # Показываем уведомление
            page.snack_bar = ft.SnackBar(content=ft.Text(f"Показаны детали для: {url}"))
            page.snack_bar.open = True
            page.update()
    

    


    def export_summary(e):
        # Безопасно получаем данные для экспорта
        if hasattr(export_btn, 'data') and export_btn.data is not None:
            try:
                report_type, summary = export_btn.data
            except (ValueError, TypeError):
                report_type, summary = 'full', summary_area.value
        else:
            report_type, summary = 'full', summary_area.value
        
        try:
            report_path = generate_report(summary, url_input.value.strip(), report_type=report_type)
            if report_path:
                page.snack_bar = ft.SnackBar(content=ft.Text(f"✅ Отчет сохранен: {os.path.basename(report_path)}"))
                page.snack_bar.open = True
                # Обновляем список экспортов если открыта страница экспорта
                if exports_content.visible:
                    refresh_exports_list()
            else:
                page.snack_bar = ft.SnackBar(content=ft.Text("❌ Ошибка при сохранении отчета"))
                page.snack_bar.open = True
        except Exception as ex:
            page.snack_bar = ft.SnackBar(content=ft.Text(f"❌ Ошибка экспорта: {str(ex)}"))
            page.snack_bar.open = True
        page.update()
    export_btn.on_click = export_summary

    def export_summary_word(e):
        # Безопасно получаем данные для экспорта
        if hasattr(export_btn, 'data') and export_btn.data is not None:
            try:
                report_type, summary = export_btn.data
            except (ValueError, TypeError):
                report_type, summary = 'full', summary_area.value
        else:
            report_type, summary = 'full', summary_area.value
        
        try:
            report_path = generate_word_report(summary, url_input.value.strip(), report_type)
            if report_path:
                page.snack_bar = ft.SnackBar(content=ft.Text(f"✅ Отчет сохранен: {os.path.basename(report_path)}"))
                page.snack_bar.open = True
                # Обновляем список экспортов если открыта страница экспорта
                if exports_content.visible:
                    refresh_exports_list()
            else:
                page.snack_bar = ft.SnackBar(content=ft.Text("❌ Ошибка: модуль python-docx не установлен"))
                page.snack_bar.open = True
        except Exception as ex:
            page.snack_bar = ft.SnackBar(content=ft.Text(f"❌ Ошибка экспорта: {str(ex)}"))
            page.snack_bar.open = True
        page.update()
    export_word_btn.on_click = export_summary_word

    def export_sitemap_full_excel(e):
        try:
            report_path = generate_sitemap_excel_report(url_input.value.strip(), ssl_checkbox.value)
            if report_path:
                page.snack_bar = ft.SnackBar(content=ft.Text(f"✅ Полный отчет Sitemap сохранен: {os.path.basename(report_path)}"))
                page.snack_bar.open = True
                # Обновляем список экспортов если открыта страница экспорта
                if exports_content.visible:
                    refresh_exports_list()
            else:
                page.snack_bar = ft.SnackBar(content=ft.Text("❌ Нет данных sitemap для экспорта"))
                page.snack_bar.open = True
        except Exception as ex:
            page.snack_bar = ft.SnackBar(content=ft.Text(f"❌ Ошибка экспорта: {str(ex)}"))
            page.snack_bar.open = True
        page.update()
    sitemap_full_excel_btn.on_click = export_sitemap_full_excel

    def export_sitemap_full_word(e):
        try:
            report_path = generate_sitemap_word_report(url_input.value.strip(), ssl_checkbox.value)
            if report_path:
                page.snack_bar = ft.SnackBar(content=ft.Text(f"✅ Полный отчет Sitemap сохранен: {os.path.basename(report_path)}"))
                page.snack_bar.open = True
                # Обновляем список экспортов если открыта страница экспорта
                if exports_content.visible:
                    refresh_exports_list()
            else:
                page.snack_bar = ft.SnackBar(content=ft.Text("❌ Нет данных sitemap для экспорта или модуль python-docx не установлен"))
                page.snack_bar.open = True
        except Exception as ex:
            page.snack_bar = ft.SnackBar(content=ft.Text(f"❌ Ошибка экспорта: {str(ex)}"))
            page.snack_bar.open = True
        page.update()
    sitemap_full_word_btn.on_click = export_sitemap_full_word

    def run_main_test(e):
        # Показываем кнопку остановки и скрываем кнопку запуска
        stop_btn.visible = True
        run_btn.visible = False
        page.update()
        
        # Создаем событие для остановки
        page.data['stop_event'] = threading.Event()
        
        run_test(
            url_input.value.strip(),
            summary_area,  # Используем только summary_area для вывода
            page,
            progress_bar,
            ssl_checkbox.value,
            ""  # Убираем ключевые слова
        )
    run_btn.on_click = run_main_test

    def stop_main_test(e):
        """Останавливает основной тест."""
        if 'stop_event' in page.data:
            page.data['stop_event'].set()
        
        # Скрываем кнопку остановки и показываем кнопку запуска
        stop_btn.visible = False
        run_btn.visible = True
        page.update()
        
        # Показываем уведомление
        page.snack_bar = ft.SnackBar(content=ft.Text("Тест остановлен"))
        page.snack_bar.open = True
        page.update()
    stop_btn.on_click = stop_main_test

    def update_links_export_buttons():
        """Обновляет видимость кнопок экспорта на основе флагов в page.data."""
        if page.data.get('links_export_btn_visible', False):
            links_export_btn.visible = True
        else:
            links_export_btn.visible = False
            
        if page.data.get('links_export_word_btn_visible', False):
            links_export_word_btn.visible = True
        else:
            links_export_word_btn.visible = False
        
        page.update()

    def run_links_test_handler(e):
        """Запускает проверку ссылок."""
        multiple_urls_text = links_multiple_input.value.strip()
        
        # Показываем кнопку остановки и скрываем кнопку запуска
        links_stop_btn.visible = True
        links_run_btn.visible = False
        page.update()
        
        # Создаем событие для остановки
        page.data['stop_event'] = threading.Event()
        
        # Проверяем, есть ли ссылки для проверки
        if multiple_urls_text:
            urls = [url.strip() for url in multiple_urls_text.split('\n') if url.strip()]
            if not urls:
                page.snack_bar = ft.SnackBar(content=ft.Text("❌ Введите ссылки для проверки"))
                page.snack_bar.open = True
                page.update()
                return
            
            links_progress_bar.value = 0.0
            links_summary_area.value = f"🔄 Запуск проверки {len(urls)} ссылок..."
            page.data['links_export_btn_visible'] = False
            page.data['links_export_word_btn_visible'] = False
            update_links_export_buttons()
            
            # Запуск проверки множественных ссылок в отдельном потоке
            threading.Thread(
                target=run_multiple_links_test,
                args=(urls, links_summary_area, page, links_progress_bar, links_ssl_checkbox.value, ""),
                daemon=True
            ).start()
        else:
            page.snack_bar = ft.SnackBar(content=ft.Text("❌ Введите ссылки для проверки"))
            page.snack_bar.open = True
            page.update()
            return
    links_run_btn.on_click = run_links_test_handler

    def stop_links_test(e):
        """Останавливает проверку ссылок."""
        if 'stop_event' in page.data:
            page.data['stop_event'].set()
        
        # Скрываем кнопку остановки и показываем кнопку запуска
        links_stop_btn.visible = False
        links_run_btn.visible = True
        
        # Скрываем кнопки экспорта
        page.data['links_export_btn_visible'] = False
        page.data['links_export_word_btn_visible'] = False
        update_links_export_buttons()
        
        # Показываем уведомление
        page.snack_bar = ft.SnackBar(content=ft.Text("Проверка ссылок остановлена"))
        page.snack_bar.open = True
        page.update()
    links_stop_btn.on_click = stop_links_test

    def clear_summary(e):
        # Очищаем все данные в page.data
        page.data.clear()
        
        # Очищаем область сводки
        summary_area.value = ""
        
        # Скрываем кнопки экспорта
        export_btn.visible = False
        export_word_btn.visible = False
        sitemap_full_excel_btn.visible = False
        sitemap_full_word_btn.visible = False
        
        # Сбрасываем прогресс-бар
        progress_bar.value = 0.0
        
        # Переключаемся на главную страницу
        switch_page(0)
        
        # Обновляем интерфейс
        page.update()
        
        # Показываем уведомление
        page.snack_bar = ft.SnackBar(content=ft.Text("Сводка очищена"))
        page.snack_bar.open = True
    clear_btn.on_click = clear_summary

    # --- Обработчики для SERP Tracker ---
    def run_serp_tracking(e):
        """Запускает трекинг позиций или анализ сайта."""
        domain = serp_domain_input.value.strip()
        keywords_text = serp_keywords_input.value.strip()
        engines = serp_engines_dropdown.value
        
        if not domain:
            page.snack_bar = ft.SnackBar(content=ft.Text("❌ Заполните домен сайта"))
            page.snack_bar.open = True
            page.update()
            return
        
        # Определяем поисковые системы
        if engines == "google":
            search_engines = ["google"]
        elif engines == "yandex":
            search_engines = ["yandex"]
        else:
            search_engines = ["google", "yandex"]
        
        # Если ключевые слова не указаны, запускаем подробный анализ сайта
        if not keywords_text:
            run_detailed_site_analysis_ui(domain, search_engines)
            return
        
        # Показываем кнопку остановки и скрываем кнопку запуска
        serp_stop_btn.visible = True
        serp_run_btn.visible = False
        page.update()
        
        # Создаем событие для остановки
        page.data['serp_stop_event'] = threading.Event()
        
        # Парсим ключевые слова
        keywords_list = [kw.strip() for kw in keywords_text.split('\n') if kw.strip()]
        
        serp_progress_bar.value = 0.0
        serp_results_area.value = f"🔄 Запуск трекинга позиций для {len(keywords_list)} ключевых слов в {', '.join(search_engines)}..."
        serp_export_btn.visible = False
        page.update()
        
        def update_serp_progress(current, total, message):
            if total > 0:
                serp_progress_bar.value = current / total
            serp_results_area.value = message
            page.update()
        
        # Запуск трекинга в отдельном потоке
        threading.Thread(
            target=lambda: run_serp_tracking_worker(keywords_list, domain, search_engines, update_serp_progress, serp_results_area, serp_export_btn, page, serp_stop_btn, serp_run_btn),
            daemon=True
        ).start()
    
    def run_serp_tracking_worker(keywords_list, domain, search_engines, update_callback, results_area, export_btn, page_ref, stop_btn, run_btn):
        """Рабочая функция для трекинга позиций."""
        try:
            # Проверяем, доступен ли модуль
            if run_serp_tracking is None:
                results_area.value = "❌ Модуль SERP Tracker недоступен"
                return
            
            # Запускаем трекинг
            result = run_serp_tracking(keywords_list, domain, search_engines, update_callback)
            
            if "error" in result:
                results_area.value = f"❌ Ошибка: {result['error']}"
            else:
                # Формируем отчет
                report = "📊 РЕЗУЛЬТАТЫ ТРЕКИНГА ПОЗИЦИЙ\n"
                report += "=" * 50 + "\n\n"
                
                for res in result['results']:
                    keyword = res['keyword']
                    engine = res['search_engine']
                    position = res['position']
                    url = res['url']
                    title = res['title']
                    status = res['status']
                    
                    if status == "success":
                        if position > 0:
                            report += f"✅ '{keyword}' в {engine}: позиция {position}\n"
                            if url:
                                report += f"   URL: {url}\n"
                            if title:
                                report += f"   Заголовок: {title}\n"
                        else:
                            report += f"❌ '{keyword}' в {engine}: не найдено в топ-10\n"
                    else:
                        report += f"❌ '{keyword}' в {engine}: ошибка - {status}\n"
                    report += "\n"
                
                # Добавляем статистику
                stats = result['statistics']
                report += "📈 СТАТИСТИКА:\n"
                report += f"Всего ключевых слов: {stats['total_keywords']}\n"
                report += f"В топ-3: {stats['top_3']}\n"
                report += f"В топ-10: {stats['top_10']}\n"
                report += f"Не найдено: {stats['not_found']}\n"
                
                results_area.value = report
                export_btn.visible = True
                page_ref.data['serp_results'] = result
                
        except Exception as e:
            results_area.value = f"❌ Ошибка трекинга: {str(e)}"
        finally:
            # Скрываем кнопку остановки и показываем кнопку запуска
            stop_btn.visible = False
            run_btn.visible = True
            page_ref.update()
    
    def run_detailed_site_analysis_ui(domain, search_engines):
        """Запускает подробный анализ сайта с автоматическим поиском ключевых слов."""
        # Показываем кнопку остановки и скрываем кнопку запуска
        serp_stop_btn.visible = True
        serp_run_btn.visible = False
        page.update()
        
        # Создаем событие для остановки
        page.data['serp_stop_event'] = threading.Event()
        
        serp_progress_bar.value = 0.0
        serp_results_area.value = f"🔄 Запуск подробного анализа сайта {domain}...\n\n🔍 Автоматический поиск ключевых слов...\n📊 Проверка позиций в {', '.join(search_engines)}..."
        serp_export_btn.visible = False
        page.update()
        
        def update_site_progress(current, total, message):
            if total > 0:
                serp_progress_bar.value = current / total
            serp_results_area.value = message
            page.update()
        
        # Запуск анализа в отдельном потоке
        threading.Thread(
            target=lambda: run_detailed_site_analysis_worker(domain, search_engines, update_site_progress, serp_results_area, serp_export_btn, page, serp_stop_btn, serp_run_btn),
            daemon=True
        ).start()
    
    def run_detailed_site_analysis_worker(domain, search_engines, update_callback, results_area, export_btn, page_ref, stop_btn, run_btn):
        """Рабочая функция для подробного анализа сайта."""
        try:
            # Проверяем, доступен ли модуль
            if run_detailed_site_analysis is None:
                results_area.value = "❌ Модуль SERP Tracker недоступен"
                return
            
            # Запускаем подробный анализ
            import serp_tracker
            result = serp_tracker.run_detailed_site_analysis(domain, search_engines, update_callback)
            
            if "error" in result:
                results_area.value = f"❌ Ошибка: {result['error']}"
            else:
                # Формируем подробный отчет
                report = "🌐 ПОДРОБНЫЙ АНАЛИЗ САЙТА\n"
                report += "=" * 60 + "\n\n"
                
                report += f"📋 ОСНОВНАЯ ИНФОРМАЦИЯ:\n"
                report += f"Домен: {result['domain']}\n"
                report += f"Проверено ключевых слов: {len(result['keywords_checked'])}\n"
                report += f"Поисковые системы: {', '.join(search_engines)}\n\n"
                
                # Показываем найденные ключевые слова
                report += f"🔍 ПРОВЕРЕННЫЕ КЛЮЧЕВЫЕ СЛОВА:\n"
                for i, keyword in enumerate(result['keywords_checked'][:10], 1):
                    report += f"{i}. {keyword}\n"
                if len(result['keywords_checked']) > 10:
                    report += f"... и еще {len(result['keywords_checked']) - 10} ключевых слов\n"
                report += "\n"
                
                # Показываем результаты позиций
                report += f"📊 РЕЗУЛЬТАТЫ ПОЗИЦИЙ:\n"
                report += "-" * 40 + "\n"
                
                found_positions = []
                not_found = []
                
                for res in result['detailed_results']:
                    keyword = res['keyword']
                    engine = res['search_engine']
                    position = res['position']
                    url = res['url']
                    title = res['title']
                    snippet = res.get('snippet', '')
                    
                    if position > 0:
                        found_positions.append(res)
                        report += f"✅ '{keyword}' в {engine}: позиция {position}\n"
                        if url:
                            report += f"   URL: {url}\n"
                        if title:
                            report += f"   Заголовок: {title}\n"
                        if snippet:
                            report += f"   Сниппет: {snippet[:100]}...\n"
                        report += "\n"
                    else:
                        not_found.append(res)
                        report += f"❌ '{keyword}' в {engine}: не найдено в топ-10\n\n"
                
                # Показываем детальную информацию о топ-результатах
                if found_positions:
                    report += f"🏆 ДЕТАЛЬНАЯ ИНФОРМАЦИЯ О НАЙДЕННЫХ ПОЗИЦИЯХ:\n"
                    report += "-" * 50 + "\n"
                    
                    for res in found_positions[:5]:  # Показываем топ-5
                        keyword = res['keyword']
                        engine = res['search_engine']
                        position = res['position']
                        search_results = res.get('search_results', {})
                        
                        report += f"🔍 '{keyword}' в {engine} (позиция {position}):\n"
                        
                        # Показываем топ-5 результатов поиска
                        all_results = search_results.get('all_results', [])
                        for i, result in enumerate(all_results[:5], 1):
                            result_url = result.get('url', '')
                            result_title = result.get('title', '')
                            is_ours = domain in result_url
                            
                            if is_ours:
                                report += f"   {i}. 🎯 {result_title}\n"
                                report += f"      {result_url}\n"
                            else:
                                report += f"   {i}. {result_title}\n"
                                report += f"      {result_url}\n"
                        report += "\n"
                
                # Добавляем статистику
                stats = result['statistics']
                report += f"📈 СТАТИСТИКА:\n"
                report += f"Всего ключевых слов: {stats['total_keywords']}\n"
                report += f"В топ-3: {stats['top_3']}\n"
                report += f"В топ-10: {stats['top_10']}\n"
                report += f"Не найдено: {stats['not_found']}\n\n"
                
                # Информация о графиках
                if result.get('charts'):
                    report += f"📊 ГРАФИКИ ДВИЖЕНИЯ ПОЗИЦИЙ:\n"
                    report += f"Сгенерировано графиков: {len(result['charts'])}\n"
                    report += f"Данные сохранены в базе для отслеживания изменений\n\n"
                
                report += f"✅ Подробный анализ завершен!"
                
                results_area.value = report
                export_btn.visible = True
                page_ref.data['detailed_site_analysis'] = result
                
        except Exception as e:
            results_area.value = f"❌ Ошибка подробного анализа: {str(e)}"
        finally:
            # Скрываем кнопку остановки и показываем кнопку запуска
            stop_btn.visible = False
            run_btn.visible = True
            page_ref.update()
    
    def stop_serp_tracking(e):
        """Останавливает трекинг позиций."""
        if 'serp_stop_event' in page.data:
            page.data['serp_stop_event'].set()
        
        # Скрываем кнопку остановки и показываем кнопку запуска
        serp_stop_btn.visible = False
        serp_run_btn.visible = True
        page.update()
        
        # Показываем уведомление
        page.snack_bar = ft.SnackBar(content=ft.Text("Трекинг позиций остановлен"))
        page.snack_bar.open = True
        page.update()
    
    def export_serp_results(e):
        """Экспортирует результаты трекинга или анализа в Excel."""
        # Проверяем, есть ли данные для экспорта
        if 'serp_results' not in page.data and 'site_analysis' not in page.data and 'detailed_site_analysis' not in page.data:
            page.snack_bar = ft.SnackBar(content=ft.Text("❌ Нет данных для экспорта"))
            page.snack_bar.open = True
            page.update()
            return
        
        try:
            # Экспортируем результаты подробного анализа сайта
            if 'detailed_site_analysis' in page.data:
                import pandas as pd
                from datetime import datetime
                
                analysis_data = page.data['detailed_site_analysis']
                
                # Генерируем имя файла
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"reports/detailed_analysis_{analysis_data['domain'].replace('.', '_')}_{timestamp}.xlsx"
                
                with pd.ExcelWriter(filename, engine='openpyxl') as writer:
                    # Лист с основной информацией
                    basic_data = {
                        'Параметр': [
                            'Домен',
                            'Проверено ключевых слов',
                            'Поисковые системы',
                            'Всего позиций',
                            'В топ-3',
                            'В топ-10',
                            'Не найдено'
                        ],
                        'Значение': [
                            analysis_data['domain'],
                            len(analysis_data['keywords_checked']),
                            ', '.join(set([res['search_engine'] for res in analysis_data['detailed_results']])),
                            analysis_data['statistics']['total_keywords'],
                            analysis_data['statistics']['top_3'],
                            analysis_data['statistics']['top_10'],
                            analysis_data['statistics']['not_found']
                        ]
                    }
                    basic_df = pd.DataFrame(basic_data)
                    basic_df.to_excel(writer, sheet_name='Основная информация', index=False)
                    
                    # Лист с результатами позиций
                    positions_data = []
                    for res in analysis_data['detailed_results']:
                        positions_data.append({
                            'Ключевое слово': res['keyword'],
                            'Поисковая система': res['search_engine'],
                            'Позиция': res['position'],
                            'URL': res['url'],
                            'Заголовок': res['title'],
                            'Сниппет': res.get('snippet', ''),
                            'Статус': res['status']
                        })
                    
                    if positions_data:
                        positions_df = pd.DataFrame(positions_data)
                        positions_df.to_excel(writer, sheet_name='Результаты позиций', index=False)
                    
                    # Лист с ключевыми словами
                    keywords_data = {
                        'Ключевое слово': analysis_data['keywords_checked']
                    }
                    keywords_df = pd.DataFrame(keywords_data)
                    keywords_df.to_excel(writer, sheet_name='Ключевые слова', index=False)
                
                page.snack_bar = ft.SnackBar(content=ft.Text(f"✅ Подробный анализ экспортирован в {filename}"))
                page.snack_bar.open = True
                page.update()
                return
            
            # Экспортируем результаты анализа сайта
            if 'site_analysis' in page.data:
                import pandas as pd
                from datetime import datetime
                
                analysis_data = page.data['site_analysis']
                
                # Создаем DataFrame для экспорта
                data = {
                    'Параметр': [
                        'Домен',
                        'Протокол', 
                        'Главная страница',
                        'HTTP статус',
                        'robots.txt',
                        'sitemap.xml'
                    ],
                    'Значение': [
                        analysis_data['domain_info']['domain'],
                        analysis_data['domain_info']['protocol'],
                        analysis_data['domain_info']['path'],
                        str(analysis_data['status_code']),
                        analysis_data['robots_info'],
                        analysis_data['sitemap_info']
                    ]
                }
                
                df = pd.DataFrame(data)
                
                # Генерируем имя файла
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"reports/site_analysis_{analysis_data['domain'].replace('.', '_')}_{timestamp}.xlsx"
                
                # Экспортируем в Excel
                df.to_excel(filename, index=False, sheet_name='Анализ сайта')
                
                page.snack_bar = ft.SnackBar(content=ft.Text(f"✅ Результаты анализа экспортированы в {filename}"))
                page.snack_bar.open = True
                page.update()
                return
            
            # Экспортируем результаты трекинга позиций
            if 'serp_results' in page.data:
                if SERPTracker is None:
                    page.snack_bar = ft.SnackBar(content=ft.Text("❌ Модуль SERP Tracker недоступен"))
                    page.snack_bar.open = True
                    page.update()
                    return
                
                tracker = SERPTracker()
                filename = tracker.export_to_excel()
                
                page.snack_bar = ft.SnackBar(content=ft.Text(f"✅ Результаты трекинга экспортированы в {filename}"))
                page.snack_bar.open = True
                page.update()
            
        except Exception as e:
            page.snack_bar = ft.SnackBar(content=ft.Text(f"❌ Ошибка экспорта: {str(e)}"))
            page.snack_bar.open = True
            page.update()
    
    serp_run_btn.on_click = run_serp_tracking
    serp_stop_btn.on_click = stop_serp_tracking
    serp_export_btn.on_click = export_serp_results

    # --- Парсер всех страниц ---
    parser_url_input = ft.TextField(
        label="Главная страница сайта", 
        width=400, 
        filled=True, 
        border_radius=10,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color=get_input_text_color(),
        label_style=ft.TextStyle(color=get_label_color())
    )
    parser_ssl_checkbox = ft.Checkbox(label="Игнорировать SSL", value=True)
    parser_max_pages = ft.TextField(
        label="Макс. страниц", 
        value="1000", 
        width=100, 
        filled=True, 
        border_radius=10,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color=get_input_text_color(),
        label_style=ft.TextStyle(color=get_label_color())
    )
    parser_progress = ft.ProgressBar(width=400, color="#F2E307", bgcolor="#394459", value=0.0, height=10, border_radius=20)
    parser_table = ft.DataTable(
        columns=[
            ft.DataColumn(label=ft.Text("Ссылка", color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("HTTP", color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("Редирект", color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("SEO Информация", color=get_input_text_color())),
        ],
        rows=[],
        horizontal_lines=ft.BorderSide(1, "#394459"),
        vertical_lines=ft.BorderSide(1, "#394459"),
        bgcolor="#F2F2F2",
        border=ft.border.all(1, "#394459"),
        border_radius=10,
    )
    parser_export_btn = ft.ElevatedButton(
        "Экспорт в Excel", 
        icon=ft.Icons.DOWNLOAD, 
        visible=False,
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    parser_export_word_btn = ft.ElevatedButton(
        "Экспорт в Word", 
        icon=ft.Icons.DESCRIPTION, 
        visible=False,
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    parser_status = ft.Text(visible=False)
    
    # Переменные для управления парсером
    parser_stop_event = threading.Event()
    parser_thread = None

    def stop_parser(e):
        """Останавливает парсер."""
        if parser_thread and parser_thread.is_alive():
            parser_stop_event.set()
            parser_status.value = "Останавливаем парсер..."
            parser_status.visible = True
            parser_run_btn.visible = True
            parser_stop_btn.visible = False
            
            # Показываем уведомление
            page.snack_bar = ft.SnackBar(content=ft.Text("Парсер остановлен"))
            page.snack_bar.open = True
            page.update()

    def parser_update(visited_count, found_count):
        if parser_stop_event.is_set():
            parser_status.value = f"Останавливаем... Обработано страниц: {visited_count}"
        else:
            parser_progress.value = min(1.0, found_count / max(visited_count, 1))
        page.update()



    def parser_done(results):
        # Проверяем, была ли остановка
        if parser_stop_event.is_set():
            parser_status.value = f"Парсер остановлен. Обработано страниц: {len(results)}"
        else:
            parser_status.value = f"Парсинг завершен. Найдено страниц: {len(results)}"
        
        parser_table.rows = []
        for r in results:
            # Формируем читаемую SEO информацию
            title = r['Title'] if r['Title'] else 'Отсутствует'
            description = r['Meta_Description'] if r['Meta_Description'] else 'Отсутствует'
            h1 = r['H1'] if r['H1'] else 'Отсутствует'
            
            # Создаем многострочный текст с правильным форматированием
            seo_text = ft.Text(
                f"📄 Title: {title}\n"
                f"📝 Description: {description}\n"
                f"🔤 H1: {h1}\n"
                f"✅ Статус: {r['SEO']}",
                size=10,
                selectable=True,
                max_lines=8,  # Ограничиваем количество строк для компактности
                overflow=ft.TextOverflow.ELLIPSIS,  # Добавляем многоточие для длинного текста
                weight=ft.FontWeight.NORMAL,
                color=get_input_text_color()
            )
            
            row = ft.DataRow(cells=[
                ft.DataCell(ft.Text(r['Ссылка'], color=get_input_text_color())),
                ft.DataCell(ft.Text(str(r['HTTP']), color=get_input_text_color())),
                ft.DataCell(ft.Text(r['Редирект'], color=get_input_text_color())),
                ft.DataCell(seo_text)
            ])
            parser_table.rows.append(row)
        # Восстанавливаем состояние кнопок
        parser_run_btn.visible = True
        parser_stop_btn.visible = False
        
        parser_export_btn.visible = True if results else False
        parser_export_word_btn.visible = True if results else False
        parser_status.visible = True
        
        # Сохраняем данные для экспорта
        page.data['parser_results'] = results
        page.data['parser_site_url'] = parser_url_input.value.strip()
        
        # Сохраняем в Excel
        import pandas as pd
        df = pd.DataFrame(results)
        fname = f"reports/allpages_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        df.to_excel(fname, index=False)
        
        # Показываем уведомление о сохранении
        page.snack_bar = ft.SnackBar(content=ft.Text(f"✅ Excel отчет сохранен: {os.path.basename(fname)}"))
        page.snack_bar.open = True
        
        # Обновляем список экспортов если открыта страница экспорта
        if exports_content.visible:
            refresh_exports_list()
        
        page.update()

    def export_to_word(e):
        """Экспортирует результаты парсера в Word."""
        if 'parser_results' in page.data and 'parser_site_url' in page.data:
            results = page.data['parser_results']
            site_url = page.data['parser_site_url']
            
            try:
                report_path = generate_word_report(results, site_url, 'parser')
                if report_path:
                    page.snack_bar = ft.SnackBar(content=ft.Text(f"✅ Отчет сохранен: {os.path.basename(report_path)}"))
                    page.snack_bar.open = True
                    # Обновляем список экспортов если открыта страница экспорта
                    if exports_content.visible:
                        refresh_exports_list()
                else:
                    page.snack_bar = ft.SnackBar(content=ft.Text("❌ Ошибка: модуль python-docx не установлен"))
                    page.snack_bar.open = True
            except Exception as ex:
                page.snack_bar = ft.SnackBar(content=ft.Text(f"❌ Ошибка экспорта: {str(ex)}"))
                page.snack_bar.open = True
            page.update()
        else:
            page.snack_bar = ft.SnackBar(content=ft.Text("❌ Нет данных для экспорта"))
            page.snack_bar.open = True
            page.update()

    def parser_run(e):
        url = parser_url_input.value.strip()
        if not url.startswith('http'):
            parser_status.value = "Введите корректный URL!"
            parser_status.visible = True
            page.update()
            return
        
        # Сбрасываем событие остановки
        parser_stop_event.clear()
        
        parser_progress.value = 0.0
        parser_table.rows = []
        parser_export_btn.visible = False
        parser_status.value = "Начинаем парсинг сайта..."
        parser_status.visible = True
        parser_run_btn.visible = False
        parser_stop_btn.visible = True
        page.update()
        
        try:
            max_pages = int(parser_max_pages.value)
        except ValueError:
            max_pages = 15000
        
        def parser_worker():
            crawl_site_without_sitemap(url, parser_ssl_checkbox.value, parser_update, parser_done, parser_stop_event, max_pages=max_pages)
        
        parser_thread = threading.Thread(target=parser_worker)
        parser_thread.start()

    parser_run_btn = ft.ElevatedButton(
        "Запустить парсер", 
        icon=ft.Icons.PLAY_ARROW, 
        on_click=parser_run,
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=10),
            elevation=5
        )
    )
    parser_stop_btn = ft.ElevatedButton(
        "Остановить", 
        icon=ft.Icons.STOP, 
        visible=False, 
        on_click=lambda e: stop_parser(e),
        style=ft.ButtonStyle(
            bgcolor="#FF6B6B",
            color="#FFFFFF",
            shape=ft.RoundedRectangleBorder(radius=10),
            elevation=5
        )
    )
    parser_export_btn.on_click = lambda e: None  # Уже экспортируется автоматически
    parser_export_word_btn.on_click = export_to_word
    parser_content.content = ft.Column([
        ft.Text("Парсер всех страниц сайта", size=20, weight=ft.FontWeight.BOLD, color=get_text_color()),
        ft.Text("Введите главную страницу сайта. Парсер найдет все внутренние ссылки и покажет SEO информацию в таблице.", size=14, color=get_secondary_text_color()),
        ft.Row([parser_url_input, parser_ssl_checkbox, parser_max_pages, parser_run_btn, parser_stop_btn]),
        parser_progress,
        ft.Container(
            ft.Row([
                ft.Container(parser_table, width=2500, expand=True)
            ], scroll=ft.ScrollMode.ALWAYS),
            expand=True,
            bgcolor="#F2F2F2",
            border_radius=10,
            padding=5,
            alignment=ft.alignment.center,
            border=ft.border.all(1, "#394459"),
        ),
        ft.Row([parser_export_btn, parser_export_word_btn]),
        parser_status
    ], expand=True)
    
    # --- Анализ текста ---
    text_analysis_url_input = ft.TextField(
        label="URL страницы для анализа", 
        width=400, 
        filled=True, 
        border_radius=10,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color=get_input_text_color(),
        label_style=ft.TextStyle(color=get_label_color())
    )
    text_analysis_keywords_input = ft.TextField(
        label="Ключевые слова (через запятую)", 
        width=300, 
        filled=True, 
        border_radius=10,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color=get_input_text_color(),
        label_style=ft.TextStyle(color=get_label_color())
    )
    text_analysis_ssl_checkbox = ft.Checkbox(label="Игнорировать SSL", value=True)
    text_analysis_progress = ft.ProgressBar(width=400, color="#F2E307", bgcolor="#394459", value=0.0, height=10, border_radius=20)
    text_analysis_status = ft.Text(visible=False)
    
    # Области для вывода результатов
    text_analysis_summary = ft.TextField(
        label="SEO Анализ текста", 
        multiline=True, 
        min_lines=20, 
        max_lines=40, 
        width=1000, 
        filled=True, 
        border_radius=10, 
        expand=True,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color=get_input_text_color(),
        label_style=ft.TextStyle(color=get_label_color())
    )
    text_analysis_keywords = ft.TextField(
        label="Ключевые слова и плотность", 
        multiline=True, 
        min_lines=25, 
        max_lines=50, 
        width=1000, 
        filled=True, 
        border_radius=10, 
        expand=True,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color=get_input_text_color(),
        label_style=ft.TextStyle(color=get_label_color())
    )
    text_analysis_structure = ft.TextField(
        label="Структура заголовков", 
        multiline=True, 
        min_lines=8, 
        max_lines=15, 
        width=1000, 
        filled=True, 
        border_radius=10, 
        expand=True,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color=get_input_text_color(),
        label_style=ft.TextStyle(color=get_label_color())
    )
    text_analysis_full_text = ft.TextField(
        label="Полный текст страницы", 
        multiline=True, 
        min_lines=15, 
        max_lines=50, 
        width=1000, 
        filled=True, 
        border_radius=10, 
        expand=True,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color=get_input_text_color(),
        label_style=ft.TextStyle(color=get_label_color())
    )
    text_analysis_declensions = ft.TextField(
        label="Анализ с учетом склонений", 
        multiline=True, 
        min_lines=25, 
        max_lines=50, 
        width=1000, 
        filled=True, 
        border_radius=10, 
        expand=True,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color=get_input_text_color(),
        label_style=ft.TextStyle(color=get_label_color())
    )
    
    def text_analysis_run(e):
        url = text_analysis_url_input.value.strip()
        if not url.startswith('http'):
            text_analysis_status.value = "Введите корректный URL!"
            text_analysis_status.visible = True
            page.update()
            return
        
        text_analysis_progress.value = 0.0
        text_analysis_status.value = "Анализируем текст (улучшенный анализ с Selenium)..."
        text_analysis_status.visible = True
        text_analysis_summary.value = ""
        text_analysis_keywords.value = ""
        text_analysis_structure.value = ""
        text_analysis_full_text.value = ""
        text_analysis_declensions.value = ""
        page.update()
        
        def worker():
            try:
                # Получаем HTML страницы через Selenium (как в функции анализа склонений)
                from selenium import webdriver
                from selenium.webdriver.chrome.options import Options
                
                chrome_options = Options()
                chrome_options.add_argument('--headless')
                chrome_options.add_argument('--no-sandbox')
                chrome_options.add_argument('--disable-dev-shm-usage')
                
                driver = webdriver.Chrome(options=chrome_options)
                driver.get(url)
                
                # Получаем HTML после загрузки страницы
                html_content = driver.page_source
                driver.quit()
                
                # Анализируем текст
                analysis = analyze_text_content(html_content, url)
                
                # Формируем сводку
                summary_parts = []
                summary_parts.append("📊 SEO АНАЛИЗ ТЕКСТА (УЛУЧШЕННЫЙ)")
                summary_parts.append("=" * 50)
                summary_parts.append("🔍 Анализ включает основной контент на странице:")
                summary_parts.append("   • Основной текст")
                summary_parts.append("   • Текст в анкорах (ссылках)")
                summary_parts.append("   • Любой другой видимый текст")
                summary_parts.append("   • ИСКЛЮЧАЕТ: header, footer, атрибуты (alt, title, placeholder, aria-label)")
                summary_parts.append("")
                summary_parts.append(f"URL: {analysis['url']}")
                summary_parts.append(f"Title: {analysis['title']}")
                summary_parts.append(f"Meta Description: {analysis['meta_description']}")
                summary_parts.append("")
                
                summary_parts.append("📈 СТАТИСТИКА:")
                summary_parts.append(f"• Всего слов: {analysis['total_words']}")
                summary_parts.append(f"• H1 заголовков: {analysis['h1_count']}")
                summary_parts.append(f"• H2 заголовков: {analysis['h2_count']}")
                summary_parts.append(f"• H3 заголовков: {analysis['h3_count']}")
                summary_parts.append(f"• Средняя длина предложения: {analysis['avg_sentence_length']} слов")
                summary_parts.append(f"• Средняя длина абзаца: {analysis['avg_paragraph_length']} слов")
                summary_parts.append(f"• Изображений: {analysis['images_total']} (с alt: {analysis['images_with_alt']}, без alt: {analysis['images_without_alt']})")
                summary_parts.append(f"• Внутренних ссылок: {analysis['internal_links']}")
                summary_parts.append(f"• Внешних ссылок: {analysis['external_links']}")
                summary_parts.append(f"• Оценка структуры: {analysis['structure_score']}/100")
                summary_parts.append("")
                
                if analysis['positives']:
                    summary_parts.append("✅ ПОЛОЖИТЕЛЬНЫЕ МОМЕНТЫ:")
                    for positive in analysis['positives']:
                        summary_parts.append(f"• {positive}")
                    summary_parts.append("")
                
                if analysis['recommendations']:
                    summary_parts.append("⚠️ РЕКОМЕНДАЦИИ:")
                    for rec in analysis['recommendations']:
                        summary_parts.append(f"• {rec}")
                    summary_parts.append("")
                
                summary_parts.append("📝 ПРЕДВАРИТЕЛЬНЫЙ ПРОСМОТР ТЕКСТА:")
                summary_parts.append(analysis['text_preview'])
                
                # Формируем ключевые слова с подробной информацией о плотности
                keywords_text = "🔑 ТОП-20 КЛЮЧЕВЫХ СЛОВ И ПЛОТНОСТЬ (УЛУЧШЕННЫЙ АНАЛИЗ):\n"
                keywords_text += "=" * 60 + "\n"
                keywords_text += "📊 Анализ включает основной контент на странице (включая анкоры, исключая header/footer и атрибуты)\n"
                keywords_text += f"📈 Всего слов в тексте: {analysis['total_words']}\n"
                keywords_text += "=" * 60 + "\n\n"
                
                # Сортируем ключевые слова по плотности
                sorted_keywords = sorted(analysis['keyword_density'].items(), key=lambda x: x[1], reverse=True)
                
                keywords_text += "📊 ПО ПЛОТНОСТИ (от высокой к низкой):\n"
                keywords_text += "-" * 40 + "\n"
                for i, (word, density) in enumerate(sorted_keywords[:20], 1):
                    count = next((count for w, count in analysis['top_keywords'] if w == word), 0)
                    # Безопасное форматирование плотности
                    if isinstance(density, (int, float)):
                        # Добавляем цветовую индикацию плотности
                        if density > 3.0:
                            indicator = "🔴"  # Очень высокая плотность
                        elif density > 2.0:
                            indicator = "🟡"  # Высокая плотность
                        elif density > 1.0:
                            indicator = "🟢"  # Нормальная плотность
                        else:
                            indicator = "⚪"  # Низкая плотность
                        
                        keywords_text += f"{i:2d}. {indicator} {word}\n"
                        keywords_text += f"    Частота: {count} раз | Плотность: {density:.2f}%\n\n"
                    else:
                        keywords_text += f"{i:2d}. ⚪ {word}\n"
                        keywords_text += f"    Частота: {count} раз | Плотность: {density}\n\n"
                
                keywords_text += "\n📈 ПО ЧАСТОТЕ (от частой к редкой):\n"
                keywords_text += "-" * 40 + "\n"
                for i, (word, count) in enumerate(analysis['top_keywords'], 1):
                    density = analysis['keyword_density'].get(word, 0)
                    # Безопасное форматирование плотности
                    if isinstance(density, (int, float)):
                        # Добавляем цветовую индикацию плотности
                        if density > 3.0:
                            indicator = "🔴"  # Очень высокая плотность
                        elif density > 2.0:
                            indicator = "🟡"  # Высокая плотность
                        elif density > 1.0:
                            indicator = "🟢"  # Нормальная плотность
                        else:
                            indicator = "⚪"  # Низкая плотность
                        
                        keywords_text += f"{i:2d}. {indicator} {word}\n"
                        keywords_text += f"    Частота: {count} раз | Плотность: {density:.2f}%\n\n"
                    else:
                        keywords_text += f"{i:2d}. ⚪ {word}\n"
                        keywords_text += f"    Частота: {count} раз | Плотность: {density}\n\n"
                
                # Добавляем рекомендации по плотности
                keywords_text += "\n💡 РЕКОМЕНДАЦИИ ПО ПЛОТНОСТИ:\n"
                keywords_text += "-" * 40 + "\n"
                high_density_words = [(word, density) for word, density in sorted_keywords if density > 3.0]
                low_density_words = [(word, density) for word, density in sorted_keywords if density < 0.5]
                
                if high_density_words:
                    keywords_text += "🔴 Слишком высокая плотность (>3%):\n"
                    for word, density in high_density_words[:5]:
                        if isinstance(density, (int, float)):
                            keywords_text += f"   • {word}: {density:.2f}% - снизьте использование\n"
                        else:
                            keywords_text += f"   • {word}: {density} - снизьте использование\n"
                    keywords_text += "\n"
                
                if low_density_words:
                    keywords_text += "⚪ Слишком низкая плотность (<0.5%):\n"
                    for word, density in low_density_words[:5]:
                        if isinstance(density, (int, float)):
                            keywords_text += f"   • {word}: {density:.2f}% - увеличьте использование\n"
                        else:
                            keywords_text += f"   • {word}: {density} - увеличьте использование\n"
                    keywords_text += "\n"
                
                keywords_text += "✅ Оптимальная плотность: 1-2% от общего количества слов\n"
                keywords_text += "⚠️ Избегайте переспама ключевыми словами\n"
                keywords_text += "\n🔧 ТЕХНИЧЕСКАЯ ИНФОРМАЦИЯ:\n"
                keywords_text += "• Анализ выполнен с использованием Selenium WebDriver\n"
                keywords_text += "• Получен основной контент со страницы\n"
                keywords_text += "• Исключены: header, footer, атрибуты (alt, title, placeholder, aria-label)\n"
                
                # Формируем структуру заголовков
                structure_text = "📋 СТРУКТУРА ЗАГОЛОВКОВ:\n"
                structure_text += "=" * 30 + "\n"
                
                if analysis['h1_texts']:
                    structure_text += "H1:\n"
                    for h1 in analysis['h1_texts']:
                        structure_text += f"• {h1}\n"
                    structure_text += "\n"
                
                if analysis['h2_texts']:
                    structure_text += "H2:\n"
                    for h2 in analysis['h2_texts']:
                        structure_text += f"• {h2}\n"
                    structure_text += "\n"
                
                if analysis['h3_texts']:
                    structure_text += "H3:\n"
                    for h3 in analysis['h3_texts']:
                        structure_text += f"• {h3}\n"
                
                # Формируем полный текст
                full_text = "📄 ПОЛНЫЙ ТЕКСТ СТРАНИЦЫ (УЛУЧШЕННЫЙ АНАЛИЗ):\n"
                full_text += "=" * 60 + "\n"
                full_text += "🔍 Включает основной контент: текст, анкоры (исключает header/footer и атрибуты)\n"
                full_text += f"📏 Длина текста: {len(analysis['full_text'])} символов\n"
                full_text += f"📊 Количество слов: {analysis['total_words']}\n"
                full_text += "=" * 60 + "\n\n"
                full_text += analysis['full_text']
                
                # Обновляем интерфейс
                text_analysis_summary.value = "\n".join(summary_parts)
                text_analysis_keywords.value = keywords_text
                text_analysis_structure.value = structure_text
                text_analysis_full_text.value = full_text
                text_analysis_status.value = "Улучшенный анализ завершен! (основной контент, исключены header/footer и атрибуты)"
                text_analysis_progress.value = 1.0
                
            except Exception as ex:
                text_analysis_status.value = f"Ошибка анализа: {str(ex)}"
                text_analysis_progress.value = 0.0
            
            page.update()
        
        threading.Thread(target=worker).start()
    
    text_analysis_run_btn = ft.ElevatedButton(
        "Анализировать текст (УЛУЧШЕННЫЙ)", 
        icon=ft.Icons.PLAY_ARROW, 
        on_click=text_analysis_run,
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=10),
            elevation=5
        )
    )
    
    def text_analysis_declensions_run(e):
        url = text_analysis_url_input.value.strip()
        keywords = text_analysis_keywords_input.value.strip()
        
        if not url.startswith('http'):
            text_analysis_status.value = "Введите корректный URL!"
            text_analysis_status.visible = True
            page.update()
            return
        
        if not keywords:
            text_analysis_status.value = "Введите ключевые слова для анализа!"
            text_analysis_status.visible = True
            page.update()
            return
        
        text_analysis_progress.value = 0.0
        text_analysis_status.value = "Анализируем текст с учетом склонений..."
        text_analysis_status.visible = True
        text_analysis_declensions.value = ""
        page.update()
        
        def worker():
            try:
                # Получаем HTML страницы
                response = requests.get(url, timeout=15, verify=not text_analysis_ssl_checkbox.value)
                if response.status_code != 200:
                    text_analysis_status.value = f"Ошибка загрузки страницы: {response.status_code}"
                    page.update()
                    return
                
                # Создаем драйвер для анализа ключевых слов
                from selenium import webdriver
                from selenium.webdriver.chrome.options import Options
                
                chrome_options = Options()
                chrome_options.add_argument('--headless')
                chrome_options.add_argument('--no-sandbox')
                chrome_options.add_argument('--disable-dev-shm-usage')
                
                driver = webdriver.Chrome(options=chrome_options)
                driver.get(url)
                
                # Анализируем ключевые слова с учетом склонений
                keywords_result, density, target_analysis = analyze_keywords(driver, url, keywords)
                
                driver.quit()
                
                # Формируем сводку анализа склонений
                declensions_text = []
                declensions_text.append("🔍 АНАЛИЗ ЦЕЛЕВЫХ КЛЮЧЕВЫХ СЛОВ С УЧЕТОМ СКЛОНЕНИЙ")
                declensions_text.append("=" * 60)
                declensions_text.append(f"URL: {url}")
                declensions_text.append(f"Ключевые слова: {keywords}")
                declensions_text.append("")
                
                if isinstance(target_analysis, dict) and target_analysis:
                    for tkw, data in target_analysis.items():
                        declensions_text.append(f"🎯 ЦЕЛЕВОЕ КЛЮЧЕВОЕ СЛОВО: '{tkw}'")
                        declensions_text.append(f"📊 ОБЩАЯ ЧАСТОТА (со склонениями): {data['freq']} раз")
                        # Безопасное форматирование плотности
                        density_value = data['density']
                        if isinstance(density_value, (int, float)):
                            declensions_text.append(f"📈 ПЛОТНОСТЬ: {density_value:.2%}")
                        else:
                            declensions_text.append(f"📈 ПЛОТНОСТЬ: {density_value}")
                        declensions_text.append("")
                        
                        # Показываем найденные склонения
                        if 'declensions_found' in data and data['declensions_found']:
                            declensions_text.append("📝 НАЙДЕННЫЕ СКЛОНЕНИЯ:")
                            for declension, count in data['declensions_found'].items():
                                declensions_text.append(f"  ✅ '{declension}': {count} раз")
                            declensions_text.append("")
                            
                            # Показываем полный текст с подсветкой
                            declensions_text.append("📄 ПОЛНЫЙ ТЕКСТ С ПОДСВЕТКОЙ:")
                            declensions_text.append("-" * 40)
                            
                            # Получаем текст страницы
                            soup = BeautifulSoup(response.text, 'html.parser')
                            text = soup.get_text(separator=' ', strip=True)
                            
                            # Подсвечиваем найденные склонения
                            highlighted_text = text
                            for declension, count in data['declensions_found'].items():
                                if count > 0:
                                    # Заменяем найденные склонения на подсвеченные версии
                                    pattern = r'\b' + re.escape(declension) + r'\b'
                                    highlighted_text = re.sub(pattern, f"【{declension}】", highlighted_text, flags=re.IGNORECASE)
                            
                            # Показываем первые 500 символов с подсветкой
                            preview = highlighted_text[:500] + "..." if len(highlighted_text) > 500 else highlighted_text
                            declensions_text.append(preview)
                            declensions_text.append("")
                            
                            # Добавляем инструкции для проверки через DevTools
                            declensions_text.append("🔍 ПРОВЕРКА ЧЕРЕЗ DEVTOOLS:")
                            declensions_text.append("-" * 40)
                            declensions_text.append("1. Откройте DevTools (F12)")
                            declensions_text.append("2. Нажмите Ctrl+F для поиска")
                            declensions_text.append("3. Проверьте каждое склонение:")
                            for declension, count in data['declensions_found'].items():
                                if count > 0:
                                    declensions_text.append(f"   • '{declension}' - должно быть {count} раз")
                            declensions_text.append("")
                            
                            # Показываем только найденные склонения
                            declensions_text.append("📚 НАЙДЕННЫЕ СКЛОНЕНИЯ:")
                            declensions_text.append("-" * 50)
                            
                            # Показываем только найденные склонения с галочками
                            for declension, count in data['declensions_found'].items():
                                if count > 0:
                                    declensions_text.append(f"  ✅ '{declension}' - найдено {count} раз")
                            
                            declensions_text.append("")
                            
                        else:
                            declensions_text.append("❌ Склонения не найдены")
                            declensions_text.append("")
                        
                        declensions_text.append("=" * 60)
                        declensions_text.append("")
                else:
                    declensions_text.append("❌ Целевые ключевые слова не найдены на странице")
                    declensions_text.append("")
                
                text_analysis_declensions.value = "\n".join(declensions_text)
                text_analysis_progress.value = 1.0
                text_analysis_status.value = "Анализ завершен!"
                page.update()
                
            except Exception as ex:
                text_analysis_status.value = f"Ошибка анализа: {str(ex)}"
                text_analysis_progress.value = 0.0
                page.update()
        
        threading.Thread(target=worker).start()
    
    text_analysis_declensions_btn = ft.ElevatedButton(
        "Анализ склонений", 
        icon=ft.Icons.SEARCH, 
        on_click=text_analysis_declensions_run,
        style=ft.ButtonStyle(
            bgcolor="#F2CC0C",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    
    text_analysis_content.content = ft.Column([
        ft.Text("Анализ текста для SEO (УЛУЧШЕННЫЙ)", size=24, weight=ft.FontWeight.BOLD, color=get_text_color()),
        ft.Text("Введите URL страницы для детального анализа основного текстового содержимого (включая анкоры, исключая header/footer и атрибуты)", size=16, color=get_secondary_text_color()),
        ft.Row([text_analysis_url_input, text_analysis_keywords_input, text_analysis_ssl_checkbox], spacing=10),
        ft.Row([text_analysis_run_btn, text_analysis_declensions_btn], spacing=10),
        text_analysis_progress,
        text_analysis_status,
        ft.Column([
            text_analysis_summary,
            ft.Container(height=20),
            text_analysis_keywords,
            ft.Container(height=20),
            text_analysis_structure,
            ft.Container(height=20),
            text_analysis_full_text
        ], expand=True),
        ft.Container(height=20),
        text_analysis_declensions
    ], expand=True)
    
    # --- Анализ кода ---
    code_analysis_url_input = ft.TextField(
        label="URL страницы для анализа кода", 
        width=400, 
        filled=True, 
        border_radius=10,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color=get_input_text_color(),
        label_style=ft.TextStyle(color=get_label_color())
    )
    code_analysis_ssl_checkbox = ft.Checkbox(label="Игнорировать SSL", value=True)
    code_analysis_progress = ft.ProgressBar(width=400, color="#F2E307", bgcolor="#394459", value=0.0, height=10, border_radius=20)
    code_analysis_status = ft.Text(visible=False)
    
    # Области для вывода результатов
    code_analysis_summary = ft.TextField(
        label="Общий анализ кода", 
        multiline=True, 
        min_lines=15, 
        max_lines=30, 
        width=1000, 
        filled=True, 
        border_radius=10, 
        expand=True,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color=get_input_text_color(),
        label_style=ft.TextStyle(color=get_label_color())
    )
    code_analysis_details = ft.TextField(
        label="Детальный анализ", 
        multiline=True, 
        min_lines=20, 
        max_lines=40, 
        width=1000, 
        filled=True, 
        border_radius=10,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color=get_input_text_color(),
        label_style=ft.TextStyle(color=get_label_color())
    )
    
    def code_analysis_run(e):
        url = code_analysis_url_input.value.strip()
        if not url.startswith('http'):
            code_analysis_status.value = "Введите корректный URL!"
            code_analysis_status.visible = True
            page.update()
            return
        
        code_analysis_progress.value = 0.0
        code_analysis_status.value = "Анализируем код..."
        code_analysis_status.visible = True
        code_analysis_summary.value = ""
        code_analysis_details.value = ""
        page.update()
        
        def worker():
            try:
                # Получаем HTML страницы
                response = requests.get(url, timeout=15, verify=not code_analysis_ssl_checkbox.value)
                if response.status_code != 200:
                    code_analysis_status.value = f"Ошибка загрузки страницы: {response.status_code}"
                    page.update()
                    return
                
                # Анализируем код
                analysis = analyze_code_content(response.text, url)
                
                # Формируем общую сводку
                summary_parts = []
                summary_parts.append("🔍 ДЕТАЛЬНЫЙ АНАЛИЗ КОДА СТРАНИЦЫ")
                summary_parts.append("=" * 60)
                summary_parts.append(f"URL: {analysis['url']}")
                summary_parts.append("")
                
                summary_parts.append("📊 ОБЩАЯ СТАТИСТИКА:")
                summary_parts.append(f"• Оценка качества кода: {analysis['quality_score']}/100")
                summary_parts.append(f"• Всего ошибок: {analysis['total_errors']}")
                summary_parts.append(f"• Всего предупреждений: {analysis['total_warnings']}")
                summary_parts.append("")
                
                # HTML статистика
                if 'html_stats' in analysis:
                    html_stats = analysis['html_stats']
                    summary_parts.append("🌐 HTML СТАТИСТИКА:")
                    summary_parts.append(f"• Всего тегов: {html_stats.get('total_tags', 0)}")
                    summary_parts.append(f"• Уникальных тегов: {html_stats.get('unique_tags', 0)}")
                    summary_parts.append(f"• Всего ссылок: {html_stats.get('total_links', 0)}")
                    summary_parts.append(f"• Внешних ссылок: {html_stats.get('external_links', 0)}")
                    summary_parts.append(f"• Внутренних ссылок: {html_stats.get('internal_links', 0)}")
                    summary_parts.append(f"• Всего изображений: {html_stats.get('total_images', 0)}")
                    summary_parts.append(f"• Покрытие alt: {html_stats.get('alt_coverage', 0):.1f}%")
                    if 'title_length' in html_stats:
                        summary_parts.append(f"• Длина title: {html_stats['title_length']} символов")
                    summary_parts.append("")
                
                # CSS статистика
                if 'css_stats' in analysis:
                    css_stats = analysis['css_stats']
                    summary_parts.append("🎨 CSS СТАТИСТИКА:")
                    summary_parts.append(f"• Блоков стилей: {css_stats.get('style_blocks', 0)}")
                    summary_parts.append(f"• Встроенных стилей: {css_stats.get('inline_styles', 0)}")
                    summary_parts.append(f"• Внешних CSS файлов: {css_stats.get('external_css', 0)}")
                    summary_parts.append(f"• CSS правил: {css_stats.get('css_rules', 0)}")
                    summary_parts.append(f"• Медиа-запросов: {css_stats.get('media_queries', 0)}")
                    summary_parts.append(f"• Анимаций: {css_stats.get('animations', 0)}")
                    summary_parts.append(f"• Переходов: {css_stats.get('transitions', 0)}")
                    summary_parts.append("")
                
                # JavaScript статистика
                if 'js_stats' in analysis:
                    js_stats = analysis['js_stats']
                    summary_parts.append("⚡ JAVASCRIPT СТАТИСТИКА:")
                    summary_parts.append(f"• Блоков скриптов: {js_stats.get('script_blocks', 0)}")
                    summary_parts.append(f"• Встроенных скриптов: {js_stats.get('inline_scripts', 0)}")
                    summary_parts.append(f"• Внешних JS файлов: {js_stats.get('external_scripts', 0)}")
                    summary_parts.append(f"• Функций: {js_stats.get('functions', 0)}")
                    summary_parts.append(f"• Стрелочных функций: {js_stats.get('arrow_functions', 0)}")
                    summary_parts.append(f"• Переменных (var): {js_stats.get('var_declarations', 0)}")
                    summary_parts.append(f"• Переменных (let): {js_stats.get('let_declarations', 0)}")
                    summary_parts.append(f"• Констант (const): {js_stats.get('const_declarations', 0)}")
                    summary_parts.append(f"• Console.log: {js_stats.get('console_logs', 0)}")
                    summary_parts.append("")
                
                # PHP статистика
                if 'php_stats' in analysis:
                    php_stats = analysis['php_stats']
                    summary_parts.append("🐘 PHP СТАТИСТИКА:")
                    summary_parts.append(f"• PHP блоков: {php_stats.get('php_blocks', 0)}")
                    summary_parts.append(f"• PHP строк: {php_stats.get('total_php_lines', 0)}")
                    summary_parts.append("")
                
                # SEO статистика
                if 'seo_stats' in analysis:
                    seo_stats = analysis['seo_stats']
                    summary_parts.append("🔍 SEO СТАТИСТИКА:")
                    if 'charset' in seo_stats:
                        summary_parts.append(f"• Кодировка: {seo_stats['charset']}")
                    if 'description_length' in seo_stats:
                        summary_parts.append(f"• Meta description: {seo_stats['description_length']} символов")
                    summary_parts.append(f"• Open Graph тегов: {seo_stats.get('og_tags', 0)}")
                    summary_parts.append(f"• Twitter Cards: {seo_stats.get('twitter_tags', 0)}")
                    if 'canonical' in seo_stats:
                        summary_parts.append("• Canonical URL: ✅")
                    summary_parts.append("")
                
                # Производительность
                if 'performance_stats' in analysis:
                    perf_stats = analysis['performance_stats']
                    summary_parts.append("⚡ ПРОИЗВОДИТЕЛЬНОСТЬ:")
                    summary_parts.append(f"• HTML размер: {perf_stats.get('html_size_kb', 0):.1f} KB")
                    summary_parts.append(f"• CSS размер: {perf_stats.get('css_size_kb', 0):.1f} KB")
                    summary_parts.append(f"• JS размер: {perf_stats.get('js_size_kb', 0):.1f} KB")
                    summary_parts.append("")
                
                # Оценка качества
                if analysis['quality_score'] >= 90:
                    summary_parts.append("🟢 ОТЛИЧНОЕ КАЧЕСТВО КОДА")
                elif analysis['quality_score'] >= 70:
                    summary_parts.append("🟡 ХОРОШЕЕ КАЧЕСТВО КОДА")
                elif analysis['quality_score'] >= 50:
                    summary_parts.append("🟠 СРЕДНЕЕ КАЧЕСТВО КОДА")
                else:
                    summary_parts.append("🔴 ПЛОХОЕ КАЧЕСТВО КОДА")
                summary_parts.append("")
                
                if analysis['positives']:
                    summary_parts.append("✅ ПОЛОЖИТЕЛЬНЫЕ МОМЕНТЫ:")
                    for positive in analysis['positives']:
                        summary_parts.append(f"• {positive}")
                    summary_parts.append("")
                
                # Формируем детальный анализ
                details_parts = []
                details_parts.append("🔍 ДЕТАЛЬНЫЙ АНАЛИЗ КОДА")
                details_parts.append("=" * 60)
                details_parts.append("")
                
                # HTML анализ
                details_parts.append("🌐 HTML АНАЛИЗ:")
                details_parts.append("-" * 30)
                if analysis['html_errors']:
                    details_parts.append("❌ ОШИБКИ:")
                    for error in analysis['html_errors']:
                        details_parts.append(f"  {error}")
                    details_parts.append("")
                
                if analysis['html_warnings']:
                    details_parts.append("⚠️ ПРЕДУПРЕЖДЕНИЯ:")
                    for warning in analysis['html_warnings']:
                        details_parts.append(f"  {warning}")
                    details_parts.append("")
                
                # Топ тегов
                if 'html_stats' in analysis and 'top_tags' in analysis['html_stats']:
                    details_parts.append("📊 ТОП-10 ТЕГОВ:")
                    for tag, count in analysis['html_stats']['top_tags']:
                        details_parts.append(f"  • <{tag}>: {count} раз")
                    details_parts.append("")
                
                # CSS анализ
                details_parts.append("🎨 CSS АНАЛИЗ:")
                details_parts.append("-" * 30)
                if analysis['css_errors']:
                    details_parts.append("❌ ОШИБКИ:")
                    for error in analysis['css_errors']:
                        details_parts.append(f"  {error}")
                    details_parts.append("")
                
                if analysis['css_warnings']:
                    details_parts.append("⚠️ ПРЕДУПРЕЖДЕНИЯ:")
                    for warning in analysis['css_warnings']:
                        details_parts.append(f"  {warning}")
                    details_parts.append("")
                
                # JavaScript анализ
                details_parts.append("⚡ JAVASCRIPT АНАЛИЗ:")
                details_parts.append("-" * 30)
                if analysis['js_errors']:
                    details_parts.append("❌ ОШИБКИ:")
                    for error in analysis['js_errors']:
                        details_parts.append(f"  {error}")
                    details_parts.append("")
                
                if analysis['js_warnings']:
                    details_parts.append("⚠️ ПРЕДУПРЕЖДЕНИЯ:")
                    for warning in analysis['js_warnings']:
                        details_parts.append(f"  {warning}")
                    details_parts.append("")
                
                # PHP анализ
                details_parts.append("🐘 PHP АНАЛИЗ:")
                details_parts.append("-" * 30)
                if analysis['php_errors']:
                    details_parts.append("❌ ОШИБКИ:")
                    for error in analysis['php_errors']:
                        details_parts.append(f"  {error}")
                    details_parts.append("")
                
                if analysis['php_warnings']:
                    details_parts.append("⚠️ ПРЕДУПРЕЖДЕНИЯ:")
                    for warning in analysis['php_warnings']:
                        details_parts.append(f"  {warning}")
                    details_parts.append("")
                
                # SEO анализ
                details_parts.append("🔍 SEO АНАЛИЗ:")
                details_parts.append("-" * 30)
                if analysis.get('seo_errors'):
                    details_parts.append("❌ ОШИБКИ:")
                    for error in analysis['seo_errors']:
                        details_parts.append(f"  {error}")
                    details_parts.append("")
                
                if analysis.get('seo_warnings'):
                    details_parts.append("⚠️ ПРЕДУПРЕЖДЕНИЯ:")
                    for warning in analysis['seo_warnings']:
                        details_parts.append(f"  {warning}")
                    details_parts.append("")
                
                # Рекомендации по улучшению
                details_parts.append("💡 РЕКОМЕНДАЦИИ ПО УЛУЧШЕНИЮ:")
                details_parts.append("-" * 40)
                
                if analysis['total_errors'] > 0:
                    details_parts.append("• Исправьте все критические ошибки в первую очередь")
                
                if analysis['html_warnings']:
                    details_parts.append("• Добавьте alt атрибуты ко всем изображениям")
                    details_parts.append("• Проверьте и исправьте незакрытые HTML теги")
                    details_parts.append("• Укажите атрибут lang в теге <html>")
                
                if analysis['css_warnings']:
                    details_parts.append("• Проверьте синтаксис CSS на наличие незакрытых скобок")
                    details_parts.append("• Добавьте точки с запятой где необходимо")
                    details_parts.append("• Рассмотрите возможность вынесения CSS в отдельные файлы")
                
                if analysis['js_warnings']:
                    details_parts.append("• Проверьте синтаксис JavaScript")
                    details_parts.append("• Добавьте точки с запятой в конце выражений")
                    details_parts.append("• Уберите console.log из продакшен кода")
                    details_parts.append("• Рассмотрите возможность вынесения JS в отдельные файлы")
                
                if analysis['php_warnings']:
                    details_parts.append("• Проверьте синтаксис PHP кода")
                    details_parts.append("• Убедитесь в правильности закрытия всех скобок")
                
                if analysis.get('seo_warnings'):
                    details_parts.append("• Добавьте meta description")
                    details_parts.append("• Укажите canonical URL")
                    details_parts.append("• Добавьте Open Graph теги")
                    details_parts.append("• Укажите viewport meta тег")
                
                # Рекомендации по производительности
                if 'performance_stats' in analysis:
                    perf_stats = analysis['performance_stats']
                    if perf_stats.get('html_size_kb', 0) > 100:
                        details_parts.append("• Оптимизируйте размер HTML (сейчас > 100KB)")
                    if perf_stats.get('css_size_kb', 0) > 50:
                        details_parts.append("• Оптимизируйте размер CSS (сейчас > 50KB)")
                    if perf_stats.get('js_size_kb', 0) > 100:
                        details_parts.append("• Оптимизируйте размер JavaScript (сейчас > 100KB)")
                
                if analysis['quality_score'] < 70:
                    details_parts.append("• Рассмотрите возможность рефакторинга кода")
                    details_parts.append("• Используйте валидаторы HTML/CSS/JS для проверки")
                    details_parts.append("• Внедрите линтеры для автоматической проверки кода")
                
                # Обновляем интерфейс
                code_analysis_summary.value = "\n".join(summary_parts)
                code_analysis_details.value = "\n".join(details_parts)
                code_analysis_status.value = "Анализ кода завершен!"
                code_analysis_progress.value = 1.0
                
            except Exception as ex:
                code_analysis_status.value = f"Ошибка анализа: {str(ex)}"
                code_analysis_progress.value = 0.0
            
            page.update()
        
        threading.Thread(target=worker).start()
    
    code_analysis_run_btn = ft.ElevatedButton(
        "Анализировать код (УЛУЧШЕННЫЙ)", 
        icon=ft.Icons.PLAY_ARROW, 
        on_click=code_analysis_run,
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=10),
            elevation=5
        )
    )
    
    code_analysis_content.content = ft.Column([
        ft.Text("Детальный анализ кода страницы", size=24, weight=ft.FontWeight.BOLD, color=get_text_color()),
        ft.Text("Введите URL страницы для детального анализа HTML, CSS, JavaScript, PHP и SEO", size=16, color=get_secondary_text_color()),
        ft.Row([code_analysis_url_input, code_analysis_ssl_checkbox, code_analysis_run_btn], spacing=10),
        code_analysis_progress,
        code_analysis_status,
        ft.Column([
            code_analysis_summary,
            ft.Container(height=20),
            code_analysis_details
        ], expand=True)
    ], expand=True)
    
    # --- Редиректы ---
    redirects_input = ft.TextField(
        label="Список URL (по одному в строке)", 
        multiline=True, 
        min_lines=5, 
        width=600, 
        filled=True, 
        border_radius=10,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color=get_input_text_color(),
        label_style=ft.TextStyle(color=get_label_color())
    )
    redirects_ssl_checkbox = ft.Checkbox(label="Игнорировать SSL", value=True)
    redirects_progress = ft.ProgressBar(width=400, color="#F2E307", bgcolor="#394459", value=0.0, height=10, border_radius=20)
    redirects_table = ft.DataTable(
        columns=[
            ft.DataColumn(label=ft.Text("Исходный URL", color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("Редирект", color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("Конечный URL", color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("HTTP", color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("OK", color=get_input_text_color())),
        ],
        rows=[],
        horizontal_lines=ft.BorderSide(1, "#394459"),
        vertical_lines=ft.BorderSide(1, "#394459"),
        bgcolor="#F2F2F2",
        border=ft.border.all(1, "#394459"),
        border_radius=10,
    )
    redirects_export_btn = ft.ElevatedButton(
        "Экспорт в Excel", 
        icon=ft.Icons.DOWNLOAD, 
        visible=False,
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=8),
            elevation=3
        )
    )
    redirects_status = ft.Text(visible=False)

    def redirects_update(done, total):
        redirects_progress.value = min(1.0, done / max(total, 1))
        page.update()

    def redirects_done(results):
        redirects_table.rows = [
            ft.DataRow(cells=[
                ft.DataCell(ft.Text(r['Исходный URL'], color=get_input_text_color())),
                ft.DataCell(ft.Text(r['Редирект'], color=get_input_text_color())),
                ft.DataCell(ft.Text(r['Конечный URL'], color=get_input_text_color())),
                ft.DataCell(ft.Text(str(r['HTTP']), color=get_input_text_color())),
                ft.DataCell(ft.Text('🟢' if r['OK'] else '🔴', color=get_input_text_color()))
            ]) for r in results
        ]
        redirects_export_btn.visible = True
        redirects_status.value = f"Проверено: {len(results)} URL"
        redirects_status.visible = True
        # Сохраняем в Excel
        import pandas as pd
        df = pd.DataFrame(results)
        fname = f"reports/redirects_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        df.to_excel(fname, index=False)
        
        # Показываем уведомление о сохранении
        page.snack_bar = ft.SnackBar(content=ft.Text(f"✅ Excel отчет сохранен: {os.path.basename(fname)}"))
        page.snack_bar.open = True
        
        # Обновляем список экспортов если открыта страница экспорта
        if exports_content.visible:
            refresh_exports_list()
        
        page.update()

    def redirects_run(e):
        urls = [u.strip() for u in redirects_input.value.splitlines() if u.strip()]
        if not urls:
            redirects_status.value = "Введите хотя бы один URL!"
            redirects_status.visible = True
            page.update()
            return
        redirects_progress.value = 0.0
        redirects_table.rows = []
        redirects_export_btn.visible = False
        redirects_status.visible = False
        page.update()
        threading.Thread(target=lambda: check_redirects(urls, redirects_ssl_checkbox.value, redirects_update, redirects_done)).start()

    redirects_run_btn = ft.ElevatedButton(
        "Проверить редиректы", 
        icon=ft.Icons.PLAY_ARROW, 
        on_click=redirects_run,
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=10),
            elevation=5
        )
    )
    redirects_export_btn.on_click = lambda e: None  # Уже экспортируется автоматически
    redirects_content.content = ft.Column([
        ft.Text("Проверка редиректов", size=20, weight=ft.FontWeight.BOLD, color=get_text_color()),
        ft.Row([redirects_input, redirects_ssl_checkbox, redirects_run_btn]),
        redirects_progress,
        ft.Container(
            ft.Row([
                ft.Container(redirects_table, width=2000, expand=True)
            ], scroll=ft.ScrollMode.ALWAYS),
            expand=True,
            bgcolor="#F2F2F2",
            border_radius=10,
            padding=5,
            alignment=ft.alignment.center,
            border=ft.border.all(1, "#394459"),
        ),
        redirects_export_btn,
        redirects_status
    ], expand=True)

    # --- Анализ конкурентов ---
    competitors_input = ft.TextField(
        label="URL конкурентов (до 5, через запятую)", 
        width=600, 
        filled=True, 
        border_radius=10,
        bgcolor="#F2F2F2",
        border_color="#394459",
        focused_border_color="#F2E307",
        color=get_input_text_color(),
        label_style=ft.TextStyle(color=get_label_color())
    )
    competitors_btn = ft.ElevatedButton(
        "Сравнить", 
        icon=ft.Icons.SEARCH,
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=10),
            elevation=5
        )
    )
    competitors_status = ft.Text(visible=False)
    competitors_table = ft.DataTable(
        columns=[
            ft.DataColumn(label=ft.Text("URL", size=12, text_align=ft.TextAlign.CENTER, color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("Скорость\nзагрузки, сек", size=12, text_align=ft.TextAlign.CENTER, color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("H1/\nH2", size=12, text_align=ft.TextAlign.CENTER, color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("Title", size=12, text_align=ft.TextAlign.CENTER, color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("Description", size=12, text_align=ft.TextAlign.CENTER, color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("URL\nв sitemap", size=12, text_align=ft.TextAlign.CENTER, color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("Top-3\nключевых слова", size=12, text_align=ft.TextAlign.CENTER, color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("Внутр.\nссылки", size=12, text_align=ft.TextAlign.CENTER, color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("Внешн.\nссылки", size=12, text_align=ft.TextAlign.CENTER, color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("Изобр.", size=12, text_align=ft.TextAlign.CENTER, color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("Символов\nна стр.", size=12, text_align=ft.TextAlign.CENTER, color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("Meta\nkeywords", size=12, text_align=ft.TextAlign.CENTER, color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("Canonical", size=12, text_align=ft.TextAlign.CENTER, color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("OpenGraph", size=12, text_align=ft.TextAlign.CENTER, color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("Twitter\nCard", size=12, text_align=ft.TextAlign.CENTER, color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("JSON-LD", size=12, text_align=ft.TextAlign.CENTER, color=get_input_text_color())),
            ft.DataColumn(label=ft.Text("Ошибка", size=12, text_align=ft.TextAlign.CENTER, color=get_input_text_color())),
        ],
        rows=[],
        horizontal_lines=ft.BorderSide(1, "#394459"),
        vertical_lines=ft.BorderSide(1, "#394459"),
        data_row_min_height=32,
        data_row_max_height=160,
        bgcolor="#F2F2F2",
        border=ft.border.all(1, "#394459"),
        border_radius=10,
    )
    competitors_ideas = ft.Text(visible=False)

    def analyze_competitors(e):
        urls = [u.strip() for u in competitors_input.value.split(',') if u.strip()][:5]
        if not urls:
            competitors_status.value = "Введите хотя бы один URL конкурента!"
            competitors_status.visible = True
            page.update()
            return
        competitors_status.value = "Анализ конкурентов..."
        competitors_status.visible = True
        competitors_table.rows = []
        competitors_ideas.value = ""
        competitors_ideas.visible = False
        page.update()
        import threading
        def worker():
            import requests
            from bs4 import BeautifulSoup
            import xml.etree.ElementTree as ET
            results = []
            ideas = []
            all_keywords = set()
            for url in urls:
                try:
                    t0 = time.time()
                    r = requests.get(url, timeout=15, verify=False)
                    load_time = time.time() - t0
                    soup = BeautifulSoup(r.text, 'html.parser')
                    h1 = len(soup.find_all('h1'))
                    h2 = len(soup.find_all('h2'))
                    title = soup.title.string.strip() if soup.title and soup.title.string else ''
                    if len(title) > 60:
                        title = title[:57] + '...'
                    desc = ''
                    meta_desc = soup.find('meta', attrs={'name': 'description'})
                    if meta_desc:
                        desc = meta_desc.get('content', '')
                    if len(desc) > 120:
                        desc = desc[:117] + '...'
                    # sitemap
                    sitemap_url = url.rstrip('/') + '/sitemap.xml'
                    try:
                        r_s = requests.get(sitemap_url, timeout=10, verify=False)
                        if r_s.status_code == 200:
                            root = ET.fromstring(r_s.text)
                            sitemap_count = len([u for u in root.iter('{http://www.sitemaps.org/schemas/sitemap/0.9}url')])
                        else:
                            sitemap_count = 0
                    except Exception as ex_s:
                        sitemap_count = 0
                    # keywords
                    text = soup.get_text(separator=' ', strip=True).lower()
                    words = re.findall(r'\w+', text)
                    stop_words = {'и','в','на','не','с','а','о','для','по','из','к','у','от','но','как','что','это','то','или','за','при'}
                    word_freq = {}
                    for word in words:
                        if word not in stop_words and len(word) > 3:
                            word_freq[word] = word_freq.get(word, 0) + 1
                    keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:3]
                    all_keywords.update([k[0] for k in keywords])
                    keywords_str = '\n'.join([k[0] for k in keywords])
                    # Внутренние/внешние ссылки
                    domain = url.split('//')[-1].split('/')[0]
                    a_tags = soup.find_all('a', href=True)
                    internal_links = [a for a in a_tags if domain in a['href'] or a['href'].startswith('/')]
                    external_links = [a for a in a_tags if domain not in a['href'] and a['href'].startswith('http')]
                    # Изображения
                    img_count = len(soup.find_all('img'))
                    # Символы на странице
                    char_count = len(text)
                    # meta keywords
                    meta_keywords = soup.find('meta', attrs={'name': 'keywords'})
                    meta_keywords_count = len(meta_keywords.get('content', '').split(',')) if meta_keywords and meta_keywords.get('content') else 0
                    # canonical
                    canonicals = soup.find_all('link', rel='canonical')
                    canonical_count = len(canonicals)
                    # OpenGraph
                    og_tags = [m for m in soup.find_all('meta') if m.get('property', '').startswith('og:')]
                    og_count = len(og_tags)
                    # Twitter Card
                    tw_tags = [m for m in soup.find_all('meta') if m.get('name', '').startswith('twitter:')]
                    tw_count = len(tw_tags)
                    # JSON-LD
                    jsonld_blocks = soup.find_all('script', {'type': 'application/ld+json'})
                    jsonld_count = len(jsonld_blocks)
                    results.append([
                        url,
                        f"{load_time:.2f}",
                        f"{h1}/{h2}",
                        title,
                        desc,
                        str(sitemap_count),
                        keywords_str,
                        len(internal_links),
                        len(external_links),
                        img_count,
                        char_count,
                        meta_keywords_count,
                        canonical_count,
                        og_count,
                        tw_count,
                        jsonld_count,
                        ''
                    ])
                except Exception as ex:
                    results.append([url, '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', f'{ex}'])
            # Идеи для оптимизации
            if results:
                min_speed = min([float(r[1]) for r in results if r[1] not in ('Ошибка','') and not r[1].startswith('Ошибка:')], default=None)
                max_h1 = max([int(r[2].split('/')[0]) for r in results if r[2] and r[2] != 'Ошибка' and not r[2].startswith('Ошибка:')], default=None)
                max_h2 = max([int(r[2].split('/')[1]) for r in results if r[2] and r[2] != 'Ошибка' and not r[2].startswith('Ошибка:')], default=None)
                max_sitemap = max([int(r[5]) for r in results if r[5].isdigit()], default=None)
                if min_speed is not None:
                    ideas.append(f"Минимальная скорость загрузки у конкурентов: {min_speed:.2f} сек. Старайтесь быть не хуже!")
                if max_h1 is not None and max_h1 > 1:
                    ideas.append(f"У некоторых конкурентов больше 1 H1 — это плохо. Используйте только один H1!")
                if max_h2 is not None and max_h2 < 2:
                    ideas.append(f"У конкурентов мало H2 — добавьте больше подзаголовков для структуры.")
                if max_sitemap is not None:
                    ideas.append(f"У конкурентов в sitemap до {max_sitemap} URL. Проверьте полноту своей карты сайта.")
                if all_keywords:
                    ideas.append(f"Популярные ключевые слова у конкурентов: {', '.join(list(all_keywords)[:10])}")
            competitors_table.rows = [ft.DataRow(cells=[ft.DataCell(ft.Text(str(cell), color=get_input_text_color())) for cell in row]) for row in results]
            competitors_status.value = f"Анализ завершён. Конкурентов: {len(results)}"
            competitors_status.visible = True
            competitors_ideas.value = '\n'.join(ideas)
            competitors_ideas.visible = True if ideas else False
            page.update()
        threading.Thread(target=worker).start()
    competitors_btn.on_click = analyze_competitors
    competitors_content.content = ft.Column([
        ft.Text("Анализ конкурентов", size=24, weight=ft.FontWeight.BOLD, color=get_text_color()),
        ft.Text("Введите до 5 URL конкурентов через запятую. Будет выполнено сравнение по скорости, структуре, мета-тегам и ключевым словам.", color=get_secondary_text_color()),
        ft.Row([competitors_input, competitors_btn]),
        competitors_status,
        ft.Container(
            ft.Row([
                ft.Container(competitors_table, width=2400, expand=True)
            ], scroll=ft.ScrollMode.ALWAYS),
            expand=True,
            bgcolor="#F2F2F2",
            border_radius=10,
            padding=5,
            alignment=ft.alignment.center,
            border=ft.border.all(1, "#394459"),
        ),
        competitors_ideas
    ], expand=True)

    # --- Функции для работы с экспортами ---
    def get_file_icon(file_path):
        """Возвращает иконку в зависимости от типа файла."""
        if file_path.lower().endswith('.xlsx'):
            return ft.Icon(ft.Icons.TABLE_CHART, color="#217346", size=32)
        elif file_path.lower().endswith('.docx'):
            return ft.Icon(ft.Icons.DESCRIPTION, color="#2B579A", size=32)
        else:
            return ft.Icon(ft.Icons.INSERT_DRIVE_FILE, color="#666666", size=32)
    
    def get_file_size(file_path):
        """Возвращает размер файла в читаемом формате."""
        try:
            size_bytes = os.path.getsize(file_path)
            if size_bytes < 1024:
                return f"{size_bytes} Б"
            elif size_bytes < 1024 * 1024:
                return f"{size_bytes / 1024:.1f} КБ"
            else:
                return f"{size_bytes / (1024 * 1024):.1f} МБ"
        except:
            return "Неизвестно"
    
    def get_file_date(file_path):
        """Возвращает дату создания файла."""
        try:
            timestamp = os.path.getctime(file_path)
            return datetime.fromtimestamp(timestamp).strftime('%d.%m.%Y %H:%M')
        except:
            return "Неизвестно"
    
    def open_file(file_path):
        """Открывает файл в системе по умолчанию."""
        try:
            import subprocess
            import platform
            if platform.system() == 'Windows':
                os.startfile(file_path)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.run(['open', file_path])
            else:  # Linux
                subprocess.run(['xdg-open', file_path])
            page.snack_bar = ft.SnackBar(content=ft.Text(f"Файл открыт: {os.path.basename(file_path)}"))
            page.snack_bar.open = True
        except Exception as e:
            page.snack_bar = ft.SnackBar(content=ft.Text(f"Ошибка открытия файла: {str(e)}"))
            page.snack_bar.open = True
        page.update()
    
    def refresh_exports_list():
        """Обновляет список экспортированных файлов."""
        try:
            exports_list.controls.clear()
            
            if not os.path.exists(REPORT_DIR):
                exports_list.controls.append(
                    ft.Container(
                        content=ft.Text("Папка reports не найдена", color="#666666"),
                        padding=20,
                        alignment=ft.alignment.center
                    )
                )
                page.update()
                return
            
            files = []
            for filename in os.listdir(REPORT_DIR):
                if filename.endswith(('.xlsx', '.docx')):
                    file_path = os.path.join(REPORT_DIR, filename)
                    files.append({
                        'name': filename,
                        'path': file_path,
                        'size': get_file_size(file_path),
                        'date': get_file_date(file_path),
                        'type': 'Excel' if filename.endswith('.xlsx') else 'Word'
                    })
            
            # Сортируем файлы по дате создания (новые сверху)
            files.sort(key=lambda x: os.path.getctime(x['path']), reverse=True)
            
            if not files:
                exports_list.controls.append(
                    ft.Container(
                        content=ft.Text("Нет экспортированных файлов", color="#666666"),
                        padding=20,
                        alignment=ft.alignment.center
                    )
                )
            else:
                for file_info in files:
                    file_card = ft.Card(
                        content=ft.Container(
                            content=ft.Row([
                                get_file_icon(file_info['path']),
                                ft.Column([
                                    ft.Text(
                                        file_info['name'], 
                                        weight=ft.FontWeight.BOLD,
                                        color=get_text_color(),
                                        size=16
                                    ),
                                    ft.Text(
                                        f"Тип: {file_info['type']} • Размер: {file_info['size']} • Создан: {file_info['date']}",
                                        color=get_text_color(),
                                        size=12
                                    )
                                ], expand=True),
                                ft.ElevatedButton(
                                    "Открыть",
                                    icon=ft.Icons.OPEN_IN_NEW,
                                    style=ft.ButtonStyle(
                                        bgcolor="#F2E307",
                                        color="#394459",
                                        shape=ft.RoundedRectangleBorder(radius=8),
                                        elevation=3
                                    ),
                                    on_click=lambda e, path=file_info['path']: open_file(path)
                                )
                            ], spacing=15),
                            padding=20
                        ),
                        elevation=5,
                        margin=ft.Margin(0, 0, 0, 10)
                    )
                    exports_list.controls.append(file_card)
            
            page.update()
        except Exception as e:
            exports_list.controls.clear()
            exports_list.controls.append(
                ft.Container(
                    content=ft.Text(f"Ошибка загрузки файлов: {str(e)}", color="#FF6B6B"),
                    padding=20,
                    alignment=ft.alignment.center
                )
            )
            page.update()
    
    # --- Интерфейс страницы экспорта ---
    exports_list = ft.Column([], scroll=ft.ScrollMode.AUTO, expand=True)
    refresh_exports_btn = ft.ElevatedButton(
        "🔄 Обновить список",
        icon=ft.Icons.REFRESH,
        style=ft.ButtonStyle(
            bgcolor="#F2E307",
            color="#394459",
            shape=ft.RoundedRectangleBorder(radius=10),
            elevation=3
        ),
        on_click=lambda e: refresh_exports_list()
    )
    
    exports_content.content = ft.Column([
        ft.Text("📁 Экспортированные файлы", size=24, weight=ft.FontWeight.BOLD, color=get_text_color()),
        ft.Text("Здесь отображаются все экспортированные отчеты в форматах Excel и Word", color=get_secondary_text_color()),
        ft.Row([refresh_exports_btn], alignment=ft.MainAxisAlignment.END),
        ft.Container(
            exports_list,
            expand=True,
            bgcolor="#F2F2F2",
            border_radius=10,
            padding=20,
            border=ft.border.all(1, "#394459"),
        )
    ], expand=True)

    # --- Layout ---
    # Инициализируем активную страницу (главная)
    switch_page(0)
    
    page.add(
        ft.Column([
            navigation_bar,
            ft.Stack([
                main_content,
                links_check_content,
                parser_content,
                text_analysis_content,
                code_analysis_content,
                redirects_content,
                competitors_content,
                exports_content,
                serp_tracker_content
            ], expand=True)
        ], expand=True)
    )

# --- ДОБАВИТЬ ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ---
def get_core_web_vitals(driver):
    """Собирает базовые метрики Core Web Vitals через JS."""
    try:
        # LCP и CLS можно получить через PerformanceObserver, FID — только в реальном взаимодействии, но можно эмулировать
        script = '''
        var resolve = arguments[0];
        let lcp = 0, cls = 0;
        try {
            new PerformanceObserver((entryList) => {
                for (const entry of entryList.getEntries()) {
                    if (entry.entryType === 'largest-contentful-paint') {
                        lcp = entry.renderTime || entry.loadTime || entry.startTime;
                    }
                    if (entry.entryType === 'layout-shift' && !entry.hadRecentInput) {
                        cls += entry.value;
                    }
                }
            }).observe({type: 'largest-contentful-paint', buffered: true});
            new PerformanceObserver((entryList) => {
                for (const entry of entryList.getEntries()) {
                    if (entry.entryType === 'layout-shift' && !entry.hadRecentInput) {
                        cls += entry.value;
                    }
                }
            }).observe({type: 'layout-shift', buffered: true});
        } catch (e) {
            console.log('PerformanceObserver error:', e);
        }
        setTimeout(() => {
            resolve({lcp, cls});
        }, 3500);
        '''
        result = driver.execute_async_script(script)
        lcp = result.get('lcp', 0)
        cls = result.get('cls', 0)
        # FID эмулировать сложно, но можно поставить 0 (или добавить позже через Lighthouse)
        fid = 0
        return lcp, fid, cls
    except Exception as ex:
        log_to_file(f"Ошибка получения Core Web Vitals: {str(ex)}")
        return 0, 0, 0

def get_microdata(driver):
    """Собирает все виды микроразметки: Schema.org, JSON-LD, OpenGraph, Twitter Cards."""
    try:
        html = driver.page_source
        soup = BeautifulSoup(html, 'html.parser')
        # Schema.org (itemscope/itemtype)
        schema_items = []
        for tag in soup.find_all(attrs={"itemscope": True}):
            itemtype = tag.get("itemtype", "")
            schema_items.append(itemtype)
        # JSON-LD
        jsonld_blocks = []
        for script in soup.find_all("script", {"type": "application/ld+json"}):
            try:
                data = json.loads(script.string)
                jsonld_blocks.append(data)
            except Exception:
                pass
        # OpenGraph
        og_tags = {}
        for meta in soup.find_all("meta"):
            prop = meta.get("property", "")
            if prop.startswith("og:"):
                og_tags[prop] = meta.get("content", "")
        # Twitter Cards
        twitter_tags = {}
        for meta in soup.find_all("meta"):
            name = meta.get("name", "")
            if name.startswith("twitter:"):
                twitter_tags[name] = meta.get("content", "")
        return schema_items, jsonld_blocks, og_tags, twitter_tags
    except Exception as ex:
        log_to_file(f"Ошибка получения микроразметки: {str(ex)}")
        return [], [], {}, {}

def process_sitemap_recursively(sitemap_url, ignore_ssl, visited_sitemaps=None, max_depth=3):
    """
    Рекурсивно обрабатывает sitemap index файлы и собирает все URL с метаданными.
    
    Args:
        sitemap_url: URL sitemap файла
        ignore_ssl: Игнорировать SSL ошибки
        visited_sitemaps: Множество уже посещенных sitemap (для предотвращения циклов)
        max_depth: Максимальная глубина рекурсии
    
    Returns:
        dict: Словарь с информацией о найденных URL, их источниках и метаданных
    """
    if visited_sitemaps is None:
        visited_sitemaps = set()
    
    if max_depth <= 0 or sitemap_url in visited_sitemaps:
        return {
            'urls': [], 
            'sources': {}, 
            'metadata': {},  # Метаданные для каждого URL
            'errors': [f"Достигнута максимальная глубина или циклическая ссылка: {sitemap_url}"]
        }
    
    visited_sitemaps.add(sitemap_url)
    result = {'urls': [], 'sources': {}, 'metadata': {}, 'errors': []}
    
    try:
        response = requests.get(sitemap_url, timeout=10, verify=not ignore_ssl)
        if response.status_code != 200:
            result['errors'].append(f"Не удалось загрузить sitemap: {sitemap_url} (статус: {response.status_code})")
            return result
        
        root = ET.fromstring(response.text)
        
        # Проверяем, это sitemap index или обычный sitemap
        if 'sitemapindex' in root.tag:
            # Это sitemap index - обрабатываем рекурсивно
            sitemap_elems = root.findall('.//{http://www.sitemaps.org/schemas/sitemap/0.9}sitemap')
            for sitemap_elem in sitemap_elems:
                sub_sitemap_url = sitemap_elem.findtext('{http://www.sitemaps.org/schemas/sitemap/0.9}loc')
                if sub_sitemap_url:
                    sub_result = process_sitemap_recursively(
                        sub_sitemap_url, 
                        ignore_ssl, 
                        visited_sitemaps.copy(), 
                        max_depth - 1
                    )
                    result['urls'].extend(sub_result['urls'])
                    result['sources'].update(sub_result['sources'])
                    result['metadata'].update(sub_result['metadata'])
                    result['errors'].extend(sub_result['errors'])
        
        elif 'urlset' in root.tag:
            # Это обычный sitemap - извлекаем URL с метаданными
            url_elems = root.findall('.//{http://www.sitemaps.org/schemas/sitemap/0.9}url')
            for url_elem in url_elems:
                url = url_elem.findtext('{http://www.sitemaps.org/schemas/sitemap/0.9}loc')
                if url:
                    result['urls'].append(url)
                    result['sources'][url] = sitemap_url
                    
                    # Извлекаем метаданные
                    lastmod = url_elem.findtext('{http://www.sitemaps.org/schemas/sitemap/0.9}lastmod') or '-'
                    priority = url_elem.findtext('{http://www.sitemaps.org/schemas/sitemap/0.9}priority') or '-'
                    changefreq = url_elem.findtext('{http://www.sitemaps.org/schemas/sitemap/0.9}changefreq') or '-'
                    
                    result['metadata'][url] = {
                        'lastmod': lastmod,
                        'priority': priority,
                        'changefreq': changefreq
                    }
        
        else:
            result['errors'].append(f"Неизвестный формат sitemap: {sitemap_url}")
    
    except ET.ParseError as e:
        result['errors'].append(f"Ошибка парсинга XML в {sitemap_url}: {str(e)}")
    except Exception as e:
        result['errors'].append(f"Ошибка обработки {sitemap_url}: {str(e)}")
    
    return result

if __name__ == "__main__":
    ft.app(target=main)